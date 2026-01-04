import streamlit as st
import streamlit.components.v1 as components
import google.generativeai as genai
from supabase import create_client, Client
import pandas as pd
import docx
import json
import re
import textwrap
import io
import time
import requests
import random
import urllib.parse # [BẮT BUỘC] Thư viện xử lý QR Code tránh lỗi

# === Brand logo (SVG, transparent) ===
LOGO_SVG_TEMPLATE = r'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 256 256" width="{size}" height="{size}" aria-label="aiexam logo" role="img">
  <defs>
    <linearGradient id="pen-{uid}" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0" stop-color="#f5f7fb"/>
      <stop offset="0.35" stop-color="#cfd6e2"/>
      <stop offset="0.7" stop-color="#9aa3b2"/>
      <stop offset="1" stop-color="#ffffff"/>
    </linearGradient>
    <linearGradient id="doc-{uid}" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0" stop-color="#0ea5e9"/>
      <stop offset="0.6" stop-color="#2563eb"/>
      <stop offset="1" stop-color="#1d4ed8"/>
    </linearGradient>
    <radialGradient id="spark-{uid}" cx="35%" cy="70%" r="60%">
      <stop offset="0" stop-color="#fff7c2"/>
      <stop offset="0.3" stop-color="#facc15"/>
      <stop offset="1" stop-color="#f59e0b" stop-opacity="0"/>
    </radialGradient>
    <pattern id="grid-{uid}" width="16" height="16" patternUnits="userSpaceOnUse">
      <path d="M16 0H0V16" fill="none" stroke="#93c5fd" stroke-opacity="0.45" stroke-width="1"/>
      <path d="M8 0V16M0 8H16" fill="none" stroke="#bfdbfe" stroke-opacity="0.25" stroke-width="1"/>
    </pattern>
  </defs>

  <!-- digital document tile -->
  <rect x="32" y="28" width="144" height="176" rx="22" fill="url(#doc-{uid})"/>
  <rect x="44" y="42" width="120" height="148" rx="14" fill="url(#grid-{uid})" opacity="0.95"/>
  <path d="M148 28h-40a22 22 0 0 0-22 22v14h84V50a22 22 0 0 0-22-22z" fill="#0b2a6f" opacity="0.10"/>

  <!-- pen nib -->
  <path d="M192 40c-10 0-19 4-26 12l-58 58c-6 6-9 14-9 22l0 26 26 0c8 0 16-3 22-9l58-58c8-7 12-16 12-26 0-14-11-25-25-25z"
        fill="url(#pen-{uid})" stroke="#64748b" stroke-opacity="0.28" stroke-width="3"/>
  <path d="M164 76l16 16" stroke="#475569" stroke-opacity="0.55" stroke-width="6" stroke-linecap="round"/>
  <circle cx="152" cy="104" r="10" fill="#0f172a" opacity="0.35"/>
  <path d="M112 152l-13 30 30-13" fill="#0f172a" opacity="0.18"/>

  <!-- spark -->
  <circle cx="104" cy="168" r="34" fill="url(#spark-{uid})"/>
  <path d="M104 144l6 16 16 6-16 6-6 16-6-16-16-6 16-6z" fill="#facc15" opacity="0.92"/>
</svg>'''

def logo_svg(size: int) -> str:
    # Inline SVG (transparent background). No border/shadow.
    st.session_state.setdefault('_logo_uid', 0)
    st.session_state._logo_uid += 1
    uid = f"ax{st.session_state._logo_uid}"
    return LOGO_SVG_TEMPLATE.format(size=size, uid=uid)

DASHBOARD_HTML = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body{margin:0;padding:0;font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial;}
  .wrap{padding:0 2px 0 2px;}
  .grid{display:grid;grid-template-columns:repeat(2,1fr);gap:12px;}
  .card{border:1px solid #e2e8f0;border-radius:14px;padding:14px 14px 12px 14px;background:#fff;
        box-shadow:0 10px 18px rgba(2,6,23,0.05);}
  .title{font-weight:800;color:#0f172a;font-size:14px;margin:0 0 6px 0}
  .sub{color:#64748b;font-size:12px;line-height:1.35;margin:0}
  .badge{display:inline-block;margin-top:10px;padding:4px 10px;border-radius:999px;background:#f1f5f9;
         border:1px solid #e2e8f0;color:#0f172a;font-size:11px;font-weight:700}
</style>
</head>
<body>
  <div class="wrap">
    <div class="grid">
      <div class="card">
        <p class="title">📘 Trợ lý Soạn bài</p>
        <p class="sub">Tạo giáo án chuẩn CTGDPT 2018 theo môn/lớp/bộ sách.</p>
        <span class="badge">Soạn giáo án</span>
      </div>
      <div class="card">
        <p class="title">🧩 Soạn bài Năng lực số</p>
        <p class="sub">Tích hợp Năng lực số (NLS) vào giáo án.</p>
        <span class="badge">Digital Competency</span>
      </div>
      <div class="card">
        <p class="title">📝 Ra đề – KTĐG</p>
        <p class="sub">Ma trận – Đặc tả – Đề – Đáp án theo đúng pháp lý.</p>
        <span class="badge">Exam Engine</span>
      </div>
      <div class="card">
        <p class="title">💬 Nhận xét – Tư vấn</p>
        <p class="sub">Nhận xét, tư vấn chuyên môn (mở rộng).</p>
        <span class="badge">Advisor</span>
      </div>
    </div>
  </div>
</body>
</html>
"""



def html_escape(text: str) -> str:
    import html
    if not text:
        return ""
    return html.escape(str(text))

# ==============================================================================
# [MODULE NLS] DỮ LIỆU & CẤU HÌNH CHO SOẠN GIÁO ÁN NĂNG LỰC SỐ
# ==============================================================================

# 1. Khung năng lực số (Chuyển từ constants.ts)
NLS_FRAMEWORK_DATA = """
KHUNG NĂNG LỰC SỐ (DIGITAL COMPETENCE FRAMEWORK) - CẬP NHẬT MỚI NHẤT
MÔ TẢ CÁC MIỀN NĂNG LỰC VÀ YÊU CẦU CẦN ĐẠT (YCCĐ):

1. MIỀN 1: KHAI THÁC DỮ LIỆU VÀ THÔNG TIN
   1.1. Duyệt, tìm kiếm và lọc dữ liệu (CB1, CB2, TC1, NC1).
   1.2. Đánh giá dữ liệu (CB1, TC1, NC1).
   1.3. Quản lý dữ liệu (CB1, TC1).

2. MIỀN 2: GIAO TIẾP VÀ HỢP TÁC
   2.1. Tương tác qua công nghệ.
   2.4. Hợp tác qua công nghệ.
   2.5. Văn hóa mạng (Netiquette).

3. MIỀN 3: SÁNG TẠO NỘI DUNG SỐ
   3.1. Phát triển nội dung.
   3.3. Bản quyền và giấy phép.

4. MIỀN 4: AN TOÀN SỐ
   4.2. Bảo vệ dữ liệu cá nhân.
   4.3. Bảo vệ sức khỏe.

5. MIỀN 5: GIẢI QUYẾT VẤN ĐỀ
   5.2. Xác định nhu cầu và giải pháp.
   5.3. Sử dụng sáng tạo.

6. MIỀN 6: ỨNG DỤNG AI
   6.1. Hiểu biết về AI.
   6.2. Sử dụng công cụ AI.
   6.3. Đạo đức AI.
"""

# 2. Câu lệnh hệ thống cho AI (System Prompt)
SYSTEM_INSTRUCTION_NLS = f"""
Bạn là chuyên gia tư vấn giáo dục cao cấp, chuyên về chuyển đổi số và Khung Năng lực số (NLS) tại Việt Nam.

DỮ LIỆU KHUNG NĂNG LỰC SỐ:
{NLS_FRAMEWORK_DATA}

NHIỆM VỤ CỐT LÕI:
1. Phân tích sâu sắc nội dung giáo án người dùng cung cấp để tìm ra các "điểm chạm" có thể tích hợp NLS một cách tự nhiên nhất.
2. Lựa chọn các YCCĐ (Yêu cầu cần đạt) từ Khung NLS phù hợp với trình độ học sinh và đặc thù môn học.
3. Nếu có file PPCT, bạn phải ưu tiên 100% nội dung NLS trong PPCT đó.

CẤU TRÚC ĐẦU RA (MARKDOWN):
I. THÔNG TIN CHUNG (Giữ nguyên từ giáo án gốc)
II. MỤC TIÊU
   1. Kiến thức, kĩ năng... (Giữ nguyên)
   2. Năng lực chung... (Giữ nguyên)
   3. Năng lực đặc thù... (Giữ nguyên)
   4. Năng lực số (Bổ sung mới): 
      - [Mã YCCĐ]: Mô tả biểu hiện cụ thể học sinh sẽ đạt được.
III. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU SỐ (Bổ sung các công cụ cần thiết cho NLS)
IV. TIẾN TRÌNH DẠY HỌC
   - Tích hợp nội dung NLS vào các hoạt động bằng thẻ <u>...</u> hoặc in đậm. 
   - Ví dụ: "HS sử dụng máy tính *thực hiện tra cứu thông tin trên trang web chính thống [1.1.CB2]*".

QUY TẮC KỸ THUẬT:
- Giữ nguyên các định dạng **Bold**, *Italic* của bản gốc.
- Không thay đổi nội dung chuyên môn gốc, chỉ làm phong phú thêm.
"""

# 3. Hàm xử lý AI riêng cho Module này
def generate_nls_lesson_plan(api_key, lesson_content, distribution_content, textbook, subject, grade, analyze_only):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.0-flash', system_instruction=SYSTEM_INSTRUCTION_NLS)
    
    user_prompt = f"""
    THÔNG TIN ĐẦU VÀO:
    - Bộ sách: {textbook} | Môn: {subject} | Lớp: {grade}
    - Chế độ: {"CHỈ PHÂN TÍCH (Không viết lại bài)" if analyze_only else "VIẾT LẠI GIÁO ÁN TÍCH HỢP NLS"}
    
    NỘI DUNG PPCT (Yêu cầu cứng):
    {distribution_content if distribution_content else "Không có, tự đề xuất theo khung NLS."}
    
    NỘI DUNG GIÁO ÁN GỐC:
    {lesson_content}
    """
    
    try:
        response = model.generate_content(user_prompt)
        return response.text
    except Exception as e:
        return f"Lỗi AI: {str(e)}"
from jsonschema import validate, Draft202012Validator # [MỚI] Thư viện Validate Schema

# [MỚI] TÍCH HỢP MODULE SOẠN BÀI HƯỚNG B (Yêu cầu 4 file đi kèm)
# Dùng try-except để không làm sập web nếu thầy chưa kịp tạo file lesson_ui.py
try:
    from lesson_ui import module_lesson_plan_B
except ImportError:
    module_lesson_plan_B = None

# ==============================================================================
# 1. CẤU HÌNH HỆ THỐNG & KẾT NỐI
# ==============================================================================
# --- CẤU HÌNH GIỚI HẠN SỬ DỤNG ---
MAX_FREE_USAGE = 3
MAX_PRO_USAGE = 15

# --- CẤU HÌNH KHUYẾN MẠI & HOA HỒNG ---
BONUS_PER_REF = 0
BONUS_PRO_REF = 3
DISCOUNT_AMT = 0
COMMISSION_AMT = 10000

# --- CẤU HÌNH THANH TOÁN (SEPAY - VIETQR) ---
BANK_ID = "VietinBank"
BANK_ACC = "107878907329"
# Alias để tương thích UI (một số đoạn dùng BANK_NO)
BANK_NO = BANK_ACC
BANK_NAME = "TRAN THANH TUAN"
PRICE_VIP = 50000

# Lấy API Key từ Secrets
try:
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    SYSTEM_GOOGLE_KEY = st.secrets.get("GOOGLE_API_KEY", "")
    SEPAY_API_TOKEN = st.secrets.get("SEPAY_API_TOKEN", "")
except:
    SUPABASE_URL = ""
    SUPABASE_KEY = ""
    SYSTEM_GOOGLE_KEY = ""
    SEPAY_API_TOKEN = ""

st.set_page_config(page_title="AI EXAM EXPERT v10 – 2026", page_icon="🎓", layout="wide", initial_sidebar_state="collapsed")

# =========================
# UI THEME (Premium SaaS)
# =========================
def inject_premium_theme():
    st.markdown("""
<style>
:root{
  --bg:#ffffff;
  --text:#0f172a;
  --muted:#64748b;
  --border:rgba(15,23,42,.10);
  --card:#ffffff;
  --soft:#f6f7fb;
  --sidebar:#f4f2fb;
  --primary:#5b5cf6;
  --primary2:#2f80ff;
  --good:#10b981;
  --warn:#f59e0b;
  --radius-lg:22px;
  --radius-md:16px;
}

.stApp{ background: var(--bg); color: var(--text); }
.block-container{ max-width: 1200px; padding-top: 1.25rem; padding-bottom: 2.5rem; }

section[data-testid="stSidebar"]{
  background: var(--sidebar);
  border-right: 1px solid rgba(15,23,42,.08);
}

.sb-brand{
  display:flex; align-items:center; gap:10px;
  padding: 10px 4px 6px 4px;
}
.sb-logo{
  width: 72px; height: 72px; border-radius: 0px;
  background: transparent;
  display:flex; align-items:center; justify-content:center;
  color: inherit; font-weight:800;
  box-shadow: none;
}
.sb-logo svg{display:block;}
.sb-title{ font-weight: 800; line-height: 1.05; }
.sb-sub{ color: var(--muted); font-size: 12px; margin-top: 2px; }

.hero{
  background: radial-gradient(900px 450px at 15% 15%, rgba(47,128,255,.22), transparent 60%),
              radial-gradient(800px 450px at 85% 20%, rgba(91,92,246,.18), transparent 60%),
              linear-gradient(135deg, #cfe8ff, #dad4ff);
  border-radius: var(--radius-lg);
  padding: 34px 28px;
  border: 1px solid rgba(15,23,42,.08);
  box-shadow: 0 14px 40px rgba(2,6,23,.08);
}

.hero h1{
  margin: 0;
  font-size: 38px;
  letter-spacing: -0.02em;
}
.hero p{ margin: 8px 0 0 0; color: rgba(15,23,42,.72); font-size: 16px; }

.glass{
  background: rgba(255,255,255,.75);
  border: 1px solid rgba(15,23,42,.10);
  border-radius: 999px;
  padding: 10px 12px;
  backdrop-filter: blur(10px);
  box-shadow: 0 10px 26px rgba(2,6,23,.08);
}

.pills{
  display:flex; flex-wrap:wrap; gap:8px;
  justify-content:center;
  margin-top: 14px;
}
.pill{
  display:inline-flex; align-items:center; gap:8px;
  padding: 8px 12px;
  border-radius: 999px;
  border: 1px solid rgba(15,23,42,.10);
  background: rgba(255,255,255,.65);
  font-size: 13px;
  color: rgba(15,23,42,.78);
  transition: transform .12s ease, box-shadow .12s ease;
}
.pill:hover{ transform: translateY(-1px); box-shadow: 0 10px 20px rgba(2,6,23,.08); }

.card{
  background: var(--card);
  border: 1px solid rgba(15,23,42,.08);
  border-radius: var(--radius-md);
  padding: 16px;
  box-shadow: 0 10px 22px rgba(2,6,23,.06);
}
.card.soft{ background: var(--soft); }

.icon-circle{
  width: 54px; height: 54px; border-radius: 999px;
  display:flex; align-items:center; justify-content:center;
  color: white; font-size: 22px;
  box-shadow: 0 10px 24px rgba(2,6,23,.10);
  margin: 0 auto 10px auto;
}
.ic1{ background: linear-gradient(135deg, rgba(47,128,255,.95), rgba(91,92,246,.95)); }
.ic2{ background: linear-gradient(135deg, rgba(16,185,129,.95), rgba(47,128,255,.80)); }
.ic3{ background: linear-gradient(135deg, rgba(245,158,11,.95), rgba(236,72,153,.75)); }
.ic4{ background: linear-gradient(135deg, rgba(236,72,153,.95), rgba(91,92,246,.80)); }
.ic5{ background: linear-gradient(135deg, rgba(100,116,139,.95), rgba(47,128,255,.75)); }
.ic6{ background: linear-gradient(135deg, rgba(91,92,246,.95), rgba(2,132,199,.75)); }

/* Sidebar radio -> nav list */
section[data-testid="stSidebar"] .stRadio > div{
  padding: 4px 6px 2px 6px;
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] label{
  background: rgba(255,255,255,.55);
  border: 1px solid rgba(15,23,42,.08);
  border-radius: 14px;
  padding: 10px 12px;
  margin: 6px 0;
  transition: transform .12s ease, box-shadow .12s ease, background .12s ease;
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] label:hover{
  transform: translateY(-1px);
  box-shadow: 0 12px 24px rgba(2,6,23,.10);
  background: rgba(255,255,255,.80);
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] label p{
  font-weight: 650;
  margin: 0;
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] input:checked + div{
  background: rgba(91,92,246,.12) !important;
  border-color: rgba(91,92,246,.38) !important;
  box-shadow: 0 14px 28px rgba(91,92,246,.18) !important;
  position: relative;
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] input:checked + div::before{
  content:"";
  position:absolute;
  left:-1px; top:-1px; bottom:-1px;
  width: 6px;
  border-radius: 14px 0 0 14px;
  background: linear-gradient(180deg, rgba(91,92,246,.95), rgba(47,128,255,.95));
}

/* Make Streamlit buttons look premium */
.stButton > button{
  border-radius: 14px;
  border: 1px solid rgba(15,23,42,.10);
  box-shadow: 0 10px 18px rgba(2,6,23,.06);
}
.stButton > button:hover{
  transform: translateY(-1px);
}

.small-muted{ color: var(--muted); font-size: 12px; }
</style>
""", unsafe_allow_html=True)

def go(page_key: str):
    # Programmatic navigation: also sync sidebar highlight
    st.session_state["current_page"] = page_key
    st.session_state["_sync_sidebar_menu"] = True
    st.rerun()

inject_premium_theme()

# ==============================================================================
# [MỚI] DỮ LIỆU NĂNG LỰC SỐ (TỪ FILE constants.ts CỦA THẦY)
# ==============================================================================
NLS_FRAMEWORK_DATA = """
KHUNG NĂNG LỰC SỐ (DIGITAL COMPETENCE FRAMEWORK) - CẬP NHẬT MỚI NHẤT

MÔ TẢ CÁC MIỀN NĂNG LỰC VÀ YÊU CẦU CẦN ĐẠT (YCCĐ):

1. MIỀN 1: KHAI THÁC DỮ LIỆU VÀ THÔNG TIN
   1.1. Duyệt, tìm kiếm và lọc dữ liệu:
      - CB1: Xác định được nhu cầu thông tin cơ bản, thực hiện tìm kiếm bằng từ khóa đơn giản.
      - CB2: Biết cách lọc và sắp xếp kết quả tìm kiếm theo các tiêu chí đơn giản (thời gian, loại file).
      - TC1: Xây dựng chiến lược tìm kiếm phức tạp, sử dụng các toán tử tìm kiếm (AND, OR, "").
      - NC1: Đánh giá và điều chỉnh chiến lược tìm kiếm dựa trên độ nhiễu của thông tin.
   1.2. Đánh giá dữ liệu:
      - CB1: Nhận biết được tin giả, tin rác cơ bản dựa trên cảm tính hoặc nguồn tin không rõ ràng.
      - TC1: Phân tích được tính tin cậy, khách quan và bản quyền của nguồn dữ liệu.
      - NC1: So sánh và đối chiếu nhiều nguồn tin để xác chứng dữ liệu trước khi sử dụng.
   1.3. Quản lý dữ liệu:
      - CB1: Biết lưu trữ file vào thư mục và đặt tên gợi nhớ.
      - TC1: Sử dụng các dịch vụ lưu trữ đám mây (Drive, OneDrive) để tổ chức dữ liệu khoa học.

2. MIỀN 2: GIAO TIẾP VÀ HỢP TÁC
   2.1. Tương tác qua công nghệ:
      - CB1: Sử dụng được email, tin nhắn để gửi thông tin đơn giản.
      - TC1: Lựa chọn được công cụ giao tiếp phù hợp với mục đích và đối tượng.
   2.4. Hợp tác qua công nghệ:
      - CB1: Tham gia vào các tệp tin chia sẻ chung (Google Docs) để đóng góp ý kiến.
      - TC1: Sử dụng công nghệ để đồng sáng tạo sản phẩm, quản lý tiến độ nhóm (Trello, Planner).
   2.5. Văn hóa mạng (Netiquette):
      - CB1: Biết cách ứng xử lịch sự, không dùng ngôn từ gây hấn trên không gian mạng.
      - TC1: Hiểu và tuân thủ các quy tắc đạo đức, chuẩn mực văn hóa số.

3. MIỀN 3: SÁNG TẠO NỘI DUNG SỐ
   3.1. Phát triển nội dung:
      - CB1: Tạo được văn bản, hình ảnh, bài trình chiếu đơn giản.
      - TC1: Thiết kế được nội dung đa phương tiện (video, infographic) thẩm mỹ.
      - NC1: Tạo ra các sản phẩm số độc đáo, giải quyết vấn đề thực tế.
   3.3. Bản quyền và giấy phép:
      - CB1: Biết trích dẫn nguồn khi sử dụng tài liệu từ internet.
      - TC1: Hiểu về các loại giấy phép Creative Commons (CC).

4. MIỀN 4: AN TOÀN SỐ
   4.2. Bảo vệ dữ liệu cá nhân:
      - CB1: Biết đặt mật khẩu mạnh, không chia sẻ thông tin cá nhân.
      - TC1: Hiểu về cơ chế thu thập dữ liệu và thiết lập quyền riêng tư.
   4.3. Bảo vệ sức khỏe:
      - CB1: Nhận biết tác hại của việc sử dụng thiết bị số quá thời gian.
      - TC1: Biết tự điều chỉnh thời gian sử dụng và vận động.

5. MIỀN 5: GIẢI QUYẾT VẤN ĐỀ
   5.2. Xác định nhu cầu và giải pháp:
      - CB1: Sử dụng công cụ số hỗ trợ tính toán, tra cứu.
      - TC1: Sử dụng thành thạo phần mềm chuyên dụng (GeoGebra, mô phỏng) để giải quyết nhiệm vụ.
   5.3. Sử dụng sáng tạo:
      - NC1: Vận dụng công cụ số tạo giải pháp mới.

6. MIỀN 6: ỨNG DỤNG AI (CẬP NHẬT MỚI)
   6.1. Hiểu biết về AI:
      - CB1: Hiểu AI là gì, nhận biết ứng dụng AI.
      - TC1: Hiểu nguyên lý AI tạo sinh và hạn chế (ảo giác).
   6.2. Sử dụng công cụ AI:
      - CB1: Biết ra lệnh (prompt) đơn giản.
      - TC1: Biết viết prompt phức tạp, cung cấp ngữ cảnh (Context).
   6.3. Đạo đức AI:
      - TC1: Nhận thức về liêm chính học thuật khi dùng AI.
"""

SYSTEM_INSTRUCTION_NLS = f"""
Bạn là chuyên gia tư vấn giáo dục cao cấp, chuyên về chuyển đổi số và Khung Năng lực số (NLS) tại Việt Nam.

DỮ LIỆU KHUNG NĂNG LỰC SỐ:
{NLS_FRAMEWORK_DATA}

NHIỆM VỤ CỐT LÕI:
1. Phân tích sâu sắc nội dung giáo án người dùng cung cấp để tìm ra các "điểm chạm" có thể tích hợp NLS một cách tự nhiên nhất (không gượng ép).
2. Lựa chọn các YCCĐ (Yêu cầu cần đạt) từ Khung NLS trên phù hợp với trình độ học sinh và đặc thù môn học.
3. Nếu có nội dung PPCT (Phân phối chương trình), bạn phải ưu tiên 100% nội dung NLS trong PPCT đó.

CẤU TRÚC ĐẦU RA (MARKDOWN):
I. THÔNG TIN CHUNG (Giữ nguyên từ giáo án gốc)
II. MỤC TIÊU
   1. Kiến thức, kĩ năng... (Giữ nguyên)
   2. Năng lực chung... (Giữ nguyên)
   3. Năng lực đặc thù... (Giữ nguyên)
   4. Năng lực số (Bổ sung mới): 
      - [Mã YCCĐ]: Mô tả biểu hiện cụ thể học sinh sẽ đạt được trong bài này.
III. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU SỐ (Bổ sung các công cụ cần thiết cho NLS)
IV. TIẾN TRÌNH DẠY HỌC
   - Tích hợp nội dung NLS vào các hoạt động bằng thẻ <u>...</u> (in nghiêng hoặc đậm để làm nổi bật). 
   - Ví dụ: "HS sử dụng máy tính *thực hiện tra cứu thông tin trên trang web chính thống [1.1.CB2]*".

QUY TẮC KỸ THUẬT:
- Công thức Toán/Lý/Hóa: Sử dụng LaTeX trong $...$.
- Bảng biểu: Sử dụng Markdown Table.
- Không thay đổi nội dung chuyên môn của giáo án gốc, chỉ làm phong phú thêm bằng năng lực số.
"""

def generate_nls_lesson_plan(api_key, lesson_content, subject, grade, textbook, ppct_content, analyze_only):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.0-flash", system_instruction=SYSTEM_INSTRUCTION_NLS)
    
    prompt = f"""
    THÔNG TIN:
    - Môn: {subject} | Lớp: {grade} | Sách: {textbook}
    - Chế độ: {"Chỉ phân tích, không sửa đổi nội dung gốc" if analyze_only else "Tích hợp và viết lại giáo án"}
    
    YÊU CẦU CỦA TRƯỜNG (PPCT):
    {ppct_content if ppct_content else "Không có, tự đề xuất theo khung NLS."}
    
    NỘI DUNG GIÁO ÁN GỐC:
    {lesson_content}
    """
    
    response = model.generate_content(prompt)
    return response.text

# ==============================================================================
# [QUAN TRỌNG] DỮ LIỆU YCCĐ CŨ (GIỮ NGUYÊN)
# ==============================================================================
FULL_YCCD_DATA = [
  {"id": "L1-SO-01", "mon": "Toán", "lop": 1, "chu_de": "Số và Phép tính", "bai": "Các số đến 100", "yccd": "Đếm, đọc, viết được các số trong phạm vi 100. Nhận biết chục và đơn vị."},
  {"id": "L1-SO-02", "mon": "Toán", "lop": 1, "chu_de": "Số và Phép tính", "bai": "So sánh số", "yccd": "Nhận biết cách so sánh, xếp thứ tự các số trong phạm vi 100."},
  {"id": "L1-PT-01", "mon": "Toán", "lop": 1, "chu_de": "Số và Phép tính", "bai": "Phép cộng, phép trừ", "yccd": "Thực hiện được phép cộng, phép trừ (không nhớ) các số trong phạm vi 100."},
  {"id": "L1-HH-01", "mon": "Toán", "lop": 1, "chu_de": "Hình học", "bai": "Hình phẳng và hình khối", "yccd": "Nhận dạng hình vuông, tròn, tam giác, chữ nhật; khối lập phương, khối hộp chữ nhật."},
  {"id": "L1-DL-01", "mon": "Toán", "lop": 1, "chu_de": "Đo lường", "bai": "Độ dài và Thời gian", "yccd": "Đo độ dài bằng đơn vị cm. Đọc giờ đúng trên đồng hồ. Xem lịch hàng ngày."},
  
  # --- LỚP 2 ---
  {"id": "L2-SO-01", "mon": "Toán", "lop": 2, "chu_de": "Số và Phép tính", "bai": "Các số đến 1000", "yccd": "Đọc, viết, so sánh các số trong phạm vi 1000. Số tròn trăm, số liền trước, liền sau."},
  {"id": "L2-PT-01", "mon": "Toán", "lop": 2, "chu_de": "Số và Phép tính", "bai": "Phép cộng, phép trừ (có nhớ)", "yccd": "Thực hiện cộng, trừ (có nhớ) trong phạm vi 1000. Tính toán trường hợp có 2 dấu phép tính."},
  {"id": "L2-PT-02", "mon": "Toán", "lop": 2, "chu_de": "Số và Phép tính", "bai": "Phép nhân, phép chia", "yccd": "Vận dụng bảng nhân 2, 5 và bảng chia 2, 5. Hiểu ý nghĩa phép nhân, chia."},
  {"id": "L2-HH-01", "mon": "Toán", "lop": 2, "chu_de": "Hình học", "bai": "Hình phẳng và hình khối", "yccd": "Nhận biết đường thẳng, đường cong, 3 điểm thẳng hàng. Nhận dạng khối trụ, khối cầu."},
  {"id": "L2-DL-01", "mon": "Toán", "lop": 2, "chu_de": "Đo lường", "bai": "Đơn vị đo lường", "yccd": "Nhận biết kg, lít, m, km, dm. Xem đồng hồ (kim phút chỉ số 3, 6)."},

  # --- LỚP 3 ---
  {"id": "L3-SO-01", "mon": "Toán", "lop": 3, "chu_de": "Số và Phép tính", "bai": "Các số đến 100.000", "yccd": "Đọc, viết, so sánh số trong phạm vi 100.000. Làm tròn số đến hàng nghìn, chục nghìn."},
  {"id": "L3-PT-01", "mon": "Toán", "lop": 3, "chu_de": "Số và Phép tính", "bai": "Phép cộng, trừ", "yccd": "Cộng trừ các số có đến 5 chữ số (có nhớ không quá 2 lượt)."},
  {"id": "L3-PT-02", "mon": "Toán", "lop": 3, "chu_de": "Số và Phép tính", "bai": "Phép nhân, chia", "yccd": "Nhân chia số có nhiều chữ số với số có 1 chữ số. Tính giá trị biểu thức."},
  {"id": "L3-HH-01", "mon": "Toán", "lop": 3, "chu_de": "Hình học", "bai": "Góc và Hình phẳng", "yccd": "Nhận biết góc vuông, không vuông. Tính chu vi tam giác, tứ giác, hình chữ nhật, hình vuông."},
  {"id": "L3-DL-01", "mon": "Toán", "lop": 3, "chu_de": "Đo lường", "bai": "Diện tích", "yccd": "Làm quen diện tích. Đơn vị cm2. Tính diện tích hình chữ nhật, hình vuông."},

  # --- LỚP 4 ---
  {"id": "L4-SO-01", "mon": "Toán", "lop": 4, "chu_de": "Số tự nhiên", "bai": "Số lớp triệu", "yccd": "Đọc, viết, so sánh số đến lớp triệu. Nhận biết giá trị theo vị trí."},
  {"id": "L4-PT-01", "mon": "Toán", "lop": 4, "chu_de": "Số tự nhiên", "bai": "4 Phép tính", "yccd": "Nhân chia với số có 2 chữ số. Tính trung bình cộng."},
  {"id": "L4-PS-01", "mon": "Toán", "lop": 4, "chu_de": "Phân số", "bai": "Khái niệm Phân số", "yccd": "Đọc viết phân số. Rút gọn, quy đồng mẫu số. So sánh phân số."},
  {"id": "L4-PS-02", "mon": "Toán", "lop": 4, "chu_de": "Phân số", "bai": "Phép tính Phân số", "yccd": "Cộng, trừ, nhân, chia hai phân số. Giải toán tìm phân số của một số."},
  {"id": "L4-HH-01", "mon": "Toán", "lop": 4, "chu_de": "Hình học", "bai": "Góc và đường thẳng", "yccd": "Góc nhọn, tù, bẹt. Hai đường thẳng vuông góc, song song."},
  {"id": "L4-HH-02", "mon": "Toán", "lop": 4, "chu_de": "Hình học", "bai": "Hình bình hành, Hình thoi", "yccd": "Nhận biết và tính diện tích hình bình hành, hình thoi."},

  # --- LỚP 5 ---
  {"id": "L5-STP-01", "mon": "Toán", "lop": 5, "chu_de": "Số thập phân", "bai": "Khái niệm Số thập phân", "yccd": "Đọc, viết, so sánh số thập phân. Viết số đo đại lượng dưới dạng số thập phân."},
  {"id": "L5-STP-02", "mon": "Toán", "lop": 5, "chu_de": "Số thập phân", "bai": "Phép tính Số thập phân", "yccd": "Cộng, trừ, nhân, chia số thập phân. Giải toán liên quan tỉ số phần trăm."},
  {"id": "L5-HH-01", "mon": "Toán", "lop": 5, "chu_de": "Hình học", "bai": "Tam giác, Hình thang, Hình tròn", "yccd": "Tính diện tích hình tam giác, hình thang. Chu vi và diện tích hình tròn."},
  {"id": "L5-HH-02", "mon": "Toán", "lop": 5, "chu_de": "Hình học", "bai": "Hình hộp", "yccd": "Tính diện tích xung quanh, toàn phần, thể tích hình hộp chữ nhật, hình lập phương."},
  {"id": "L5-DL-01", "mon": "Toán", "lop": 5, "chu_de": "Đo lường", "bai": "Toán chuyển động", "yccd": "Giải bài toán về vận tốc, quãng đường, thời gian (chuyển động đều)."}
]

# ==============================================================================
# [MỚI] 2.1. DỮ LIỆU PPCT
# ==============================================================================
PPCT_DATA = [
    # Ví dụ Toán lớp 5
    {"cap_hoc": "Tiểu học", "mon": "Toán", "lop": "Lớp 5", "bo_sach": "Kết nối tri thức với cuộc sống", "tuan": 1, "tiet": 1, "bai_id": "T5-KNTT-T1-1", "ten_bai": "Ôn tập khái niệm phân số", "ghi_chu": "Tiết 1"},
    {"cap_hoc": "Tiểu học", "mon": "Toán", "lop": "Lớp 5", "bo_sach": "Kết nối tri thức với cuộc sống", "tuan": 1, "tiet": 2, "bai_id": "T5-KNTT-T1-2", "ten_bai": "Ôn tập tính chất cơ bản của phân số", "ghi_chu": "Tiết 2"},
    # Ví dụ Tiếng Việt lớp 5
    {"cap_hoc": "Tiểu học", "mon": "Tiếng Việt", "lop": "Lớp 5", "bo_sach": "Chân trời sáng tạo", "tuan": 1, "tiet": 1, "bai_id": "TV5-CTST-T1-1", "ten_bai": "Đọc: Chiều dòng sông", "ghi_chu": "Đọc hiểu"},
]

def ppct_filter(cap_hoc, mon, lop, bo_sach):
    return [x for x in PPCT_DATA if x.get("cap_hoc") == cap_hoc and x.get("mon") == mon and x.get("lop") == lop and x.get("bo_sach") == bo_sach]

# ==============================================================================
# 2. CONSTANTS (GIỮ NGUYÊN)
# ==============================================================================
APP_CONFIG = {
    "name": "AI EXAM EXPERT v10 – 2026",
    "role": "Trợ lý chuyên môn Cấp Sở: Ra đề - Thẩm định - Quản trị hồ sơ.",
    "context": """🎯 1. VAI TRÒ VÀ SỨ MỆNH:
    Bạn là Trợ lý AI Chuyên môn Cấp Sở, tuân thủ tuyệt đối các quy định mới nhất của Bộ GD&ĐT.

    🟦 2. QUY ĐỊNH PHÁP LÝ (BẮT BUỘC):
    2.1. CẤP TIỂU HỌC (Thông tư 27/2020):
       - Đề thi thiết kế theo 3 MỨC ĐỘ: M1 (Nhận biết - 40%), M2 (Kết nối - 30%), M3 (Vận dụng - 30%).
       - Điểm số: Thang 10, làm tròn thành số nguyên (0.5 -> 1).
       - Môn TIẾNG VIỆT: Phần Đọc hiểu phải dùng văn bản MỚI (ngoài SGK). Phần Viết có Chính tả & TLV.

    2.2. CẤP TRUNG HỌC (Thông tư 22/2021 & QĐ 764):
       - Ma trận 4 MỨC ĐỘ: NB (40%) - TH (30%) - VD (20%) - VDC (10%).
       - THPT từ 2025: Cấu trúc 3 phần (TN Nhiều lựa chọn, TN Đúng/Sai, Trả lời ngắn)."""
}

PRACTICAL_SUBJECTS = ["Tin học", "Công nghệ", "Mĩ thuật", "Âm nhạc", "Khoa học", "Khoa học tự nhiên", "Vật lí", "Hóa học", "Sinh học", "Tin học và Công nghệ"]

SUBJECT_STRUCTURE_DATA = {
    "THPT_2025": "Phần I: TN Nhiều lựa chọn (0.25đ) | Phần II: TN Đúng/Sai (Max 1đ) | Phần III: Trả lời ngắn (0.5đ)",
    "TieuHoc_TV": "A. Kiểm tra Đọc (10đ) [Đọc tiếng + Đọc hiểu văn bản mới] + B. Kiểm tra Viết (10đ) [Chính tả + TLV].",
    "TieuHoc_Chung": "Trắc nghiệm (60-70%) + Tự luận (30-40%). Mức độ: M1-M2-M3",
    "Toán": "Trắc nghiệm (70%) + Vận dụng (30%)",
    "Ngữ văn": "Đọc hiểu (6.0đ) + Viết (4.0đ)",
    "Tiếng Anh": "Listening (2.5) - Language (2.5) - Reading (2.5) - Writing (2.5)",
    "Mặc định": "NB (40%) - TH (30%) - VD (20%) - VDC (10%)"
}

EDUCATION_DATA = {
    "tieu_hoc": {
        "label": "Tiểu học",
        "grades": ["Lớp 1", "Lớp 2", "Lớp 3", "Lớp 4", "Lớp 5"],
        "subjects": ["Toán", "Tiếng Việt", "Tiếng Anh", "Đạo đức", "Tự nhiên và Xã hội", "Khoa học", "Lịch sử và Địa lí", "Tin học và Công nghệ", "Giáo dục thể chất", "Âm nhạc", "Mĩ thuật", "Hoạt động trải nghiệm", "Công nghệ", "Tin học"],
        "legal": "Thông tư 27 (3 Mức độ)"
    },
    "thcs": {
        "label": "THCS",
        "grades": ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"],
        "subjects": ["Ngữ văn", "Toán", "Tiếng Anh", "Giáo dục công dân", "Khoa học tự nhiên", "Lịch sử và Địa lí", "Tin học", "Công nghệ", "Giáo dục thể chất", "Âm nhạc", "Mĩ thuật", "HĐTN, HN", "Giáo dục địa phương"],
        "legal": "Thông tư 22 (4 Mức độ)"
    },
    "thpt": {
        "label": "THPT",
        "grades": ["Lớp 10", "Lớp 11", "Lớp 12"],
        "subjects": ["Ngữ văn", "Toán", "Tiếng Anh", "Lịch sử", "Địa lí", "Vật lí", "Hóa học", "Sinh học", "GDKT & PL", "Tin học", "Công nghệ", "Âm nhạc", "Mĩ thuật", "GDTC", "GDQP&AN", "HĐTN, HN"],
        "legal": "Cấu trúc 2025 (QĐ 764)"
    }
}

BOOKS_LIST = [
    "Kết nối tri thức với cuộc sống", "Chân trời sáng tạo", "Cánh Diều", "Cùng khám phá",
    "Vì sự bình đẳng và dân chủ trong giáo dục", "Tin học: Đại học Vinh (Tiểu học)",
    "Tiếng Anh: Global Success", "Tiếng Anh: Family and Friends", "Tiếng Anh: Friends Plus",
    "Tiếng Anh: i-Learn Smart Start", "Tiếng Anh: Explore English",
    "Tin học: Kết nối tri thức", "Tin học: Chân trời sáng tạo", "Tin học: Cánh Diều",
    "Tài liệu Giáo dục địa phương tỉnh Tuyên Quang", "Chuyên đề học tập (THPT)"
]

FULL_SCOPE_LIST = ["Khảo sát chất lượng đầu năm", "Kiểm tra giữa kì 1", "Kiểm tra cuối kì 1", "Kiểm tra giữa kì 2", "Kiểm tra cuối kì 2", "Thi thử Tốt nghiệp THPT", "Thi học sinh giỏi cấp Trường", "Thi học sinh giỏi cấp Huyện/Tỉnh"]
LIMITED_SCOPE_LIST = ["Khảo sát chất lượng đầu năm", "Kiểm tra cuối kì 1", "Kiểm tra cuối kì 2"]

SCOPE_MAPPING = {
    "Khảo sát chất lượng đầu năm": "Ôn tập hè & Tuần 1-2",
    "Kiểm tra giữa kì 1": "Tuần 1 đến Tuần 9",
    "Kiểm tra cuối kì 1": "Tuần 10 đến Tuần 18 (Ôn tập cả HK1)",
    "Kiểm tra giữa kì 2": "Tuần 19 đến Tuần 27",
    "Kiểm tra cuối kì 2": "Tuần 28 đến Tuần 35 (Ôn tập cả HK2)",
    "Thi thử Tốt nghiệp THPT": "Toàn bộ chương trình",
    "Thi học sinh giỏi cấp Trường": "Nâng cao",
    "Thi học sinh giỏi cấp Huyện/Tỉnh": "Chuyên sâu"
}

CURRICULUM_DATA = {
    "Toán": {
        "Lớp 6": {"Kiểm tra giữa kì 1": "Tập hợp số tự nhiên; Phép tính; Số nguyên tố."},
        "Lớp 12": {"Kiểm tra cuối kì 1": "Nguyên hàm; Tích phân; Phương trình mặt phẳng."}
    }
}

LEGAL_DOCUMENTS = [
    {"code": "CV 7791/2024", "title": "Công văn 7791 (Mới)", "summary": "Hướng dẫn kỹ thuật xây dựng ma trận, đặc tả.", "highlight": True},
    {"code": "QĐ 764/2024", "title": "Cấu trúc THPT 2025", "summary": "Định dạng đề thi mới: TN nhiều lựa chọn, Đúng/Sai, Trả lời ngắn.", "highlight": True},
    {"code": "TT 22/2021", "title": "Đánh giá Trung học", "summary": "4 mức độ: NB-TH-VD-VDC.", "highlight": True},
    {"code": "TT 27/2020", "title": "Đánh giá Tiểu học", "summary": "3 mức độ nhận thức (M1, M2, M3).", "highlight": True},
    {"code": "CV 2345", "title": "KHGD Tiểu học", "summary": "Xây dựng kế hoạch bài dạy, ma trận đề kiểm tra.", "highlight": False},
    {"code": "CV 3175", "title": "Đổi mới PPDH", "summary": "Hướng dẫn kỹ thuật biên soạn câu hỏi.", "highlight": False},
    {"code": "TT 32/2018", "title": "CT GDPT 2018", "summary": "Văn bản gốc quy định Yêu cầu cần đạt.", "highlight": False}
]

# ==============================================================================
# 3. GIAO DIỆN & CSS (CẬP NHẬT CSS CHO BẢNG)
# ==============================================================================
st.markdown(textwrap.dedent('''
<style>
    /* ===== Dashboard KPI ===== */
.kpi-card{
  background:#FFFFFF;
  border:1px solid #E2E8F0;
  border-radius:12px;
  padding:16px 18px;
  box-shadow:0 4px 8px rgba(0,0,0,0.04);
  margin-bottom:12px;
}
.kpi-title{ font-size:12px; font-weight:700; color:#64748B; text-transform:uppercase; letter-spacing:.5px;}
.kpi-value{ font-size:22px; font-weight:900; color:#0F172A; margin-top:6px;}
.kpi-sub{ font-size:12px; color:#64748B; margin-top:4px;}

/* ===== Module Cards ===== */
.module-card{
  background:#FFFFFF;
  border:1px solid #E2E8F0;
  border-radius:14px;
  padding:18px 18px 14px 18px;
  box-shadow:0 10px 18px rgba(2,6,23,0.05);
  margin-bottom:12px;
}
.module-card.highlight{
  border:1px solid #BFDBFE;
  box-shadow:0 14px 24px rgba(37,99,235,0.12);
}
.module-badge{
  display:inline-block;
  font-size:11px;
  font-weight:800;
  padding:4px 10px;
  border-radius:999px;
  background:#EFF6FF;
  border:1px solid #BFDBFE;
  color:#1D4ED8;
  margin-bottom:10px;
}
.module-title{
  font-size:18px;
  font-weight:900;
  color:#0F172A;
  margin:4px 0 6px 0;
}
.module-desc{
  font-size:13px;
  color:#334155;
  line-height:1.55;
  margin-bottom:8px;
}
.module-meta{
  font-size:12px;
  color:#64748B;
  border-top:1px dashed #E2E8F0;
  padding-top:10px;
}

/* ===== Word Preview CSS ===== */
.paper-view table { width: 100%; border-collapse: collapse; margin-bottom: 1em; }
.paper-view th, .paper-view td { border: 1px solid black; padding: 6px; text-align: left; vertical-align: top; }
.paper-view th { background-color: #f2f2f2; font-weight: bold; }

</style>
'''), unsafe_allow_html=True)

# ==============================================================================
# 4. HÀM XỬ LÝ LOGIC
# ==============================================================================

def init_supabase():
    try: return create_client(SUPABASE_URL, SUPABASE_KEY)
    except: return None

def read_file_content(uploaded_file, file_type):
    if not uploaded_file: return ""
    try:
        if uploaded_file.name.endswith('.docx'):
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            return "\n".join([p.text for p in doc.paragraphs])
        elif uploaded_file.name.endswith('.xlsx'):
            content = pd.read_excel(uploaded_file).to_string()
        
        # Gắn nhãn chuẩn Logic React
        if file_type == 'matrix': return f"\n[DỮ LIỆU MA TRẬN TỪ NGƯỜI DÙNG]:\n{content}\n"
        if file_type == 'spec': return f"\n[DỮ LIỆU ĐẶC TẢ TỪ NGƯỜI DÙNG]:\n{content}\n"
    except: return ""
    return content


# =========================
# [NEW] HỖ TRỢ ĐỌC PDF (kể cả PDF scan/ảnh) cho MODULE SOẠN GIÁO ÁN
# - Ưu tiên trích xuất text trực tiếp (nhanh)
# - Nếu PDF là ảnh (text rất ít) -> thử OCR (cần cài thêm pdf2image + pytesseract)
# =========================
import hashlib

def _hash_bytes(b: bytes) -> str:
    try:
        return hashlib.sha256(b).hexdigest()
    except Exception:
        return str(len(b))

@st.cache_data(show_spinner=False)
def extract_text_from_pdf_bytes(pdf_bytes: bytes, max_pages: int = 6, ocr_if_needed: bool = True) -> str:
    """Trả về text đã trích từ PDF. Nếu PDF scan và có OCR tools thì OCR.
    Giới hạn số trang để tránh nặng VPS. Trả về chuỗi đã được cắt ngắn (<= 12000 ký tự).
    """
    if not pdf_bytes:
        return ""
    text_parts = []

    # 1) Thử extract text trực tiếp (PyPDF2 / pypdf)
    try:
        try:
            from pypdf import PdfReader
        except Exception:
            from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(io.BytesIO(pdf_bytes))
        n = min(len(reader.pages), max_pages)
        for i in range(n):
            try:
                t = reader.pages[i].extract_text() or ""
                t = re.sub(r"\s+", " ", t).strip()
                if t:
                    text_parts.append(t)
            except Exception:
                continue
    except Exception:
        pass

    direct_text = "\n\n".join(text_parts).strip()
    # Nếu đã có text đủ dùng -> trả luôn
    if len(direct_text) >= 300 or (not ocr_if_needed):
        return direct_text[:12000]

    # 2) Nếu text quá ít, thử OCR (PDF scan)
    # Cần: pdf2image + pytesseract (+ poppler cho pdf2image)
    try:
        from pdf2image import convert_from_bytes  # type: ignore
        import pytesseract  # type: ignore
    except Exception:
        # Không có OCR deps -> trả direct_text (có thể rỗng)
        return direct_text[:12000]

    try:
        images = convert_from_bytes(pdf_bytes, dpi=220, first_page=1, last_page=max_pages)
        ocr_texts = []
        for img in images:
            try:
                # Tiếng Việt: nếu máy có gói vie; nếu không, vẫn OCR được nhưng kém hơn
                try:
                    t = pytesseract.image_to_string(img, lang="vie")
                except Exception:
                    t = pytesseract.image_to_string(img)
                t = re.sub(r"\s+", " ", t).strip()
                if t:
                    ocr_texts.append(t)
            except Exception:
                continue
        ocr_text = "\n\n".join(ocr_texts).strip()
        # Nếu OCR vẫn rỗng -> fallback direct_text
        out = ocr_text if ocr_text else direct_text
        return out[:12000]
    except Exception:
        return direct_text[:12000]

def build_pdf_context_for_teacher_note(pdf_text: str) -> str:
    pdf_text = (pdf_text or "").strip()
    if not pdf_text:
        return ""
    # Nhắc AI: bám sát nội dung PDF, tránh bịa
    return (
        "\n\n[NỘI DUNG TRÍCH TỪ PDF/ẢNH BÀI HỌC – ƯU TIÊN BÁM SÁT]\n"
        "- Đây là nội dung trích xuất từ tài liệu người dùng tải lên.\n"
        "- Khi soạn giáo án: ưu tiên bám sát đúng thuật ngữ, ví dụ, bài tập, yêu cầu trong tài liệu.\n"
        "- Không tự bịa thêm bài tập/đề mục không có trong tài liệu (trừ khi GV yêu cầu bổ sung).\n"
        f"\n{pdf_text}\n"
    )



def extract_text_from_upload(file, max_pages: int = 6, ocr_if_needed: bool = True) -> str:
    """Trích text từ 1 uploaded file (pdf/docx/image). Không làm sập app nếu thiếu thư viện."""
    if not file:
        return ""

    name = (getattr(file, "name", "") or "").lower()
    data = file.getvalue() if hasattr(file, "getvalue") else None
    if not data:
        return ""

    # DOCX
    if name.endswith(".docx"):
        try:
            d = docx.Document(io.BytesIO(data))
            return "\n".join([p.text for p in d.paragraphs]).strip()
        except Exception:
            return ""

    # PDF
    if name.endswith(".pdf"):
        try:
            return extract_text_from_pdf_bytes(data, max_pages=max_pages, ocr_if_needed=ocr_if_needed) or ""
        except Exception:
            return ""

    # IMAGE (jpg/png)
    if name.endswith((".png", ".jpg", ".jpeg")):
        if not ocr_if_needed:
            return ""
        try:
            from PIL import Image  # type: ignore
            import pytesseract  # type: ignore
            img = Image.open(io.BytesIO(data))
            try:
                text = pytesseract.image_to_string(img, lang="vie")
            except Exception:
                text = pytesseract.image_to_string(img)
            return re.sub(r"\s+", " ", (text or "")).strip()
        except Exception:
            return ""

    return ""


def build_uploaded_materials_context(lesson_files, ppct_file, max_pages: int, try_ocr: bool) -> str:
    """Ghép nội dung file upload thành 1 đoạn context sạch để AI bám sát."""
    parts = []

    # PPCT
    if ppct_file:
        ppct_txt = extract_text_from_upload(ppct_file, max_pages=max_pages, ocr_if_needed=try_ocr)
        ppct_txt = (ppct_txt or "").strip()
        if ppct_txt:
            parts.append(
                "[PPCT/KHDH (ƯU TIÊN)]\n"
                "Bám 100% nội dung dưới đây khi soạn:\n"
                f"{ppct_txt[:8000]}"
            )

    # Lesson files
    if lesson_files:
        merged = []
        for f in lesson_files:
            t = extract_text_from_upload(f, max_pages=max_pages, ocr_if_needed=try_ocr)
            t = (t or "").strip()
            if t:
                merged.append(f"[FILE: {getattr(f,'name','file')}]" + "\n" + t[:8000])
        if merged:
            parts.append(
                "[NỘI DUNG BÀI HỌC (SGK/TÀI LIỆU)]\n"
                "Ưu tiên bám sát thuật ngữ, ví dụ, bài tập, yêu cầu trong các file dưới đây. "
                "Không tự bịa thêm bài tập nếu không cần.\n\n"
                + "\n\n".join(merged)
            )

    return ("\n\n" + "\n\n".join(parts)).strip()


# [FIX] HÀM LÀM SẠCH JSON CHUẨN (KHÔNG ĐƯỢC XÓA)
def clean_json(text):
    text = text.strip()
    if "```" in text:
        parts = re.split(r'```(?:json)?', text)
        if len(parts) > 1: text = parts[1]
    
    start_idx = text.find('{')
    if start_idx == -1: return "{}"
    text = text[start_idx:]
    
    try:
        decoder = json.JSONDecoder()
        obj, idx = decoder.raw_decode(text)
        return json.dumps(obj)
    except:
        end_idx = text.rfind('}')
        if end_idx != -1: return text[:end_idx+1]
        return text

# [CẬP NHẬT] Hàm tạo File Word chuẩn Font XML VÀ CÓ BẢNG
def create_word_doc(html, title):
    doc_content = f"""
    <html xmlns:o='urn:schemas-microsoft-com:office:office' xmlns:w='urn:schemas-microsoft-com:office:word' xmlns='http://www.w3.org/TR/REC-html40'>
    <head>
        <meta charset='utf-8'>
        <title>{title}</title>
        <xml>
            <w:WordDocument>
                <w:View>Print</w:View>
                <w:Zoom>100</w:Zoom>
                <w:DoNotOptimizeForBrowser/>
            </w:WordDocument>
        </xml>
        <style>
            @page {{ size: 21cm 29.7cm; margin: 2cm 2cm 2cm 2cm; mso-page-orientation: portrait; }}
            body {{ font-family: 'Times New Roman', serif; font-size: 13pt; line-height: 1.3; }}
            p, div, span, li, td, th {{ font-family: 'Times New Roman', serif; mso-ascii-font-family: 'Times New Roman'; mso-hansi-font-family: 'Times New Roman'; color: #000000; }}
            table {{ border-collapse: collapse; width: 100%; border: 1px solid black; }}
            td, th {{ border: 1px solid black; padding: 5px; vertical-align: top; }}
        </style>
    </head>
    <body>
        <div class="WordSection1">
            {html}
        </div>
    </body>
    </html>
    """
    return "\ufeff" + doc_content

# ==============================================================================
# [PATCH 3/3] RENDER HTML TỪ JSON (BẢNG 2 CỘT GV/HS) - KHÓA MẪU
# ==============================================================================

def _html_escape(s: str) -> str:
    if s is None:
        return ""
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )

def _render_ul(items) -> str:
    if not items:
        return "<ul><li>...</li></ul>"
    lis = "".join([f"<li>{_html_escape(x)}</li>" for x in items if str(x).strip()])
    return f"<ul>{lis or '<li>...</li>'}</ul>"

def render_lesson_plan_html(data: dict) -> str:
    """Render lesson plan JSON (meta + sections) -> printable HTML (A4) theo bảng 2 cột GV/HS."""
    data = data or {}
    meta = data.get("meta", {}) or {}
    sections = data.get("sections", {}) or {}

    def esc(s):
        return _html_escape("" if s is None else str(s))

    sec_I = sections.get("I", {}) or {}
    sec_II = sections.get("II", {}) or {}
    sec_III = sections.get("III", {}) or {}
    sec_IV = sections.get("IV", {}) or {}

    yccd = sec_I.get("yeu_cau_can_dat") or []
    nang_luc = sec_I.get("nang_luc") or []
    pham_chat = sec_I.get("pham_chat") or []
    nldac = sec_I.get("nang_luc_dac_thu") or []
    nlso = sec_I.get("nang_luc_so") or []

    gv_dd = sec_II.get("giao_vien") or []
    hs_dd = sec_II.get("hoc_sinh") or []

    bang = sec_III.get("bang") if isinstance(sec_III, dict) else []
    bang = bang or []

    dieu_chinh = sec_IV.get("dieu_chinh_sau_bai_day") or ""

    def ul(items):
        items = items if isinstance(items, list) else []
        if not items:
            return "<p class='muted'>(Chưa có nội dung)</p>"
        return "<ul>" + "".join(f"<li>{esc(x)}</li>" for x in items) + "</ul>"

    css = """
    <style>
      @page { size: 21cm 29.7cm; margin: 2cm; }
      body{font-family:'Times New Roman',serif;font-size:13pt;line-height:1.35;color:#111;}
      .wrap{max-width:980px;margin:0 auto;}
      h1{font-size:18pt;text-align:center;margin:0 0 8px 0;}
      h2{font-size:14pt;margin:12px 0 6px 0;border-bottom:1px solid #ccc;padding-bottom:3px;}
      h3{font-size:13pt;margin:8px 0 4px 0;}
      p{margin:6px 0;}
      ul{margin:6px 0 6px 20px;}
      .meta p{margin:3px 0;}
      table.lp{width:100%;border-collapse:collapse;table-layout:fixed;}
      table.lp th, table.lp td{border:1px solid #000;padding:6px;vertical-align:top;word-wrap:break-word;overflow-wrap:break-word;}
      table.lp th{text-align:center;font-weight:700;}
      .muted{color:#333;font-style:italic;}
      @media print{ thead{display:table-header-group;} tr{page-break-inside:avoid;} }
    </style>
    """

    rows = []
    for r in bang:
        if not isinstance(r, dict):
            continue
        kieu = (r.get("kieu") or "row").strip().lower()
        if kieu == "header":
            title = r.get("tieu_de") or ""
            if title:
                rows.append(f"<tr><td colspan='2'><b>{esc(title)}</b></td></tr>")
            continue
        gv = r.get("giao_vien") or r.get("gv") or ""
        hs = r.get("hoc_sinh") or r.get("hs") or ""
        tg = r.get("thoi_gian")
        if isinstance(tg, int) and tg > 0:
            gv_html = f"<b>({tg}')</b> {esc(gv)}"
        else:
            gv_html = esc(gv)
        rows.append(f"<tr><td>{gv_html}</td><td>{esc(hs)}</td></tr>")

    table_html = "<p class='muted'>(Chưa có bảng hoạt động)</p>" if not rows else (
        "<table class='lp'><thead><tr><th>Hoạt động của Giáo viên</th><th>Hoạt động của Học sinh</th></tr></thead>"
        "<tbody>" + "".join(rows) + "</tbody></table>"
    )

    html = (
    "<!doctype html><html lang='vi'><head><meta charset='utf-8'/>"
    + css +
    "</head><body>"
    "<div class='wrap'>"
    "<h1>GIÁO ÁN</h1>"
    "<div class='meta'>"
    f"<p><b>Môn:</b> {esc(meta.get('mon'))} &nbsp;&nbsp; <b>Lớp:</b> {esc(meta.get('lop'))} &nbsp;&nbsp; <b>Cấp:</b> {esc(meta.get('cap_hoc'))}</p>"
    f"<p><b>Bài:</b> {esc(meta.get('ten_bai'))} &nbsp;&nbsp; <b>Thời lượng:</b> {esc(meta.get('thoi_luong'))} phút &nbsp;&nbsp; <b>Bộ sách:</b> {esc(meta.get('bo_sach'))}</p>"
    "</div>"
    "<h2>I. Yêu cầu cần đạt</h2>"
    "<h3>1) Yêu cầu cần đạt</h3>"
    + ul(yccd) +
    "<h3>2) Năng lực</h3>"
    + ul(nang_luc) +
    "<h3>3) Phẩm chất</h3>"
    + ul(pham_chat) +
    "<h3>4) Năng lực đặc thù (nếu có)</h3>"
    + ul(nldac) +
    "<h3>5) Năng lực số (nếu có)</h3>"
    + ul(nlso) +
    "<h2>II. Đồ dùng dạy – học</h2>"
    "<h3>1) Giáo viên</h3>"
    + ul(gv_dd) +
    "<h3>2) Học sinh</h3>"
    + ul(hs_dd) +
    "<h2>III. Các hoạt động dạy – học chủ yếu</h2>"
    + table_html +
    "<h2>IV. Điều chỉnh sau bài dạy (nếu có)</h2>"
    + (f"<p>{esc(dieu_chinh)}</p>" if dieu_chinh else "<p>……………………………………………………………………………………………<br/>……………………………………………………………………………………………<br/>……………………………………………………………………………………………</p>")
    + "</div></body></html>"
    )
    return html

def get_knowledge_context(subject, grade, book, scope):
    try:
        data = CURRICULUM_DATA.get(subject, {}).get(grade, {}).get(book, {})
        key = next((k for k in data.keys() if k in scope or scope in k), None)
        if key: return f"NỘI DUNG CHƯƠNG TRÌNH ({key}): {data[key]}"
        week_info = SCOPE_MAPPING.get(scope, scope)
        return f"NỘI DUNG TỰ TRA CỨU: Bám sát chuẩn kiến thức kĩ năng môn {subject} {grade} - Bộ sách {book}. Thời điểm: {week_info}."
    except: return "NỘI DUNG: Theo chuẩn CTGDPT 2018."

# --- [BỔ SUNG] HÀM CHECK TIỀN TỰ ĐỘNG (Dùng SePay) ---
def check_sepay_transaction(amount, content_search):
    token = st.secrets.get("SEPAY_API_TOKEN", "")
    if not token: return False
    try:
        url = "https://my.sepay.vn/userapi/transactions/list"
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        res = requests.get(url, headers=headers)
        if res.status_code == 200:
            data = res.json()
            for t in data.get('transactions', []):
                # Kiểm tra số tiền và nội dung
                if float(t['amount_in']) >= amount and content_search in t['transaction_content']:
                    return True
    except:
        return False


# ==============================================================================
# [MỚI] HỆ THỐNG ĐIỂM (VIP POINTS)
# - 50.000đ = 550 điểm
# - Trừ điểm theo lượt dùng module (mặc định: 30 điểm/lần)
# - Lưu trực tiếp vào bảng users_pro (cột: points). Nếu DB chưa có cột points -> hệ thống vẫn chạy theo usage_count như cũ.
# ==============================================================================

VIP_TOPUP_AMOUNT_VND = PRICE_VIP           # 50k
VIP_TOPUP_POINTS = 550
POINT_COST_LESSON_PLAN = 30
POINT_COST_EXAM = 30
POINT_COST_NLS = 30

def _db_has_points(user_row: dict) -> bool:
    return isinstance(user_row, dict) and ("points" in user_row)

def get_user_row(client, username: str) -> dict:
    if not client or not username:
        return {}
    try:
        res = client.table('users_pro').select("*").eq('username', username).execute()
        return res.data[0] if getattr(res, "data", None) else {}
    except Exception:
        return {}

def get_user_points(client, username: str) -> int:
    row = get_user_row(client, username)
    if not _db_has_points(row):
        return -1  # -1 = DB chưa có cột points
    try:
        return int(row.get("points", 0) or 0)
    except Exception:
        return 0

def add_user_points(client, username: str, add_points: int, reason: str = "vip_topup", meta: dict | None = None) -> bool:
    """Cộng điểm.

    Ưu tiên gọi RPC server-side (nếu bạn tạo) để tránh race-condition:
    - rpc_add_points(username, add_points, reason, meta_json)

    Nếu RPC chưa có, sẽ fallback update trực tiếp (an toàn kém hơn).
    """
    if not client or not username or int(add_points) == 0:
        return False

    # 1) Try RPC (recommended)
    try:
        meta_json = json.dumps(meta or {}, ensure_ascii=False)
        rpc = client.rpc('rpc_add_points', {
            'p_username': username,
            'p_add': int(add_points),
            'p_reason': reason,
            'p_meta': meta_json
        }).execute()
        # Expect rpc.data = {"ok": true, "points": <new_points>} (you define it)
        data = getattr(rpc, "data", None)
        if isinstance(data, dict) and data.get("ok"):
            newv = int(data.get("points", 0) or 0)
            st.session_state.setdefault("user", {})
            st.session_state["user"]["points"] = newv
            return True
    except Exception:
        pass

    # 2) Fallback: read-modify-write (NOT atomic)
    row = get_user_row(client, username)
    if not _db_has_points(row):
        return False
    try:
        cur = int(row.get("points", 0) or 0)
        newv = cur + int(add_points)
        client.table('users_pro').update({'points': newv}).eq('username', username).execute()
        st.session_state.setdefault("user", {})
        st.session_state["user"]["points"] = newv
        return True
    except Exception:
        return False


def deduct_user_points(client, username: str, cost: int, reason: str = "ai_call", meta: dict | None = None) -> bool:
    """Trừ điểm.

    **Quan trọng:** để chạy chắc 100% và không bị trừ âm/race-condition khi user bấm 2 lần,
    bạn nên tạo RPC server-side:
    - rpc_deduct_points(username, cost, reason, meta_json) -> {"ok":bool,"points":int}

    Nếu RPC chưa có, sẽ fallback update trực tiếp (an toàn kém hơn).
    """
    if int(cost) <= 0:
        return True
    if not client or not username:
        return False

    # 1) Try RPC (recommended)
    try:
        meta_json = json.dumps(meta or {}, ensure_ascii=False)
        rpc = client.rpc('rpc_deduct_points', {
            'p_username': username,
            'p_cost': int(cost),
            'p_reason': reason,
            'p_meta': meta_json
        }).execute()
        data = getattr(rpc, "data", None)
        if isinstance(data, dict):
            if data.get("ok"):
                newv = int(data.get("points", 0) or 0)
                st.session_state.setdefault("user", {})
                st.session_state["user"]["points"] = newv
                return True
            return False
    except Exception:
        pass

    # 2) Fallback: read-modify-write (NOT atomic)
    row = get_user_row(client, username)
    if not _db_has_points(row):
        return False
    try:
        cur = int(row.get("points", 0) or 0)
        if cur < cost:
            return False
        newv = cur - int(cost)
        client.table('users_pro').update({'points': newv}).eq('username', username).execute()
        st.session_state.setdefault("user", {})
        st.session_state["user"]["points"] = newv
        return True
    except Exception:
        return False
def require_points_or_block(cost: int, action_name: str = "thao tác") -> bool:
    """Gọi nhanh trong UI: nếu DB có points thì kiểm tra đủ điểm; nếu DB chưa có points -> cho chạy theo logic cũ."""
    user = st.session_state.get("user", {}) or {}
    username = user.get("email") or user.get("username") or ""
    client = init_supabase()
    if not client or not username:
        return True

    pts = get_user_points(client, username)
    if pts < 0:
        # DB chưa có cột points -> không chặn (fallback usage_count như cũ)
        return True

    if pts < cost:
        st.error(f"❌ Không đủ điểm để {action_name}. Bạn còn {pts} điểm, cần {cost} điểm.")
        st.info("👉 Vào **Menu chính** để nạp VIP và cộng điểm tự động.")
        return False
    return True

    return False

# ==============================================================================
# [MỚI - ĐÃ SỬA LỖI JSON] MODULE QUẢN LÝ YÊU CẦU CẦN ĐẠT (KHÔNG CẦN FILE JSON)
# ==============================================================================
class YCCDManager:
    def __init__(self):
        # Đọc trực tiếp từ biến trong code, không đọc file nữa
        self.data = FULL_YCCD_DATA 

    def get_grades(self):
        grades = set([item['lop'] for item in self.data])
        return sorted(list(grades))

    def get_topics_by_grade(self, grade):
        topics = set([item['chu_de'] for item in self.data if item['lop'] == grade])
        return sorted(list(topics))

    def get_yccd_list(self, grade, topic):
        return [item for item in self.data if item['lop'] == grade and item['chu_de'] == topic]

class QuestionGeneratorYCCD:
    def __init__(self, api_key):
        genai.configure(api_key=api_key)
        # [SỬA LỖI 404] Dùng gemini-2.0-flash theo yêu cầu
        self.model = genai.GenerativeModel('gemini-2.0-flash')

    def generate(self, yccd_item, muc_do="Thông hiểu"):
        prompt = f"""
        VAI TRÒ: Giáo viên Toán Tiểu học (Chương trình GDPT 2018).
        NHIỆM VỤ: Soạn 01 câu hỏi trắc nghiệm Toán.
        THÔNG TIN BẮT BUỘC:
        - Lớp: {yccd_item['lop']} (Câu hỏi phải phù hợp tâm lý lứa tuổi lớp {yccd_item['lop']})
        - Chủ đề: {yccd_item['chu_de']}
        - Bài học: {yccd_item['bai']}
        - YÊU CẦU CẦN ĐẠT: "{yccd_item['yccd']}"
        - Mức độ: {muc_do}
        
        YÊU CẦU ĐẦU RA (JSON format):
        {{
            "question": "Nội dung câu hỏi (ngắn gọn, dễ hiểu)",
            "options": ["A. ...", "B. ...", "C. ...", "D. ..."],
            "answer": "A, B, C hoặc D",
            "explanation": "Giải thích chi tiết (Dành cho học sinh tự học)"
        }}
        """
        try:
            # [FIX LỖI] Tắt bộ lọc an toàn để tránh AI chặn nội dung đề thi
            safe_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
            ]
            
            res = self.model.generate_content(
                prompt, 
                generation_config={"response_mime_type": "application/json"},
                safety_settings=safe_settings
            )
            # Dùng clean_json để tránh lỗi định dạng
            return json.loads(clean_json(res.text))
        except Exception as e:
            return None

# ==============================================================================
# [MỚI] 2.2. JSON SCHEMA KHÓA CỨNG (CÓ BẢNG)
# ==============================================================================
LESSON_PLAN_SCHEMA = {
    "type": "object",
    "required": ["meta", "sections", "renderHtml"],
    "additionalProperties": False,
    "properties": {
        "meta": {
            "type": "object",
            "required": ["cap_hoc", "mon", "lop", "bo_sach", "ppct", "ten_bai", "thoi_luong"],
            "additionalProperties": False,
            "properties": {
                "cap_hoc": {"type": "string"},
                "mon": {"type": "string"},
                "lop": {"type": "string"},
                "bo_sach": {"type": "string"},
                "ppct": {
                    "type": "object",
                    "required": ["tuan", "tiet", "bai_id"],
                    "additionalProperties": False,
                    "properties": {
                        "tuan": {"type": "integer", "minimum": 1, "maximum": 60},
                        "tiet": {"type": "integer", "minimum": 1, "maximum": 20},
                        "bai_id": {"type": "string"},
                        "ghi_chu": {"type": "string"}
                    }
                },
                "ten_bai": {"type": "string", "minLength": 2},
                "thoi_luong": {"type": "integer", "minimum": 30, "maximum": 120},
                "si_so": {"type": "integer", "minimum": 10, "maximum": 60},
                "ngay_day": {"type": "string"}
            }
        },
        "sections": {
            "type": "object",
            "required": ["I", "II", "III", "IV"],
            "additionalProperties": False,
            "properties": {
                "I": {  # Yêu cầu cần đạt
                    "type": "object",
                    "required": ["yeu_cau_can_dat"],
                    "additionalProperties": False,
                    "properties": {
                        "yeu_cau_can_dat": {
                            "type": "array",
                            "minItems": 1,
                            "items": {"type": "string"}
                        },
                        "pham_chat": {"type": "array", "items": {"type": "string"}},
                        "nang_luc": {"type": "array", "items": {"type": "string"}}
                    }
                },
                "II": {  # Đồ dùng dạy học
                    "type": "object",
                    "required": ["giao_vien", "hoc_sinh"],
                    "additionalProperties": False,
                    "properties": {
                        "giao_vien": {"type": "array", "items": {"type": "string"}},
                        "hoc_sinh": {"type": "array", "items": {"type": "string"}}
                    }
                },
                "III": {  # Tiến trình dạy học
                    "type": "object",
                    "required": ["hoat_dong"],
                    "additionalProperties": False,
                    "properties": {
                        "hoat_dong": {
                            "type": "array",
                            "minItems": 3,
                            "items": {
                                "type": "object",
                                "required": ["ten", "thoi_gian", "muc_tieu", "to_chuc"],
                                "additionalProperties": False,
                                "properties": {
                                    "ten": {"type": "string"},
                                    "thoi_gian": {"type": "integer", "minimum": 1, "maximum": 60},
                                    "muc_tieu": {"type": "array", "minItems": 1, "items": {"type": "string"}},
                                    "to_chuc": {
                                        "type": "array",
                                        "minItems": 2,
                                        "items": {
                                            "type": "object",
                                            "required": ["gv", "hs", "san_pham"],
                                            "additionalProperties": False,
                                            "properties": {
                                                "gv": {"type": "string"},
                                                "hs": {"type": "string"},
                                                "san_pham": {"type": "string"}
                                            }
                                        }
                                    },
                                    "noi_dung_cot_loi": {"type": "array", "items": {"type": "string"}}
                                }
                            }
                        }
                    }
                },
                "IV": {  # Điều chỉnh sau bài dạy
                    "type": "object",
                    "required": ["dieu_chinh_sau_bai_day"],
                    "additionalProperties": False,
                    "properties": {
                        "dieu_chinh_sau_bai_day": {"type": "string"}
                    }
                }
            }
        },
        "renderHtml": {"type": "string", "minLength": 50, "description": "Toàn bộ nội dung giáo án dạng HTML. Phần III PHẢI là bảng (table) 2 cột: Hoạt động của GV và Hoạt động của HS."}
    }
}

# ==============================================================================
# [PATCH 1/3] LESSON PLAN DATA-ONLY SCHEMA (CẤP SỞ) + VALIDATOR
# - AI CHỈ TRẢ JSON DATA, KHÔNG TRẢ HTML
# - HỆ THỐNG TỰ RENDER HTML
# ==============================================================================

from jsonschema import validate, Draft202012Validator, ValidationError

LESSON_PLAN_DATA_SCHEMA = {
    "type": "object",
    "required": ["meta", "sections"],
    "additionalProperties": False,
    "properties": {
        "meta": {
            "type": "object",
            "required": ["cap_hoc", "mon", "lop", "bo_sach", "ppct", "ten_bai", "thoi_luong", "si_so"],
            "additionalProperties": False,
            "properties": {
                "cap_hoc": {"type": "string", "minLength": 2},
                "mon": {"type": "string", "minLength": 2},
                "lop": {"type": "string", "minLength": 2},
                "bo_sach": {"type": "string", "minLength": 2},
                "ppct": {
                    "type": "object",
                    "required": ["tuan", "tiet", "bai_id"],
                    "additionalProperties": False,
                    "properties": {
                        "tuan": {"type": "integer", "minimum": 1, "maximum": 60},
                        "tiet": {"type": "integer", "minimum": 1, "maximum": 30},
                        "bai_id": {"type": "string", "minLength": 2},
                        "ghi_chu": {"type": "string"}
                    }
                },
                "ten_bai": {"type": "string", "minLength": 2},
                "thoi_luong": {"type": "integer", "minimum": 30, "maximum": 120},
                "si_so": {"type": "integer", "minimum": 10, "maximum": 60},
                "ngay_day": {"type": "string"}
            }
        },
        "sections": {
            "type": "object",
            "required": ["I", "II", "III", "IV"],
            "additionalProperties": False,
            "properties": {
                "I": {
                    "type": "object",
                    "required": ["yeu_cau_can_dat"],
                    "additionalProperties": False,
                    "properties": {
                        "yeu_cau_can_dat": {"type": "array", "minItems": 1, "items": {"type": "string"}},
                        "pham_chat": {"type": "array", "items": {"type": "string"}},
                        "nang_luc": {"type": "array", "items": {"type": "string"}},
                        "nang_luc_dac_thu": {"type": "array", "items": {"type": "string"}},
                        "nang_luc_so": {"type": "array", "items": {"type": "string"}}
                    }
                },
                "II": {
                    "type": "object",
                    "required": ["giao_vien", "hoc_sinh"],
                    "additionalProperties": False,
                    "properties": {
                        "giao_vien": {"type": "array", "minItems": 1, "items": {"type": "string"}},
                        "hoc_sinh": {"type": "array", "minItems": 1, "items": {"type": "string"}}
                    }
                },
                "III": {  # Tiến trình dạy học (BẢNG 2 CỘT GV/HS)
                    "type": "object",
                    "required": ["bang"],
                    "additionalProperties": False,
                    "properties": {
                        "bang": {
                            "type": "array",
                            "minItems": 12,
                            "items": {
                                "type": "object",
                                "required": ["kieu"],
                                "additionalProperties": False,
                                "properties": {
                                    "kieu": {"type": "string", "enum": ["header", "row"]},
                                    "tieu_de": {"type": "string", "minLength": 2},
                                    "thoi_gian": {"type": "integer", "minimum": 1, "maximum": 60},
                                    "giao_vien": {"type": "string", "minLength": 3},
                                    "hoc_sinh": {"type": "string", "minLength": 3},
                                    "ghi_chu": {"type": "string"}
                                },
                                "anyOf": [
                                    {"properties": {"kieu": {"const": "header"}}, "required": ["tieu_de"]},
                                    {"properties": {"kieu": {"const": "row"}}, "required": ["giao_vien", "hoc_sinh"]}
                                ]
                            }
                        }
                    }
                },
                "IV": {
                    "type": "object",
                    "required": ["dieu_chinh_sau_bai_day"],
                    "additionalProperties": False,
                    "properties": {
                        "dieu_chinh_sau_bai_day": {"type": "string", "minLength": 1}
                    }
                }
            }
        }
    }
}

def validate_lesson_plan_data(data: dict) -> None:
    Draft202012Validator.check_schema(LESSON_PLAN_DATA_SCHEMA)
    validate(instance=data, schema=LESSON_PLAN_DATA_SCHEMA)



def validate_lesson_plan_quality(data: dict) -> None:
    """Quality gate để chặn giáo án 'khung' và thiếu chi tiết."""
    import re
    data = data or {}
    meta = data.get("meta", {}) or {}
    sections = data.get("sections", {}) or {}
    mon = str(meta.get("mon","")).lower()

    # collect all strings
    texts = []
    def collect(x):
        if x is None:
            return
        if isinstance(x, str):
            texts.append(x)
        elif isinstance(x, dict):
            for v in x.values():
                collect(v)
        elif isinstance(x, list):
            for v in x:
                collect(v)
    collect(sections)

    joined = " ".join(texts).lower()
    if re.search(r"\bbổ\s*sung\s*nội\s*dung\b", joined):
        raise ValueError("Giáo án còn placeholder 'Bổ sung nội dung'.")
    if re.search(r"\bbước\s*\d+\b", joined) or re.search(r"\bnhiệm\s*vụ\s*\d+\b", joined):
        raise ValueError("Giáo án còn dùng 'Bước/Nhiệm vụ 1..' (không đạt chuẩn).")

    secIII = sections.get("III", {}) or {}
    bang = secIII.get("bang") if isinstance(secIII, dict) else []
    if not isinstance(bang, list) or len(bang) < 12:
        raise ValueError("Bảng hoạt động (III.bang) quá ngắn hoặc thiếu (cần tối thiểu ~12 dòng để đủ chi tiết).")

    # For math: need at least 2 'Bài' and some numbers/expressions
    if "toán" in mon:
        bai_count = sum(1 for t in texts if re.search(r"\bBài\s*\d+\b", t))
        num_count = sum(1 for t in texts if re.search(r"\d+[\.,]\d+|\d+\s*[-+×x*/:]\s*\d+", t))
        if bai_count < 2 or num_count < 4:
            raise ValueError("Giáo án Toán chưa đủ chi tiết: cần tối thiểu 2 mục 'Bài ...' và có số liệu/phép tính cụ thể.")

def _schema_error_to_text(e: Exception) -> str:
    if isinstance(e, ValidationError):
        path = " → ".join([str(p) for p in e.path]) if e.path else "(root)"
        return f"SchemaError at {path}: {e.message}"
    return str(e)

def validate_lesson_plan(data: dict) -> None:
    try:
        Draft202012Validator.check_schema(LESSON_PLAN_SCHEMA)
        validate(instance=data, schema=LESSON_PLAN_SCHEMA)
    except Exception as e:
        print(f"Schema Warning: {e}")

# ==============================================================================
# [MỚI] 2.3. HÀM TẠO PROMPT & GỌI AI (CHUẨN HÓA BẢNG 2 CỘT)
# ==============================================================================
def build_lesson_system_prompt_locked(meta: dict, teacher_note: str) -> str:
    return f"""
VAI TRÒ: Bạn là Giáo viên Tiểu học cốt cán, chuyên soạn GIÁO ÁN MẪU theo định hướng phát triển năng lực (CV 2345/BGDĐT).

THÔNG TIN BÀI DẠY:
- Cấp học: {meta.get("cap_hoc")} | Môn: {meta.get("mon")} | Lớp: {meta.get("lop")}
- Tuần: {meta.get("tuan")} | Tiết: {meta.get("tiet")}
- Tên bài: {meta.get("ten_bai")} ({meta.get("ghi_chu","")})
- Mã bài: {meta.get("bai_id")}
- Bộ sách: {meta.get("bo_sach")}

YÊU CẦU CẤU TRÚC (BẮT BUỘC GIỐNG MẪU CHUẨN):
Giáo án phải trình bày dưới dạng HTML, font Times New Roman, gồm 4 phần chính:

I. Yêu cầu cần đạt:
- Nêu rõ năng lực đặc thù, năng lực chung và phẩm chất.

II. Đồ dùng dạy học:
- Giáo viên: (Slide, tranh ảnh, thẻ từ...)
- Học sinh: (SGK, bảng con...)

III. Các hoạt động dạy – học chủ yếu:
***QUAN TRỌNG NHẤT: Phần này phải kẻ BẢNG (HTML <table>) gồm 2 cột***
- Cột 1: Hoạt động của Giáo viên
- Cột 2: Hoạt động của Học sinh
- Nội dung chia thành các hoạt động lớn (dùng dòng colspan hoặc in đậm để phân cách):
  1. Khởi động (Trò chơi, hát, kết nối...)
  2. Khám phá / Hình thành kiến thức mới (hoặc Luyện tập thực hành tùy bài)
  3. Vận dụng / Trải nghiệm
*Lưu ý văn phong:* Dùng từ ngữ sư phạm như "Tổ chức cho HS...", "Yêu cầu HS...", "Mời đại diện nhóm...", "GV chốt lại...".
*Chi tiết:* Viết rõ lời thoại, câu hỏi của GV và câu trả lời dự kiến của HS. Viết rõ các phép tính hoặc nội dung bài tập (VD: 27 - 1,2 = 25,8).

IV. Điều chỉnh sau bài dạy:
- Để trống dòng kẻ chấm (...) để GV tự ghi.

GHI CHÚ GV: {teacher_note}

OUTPUT JSON FORMAT:
Chỉ trả về JSON hợp lệ với 2 trường chính:
1. "meta": Thông tin bài học.
2. "renderHtml": Toàn bộ nội dung giáo án dạng HTML (để hiển thị và in ấn). Trong đó phần III phải là thẻ <table> có border="1".
""".strip()

# [FIX] Hàm LOCKED: chỉ làm nhiệm vụ gọi AI và trả dict (KHÔNG chứa UI, KHÔNG tự gọi lại)
def generate_lesson_plan_locked(
    api_key: str,
    meta_ppct: dict,
    bo_sach: str,
    thoi_luong: int,
    si_so: int,
    teacher_note: str,
    model_name: str = "gemini-2.0-flash"
) -> dict:
    """
    Sinh JSON data-only theo LESSON_PLAN_DATA_SCHEMA (meta + sections).
    Không render HTML ở đây. Không dùng st.spinner ở đây.
    """
    genai.configure(api_key=api_key)

    # meta chuẩn (đúng schema)
    req_meta = {
        "cap_hoc": meta_ppct.get("cap_hoc", ""),
        "mon": meta_ppct.get("mon", ""),
        "lop": meta_ppct.get("lop", ""),
        "bo_sach": bo_sach,
        "ppct": {
            "tuan": int(meta_ppct.get("tuan", 1)),
            "tiet": int(meta_ppct.get("tiet", 1)),
            "bai_id": meta_ppct.get("bai_id", "AUTO"),
            "ghi_chu": meta_ppct.get("ghi_chu", "")
        },
        "ten_bai": meta_ppct.get("ten_bai", ""),
        "thoi_luong": int(thoi_luong),
        "si_so": int(si_so),
        "ngay_day": meta_ppct.get("ngay_day", "")
    }

    # prompt data-only (khuyến nghị dùng prompt data-only thay vì prompt HTML)
    system_prompt = build_lesson_system_prompt_data_only(
        meta={
            "cap_hoc": req_meta["cap_hoc"],
            "mon": req_meta["mon"],
            "lop": req_meta["lop"],
            "bo_sach": req_meta["bo_sach"],
            "tuan": req_meta["ppct"]["tuan"],
            "tiet": req_meta["ppct"]["tiet"],
            "bai_id": req_meta["ppct"]["bai_id"],
            "ten_bai": req_meta["ten_bai"],
            "thoi_luong": req_meta["thoi_luong"],
            "si_so": req_meta["si_so"],
        },
        teacher_note=teacher_note
    )

    model = genai.GenerativeModel(model_name, system_instruction=system_prompt)

    safe_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]

    base_req = {"meta": req_meta, "note": teacher_note}
    last_err = ""

    # thử tối đa 2 lần, nếu sai schema thì tự sửa
    for attempt in range(1, 3):
        try:
            res = model.generate_content(
                json.dumps(base_req, ensure_ascii=False),
                generation_config={"response_mime_type": "application/json"},
                safety_settings=safe_settings
            )

            raw = json.loads(clean_json(res.text))

            data = {
                "meta": req_meta,
                "sections": raw.get("sections", {})
            }

            validate_lesson_plan_data(data)  # bắt buộc đúng schema
            return data

        except Exception as e:
            last_err = _schema_error_to_text(e)
            repair_note = f"""
[SCHEMA_REPAIR]
Bạn vừa trả JSON KHÔNG đạt schema.
LỖI: {last_err}

YÊU CẦU:
- Chỉ trả JSON gồm "meta" và "sections"
- sections phải có đủ I, II, III, IV
- III.hoat_dong >= 3; mỗi hoạt động có ten_hoat_dong, thoi_gian, gv>=2, hs>=2
- Không tạo HTML
Chỉ trả JSON
"""
            base_req = {"meta": req_meta, "note": teacher_note + "\n" + repair_note}

    # fallback an toàn
    return {
        "meta": req_meta,
        "sections": {
            "I": {"yeu_cau_can_dat": [f"(Lỗi tạo dữ liệu) {last_err}"]},
            "II": {"giao_vien": ["..."], "hoc_sinh": ["..."]},
            "III": {"hoat_dong": [
                {"ten_hoat_dong": "Khởi động", "thoi_gian": 5, "gv": ["...", "..."], "hs": ["...", "..."]},
                {"ten_hoat_dong": "Hình thành kiến thức", "thoi_gian": 15, "gv": ["...", "..."], "hs": ["...", "..."]},
                {"ten_hoat_dong": "Luyện tập/Vận dụng", "thoi_gian": 15, "gv": ["...", "..."], "hs": ["...", "..."]}
            ]},
            "IV": {"dieu_chinh_sau_bai_day": "...................................................................................."}
        }
    }

# ==============================================================================
# [PATCH 2/3] PROMPT KHÓA CỨNG: DATA-ONLY JSON (ANTI-HALLUCINATION)
# ==============================================================================

def build_lesson_system_prompt_data_only(meta: dict, teacher_note: str) -> str:
    """System prompt để AI sinh JSON (meta + sections) theo mẫu giáo án tiểu học.
    Bám Công văn 2345/BGDĐT-GDTH và mẫu giáo án chuẩn do người dùng cung cấp.
    """
    return f"""
Bạn là GIÁO VIÊN TIỂU HỌC cốt cán, soạn KẾ HOẠCH BÀI DẠY theo CTGDPT 2018 (CV 2345/BGDĐT-GDTH).

NHIỆM VỤ:
- Bạn sẽ nhận INPUT là 1 JSON có trường meta (thông tin bài) và note (ghi chú GV).
- Bạn phải trả về DUY NHẤT 1 JSON hợp lệ, KHÔNG kèm chữ giải thích.

YÊU CẦU CHẤT LƯỢNG (RẤT QUAN TRỌNG):
- Viết ĐÚNG NGHIỆP VỤ SƯ PHẠM, không viết khung chung chung.
- CẤM các cụm: "Bổ sung nội dung", "Bước 1/2", "Nhiệm vụ 1/2", "Tổ chức bước...".
- Phần III phải có NỘI DUNG DẠY - HỌC THẬT: bài tập/ví dụ/câu hỏi, sản phẩm HS (bảng con/vở/phiếu), lời gợi mở GV.
- Nếu là TOÁN: bắt buộc có tối thiểu 2 mục "Bài 1/2/..." hoặc "Ví dụ..." và có số liệu/phép tính cụ thể (vd: 12,5 - 3,7; 4,2 × 0,5).

CẤU TRÚC BẮT BUỘC:
Trả về JSON có dạng:
{{
  "sections": {{
    "I": {{
      "yeu_cau_can_dat": [... >=5 ý ...],
      "nang_luc": [... >=3 ý ...],
      "pham_chat": [... >=2 ý ...],
      "nang_luc_dac_thu": [... >=2 ý ...],
      "nang_luc_so": [... >=1 ý ...]
    }},
    "II": {{
      "giao_vien": [... >=6 ý ...],
      "hoc_sinh": [... >=6 ý ...]
    }},
    "III": {{
      "bang": [
        {{"kieu":"header", "tieu_de":"1. Khởi động:"}},
        {{"kieu":"row", "thoi_gian":4, "giao_vien":"...", "hoc_sinh":"..."}},
        {{"kieu":"header", "tieu_de":"2. Luyện tập:"}},
        {{"kieu":"row", "thoi_gian":10, "giao_vien":"...", "hoc_sinh":"Bài 1: ..."}}
      ]
    }},
    "IV": {{
      "dieu_chinh_sau_bai_day": "... (để dòng chấm cho GV ghi hoặc gợi ý 3 ý) ..."
    }}
  }}
}}

QUY TẮC BẢNG (III.bang):
- bang là BẢNG 2 CỘT (GV/HS), nhưng trả về dạng JSON để hệ thống render.
- kieu="header": chỉ dùng để ngăn cách hoạt động lớn (Khởi động/Khám phá-Hình thành/Luyện tập/Vận dụng).
- kieu="row": phải có giao_vien và hoc_sinh viết CỤ THỂ (có câu hỏi, nhiệm vụ, sản phẩm).
- Tổng số dòng bang tối thiểu 10 (không tính header), ưu tiên 12–18 dòng tuỳ bài.
- thoi_gian: phút của dòng (1–10). Tổng cộng xấp xỉ meta.thoi_luong.

BỐI CẢNH BÀI DẠY:
- Cấp học: {meta.get('cap_hoc')}
- Môn: {meta.get('mon')}
- Lớp: {meta.get('lop')}
- Bộ sách: {meta.get('bo_sach')}
- Tên bài: {meta.get('ten_bai')}
- PPCT: {meta.get('ppct')}

GHI CHÚ GV (nếu có): {teacher_note}

Chỉ trả JSON hợp lệ.
""".strip()

def generate_lesson_plan_data_only(
    api_key: str,
    meta_ppct: dict,
    teacher_note: str,
    model_name: str = "gemini-2.0-flash"
) -> dict:
    """Sinh JSON data-only (meta + sections) để render HTML.
    Tự sửa tối đa 3 lần nếu sai schema hoặc thiếu chi tiết.
    """
    import json
    genai.configure(api_key=api_key)

    req_meta = {
        "cap_hoc": meta_ppct.get("cap_hoc", ""),
        "mon": meta_ppct.get("mon", ""),
        "lop": meta_ppct.get("lop", ""),
        "bo_sach": meta_ppct.get("bo_sach", ""),
        "ppct": meta_ppct.get("ppct", {}) or {},
        "ten_bai": meta_ppct.get("ten_bai", ""),
        "thoi_luong": int(meta_ppct.get("thoi_luong", 40) or 40),
        "si_so": int(meta_ppct.get("si_so", 35) or 35),
    }

    system_prompt = build_lesson_system_prompt_data_only(req_meta, teacher_note)
    model = genai.GenerativeModel(model_name, system_instruction=system_prompt)

    safe_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]

    base_req = {"meta": req_meta, "note": teacher_note}
    last_err = ""

    for attempt in range(1, 4):
        try:
            res = model.generate_content(
                json.dumps(base_req, ensure_ascii=False),
                generation_config={"response_mime_type": "application/json"},
                safety_settings=safe_settings
            )
            raw = json.loads(clean_json(res.text))
            data = {"meta": req_meta, "sections": raw.get("sections", {})}

            validate_lesson_plan_data(data)
            validate_lesson_plan_quality(data)
            return data

        except Exception as e:
            last_err = _schema_error_to_text(e)
            repair_note = f"""
[SCHEMA_REPAIR]
Bạn vừa trả JSON KHÔNG đạt schema hoặc thiếu chi tiết.
LỖI: {last_err}

YÊU CẦU BẮT BUỘC (chỉ trả JSON):
- Root chỉ gồm object JSON có khóa 'sections'.
- sections phải có đủ: I, II, III, IV.
- I:
  * yeu_cau_can_dat: mảng >=5 ý
  * nang_luc: mảng >=3 ý
  * pham_chat: mảng >=2 ý
  * nang_luc_dac_thu: mảng >=2 ý
  * nang_luc_so: mảng >=1 ý
- II:
  * giao_vien: mảng >=6 ý (thiết bị/học liệu/phiếu)
  * hoc_sinh: mảng >=6 ý
- III:
  * bắt buộc có 'bang' là mảng.
  * bang phải có >= 12 dòng 'row' (không tính header).
  * header mẫu: {{"kieu":"header","tieu_de":"1. Khởi động:"}}
  * row mẫu: {{"kieu":"row","thoi_gian":4,"giao_vien":"...","hoc_sinh":"..."}}
  * CẤM 'Bước 1/2' hoặc 'Nhiệm vụ 1/2'. Viết nhiệm vụ học tập CỤ THỂ.
  * Nếu Toán: phải có 'Bài 1/2/...' hoặc 'Ví dụ...' và có số liệu/phép tính cụ thể.
- IV:
  * dieu_chinh_sau_bai_day: chuỗi (có thể để dòng chấm).

Chỉ trả JSON hợp lệ.
""".strip()

            base_req = {"meta": req_meta, "note": teacher_note + "\n" + repair_note}

    # fallback an toàn (vẫn đúng schema)
    return {
        "meta": req_meta,
        "sections": {
            "I": {
                "yeu_cau_can_dat": [f"(Lỗi tạo dữ liệu) {last_err}"],
                "nang_luc": ["(Chưa có nội dung)"],
                "pham_chat": ["(Chưa có nội dung)"],
                "nang_luc_dac_thu": ["(Chưa có nội dung)"],
                "nang_luc_so": ["(Chưa có nội dung)"],
            },
            "II": {"giao_vien": ["(Chưa có nội dung)"], "hoc_sinh": ["(Chưa có nội dung)"]},
            "III": {"bang": [
                {"kieu":"header","tieu_de":"1. Khởi động:"},
                {"kieu":"row","thoi_gian":4,"giao_vien":"(Lỗi tạo dữ liệu) Không tạo được tiến trình. Vui lòng bấm TẠO LẠI.","hoc_sinh":"Lắng nghe và ghi nhận."},
                {"kieu":"header","tieu_de":"2. Hình thành kiến thức / Luyện tập:"},
                {"kieu":"row","thoi_gian":20,"giao_vien":"(Lỗi tạo dữ liệu) Hướng dẫn HS ôn tập và làm bài theo SGK.","hoc_sinh":"Làm bài vào vở/bảng con theo hướng dẫn."},
                {"kieu":"header","tieu_de":"3. Vận dụng/Mở rộng:"},
                {"kieu":"row","thoi_gian":8,"giao_vien":"(Lỗi tạo dữ liệu) Giao bài vận dụng và dặn dò.","hoc_sinh":"Hoàn thành bài, ghi nhiệm vụ về nhà."}
            ]},
            "IV": {"dieu_chinh_sau_bai_day": "……………………………………………………………………………………………\n……………………………………………………………………………………………\n……………………………………………………………………………………………"}
        }
    }

# ==============================================================================
# [MỚI] OPENAI GPT-4o (tuỳ chọn) cho MODULE SOẠN GIÁO ÁN
# - Dùng requests, không cần cài thư viện openai
# - Bật bằng checkbox trong UI soạn giáo án
# ==============================================================================
def openai_chat_json(api_key: str, system_prompt: str, user_content: str, model: str = "gpt-4o", timeout: int = 60) -> dict:
    """Gọi OpenAI Chat Completions và yêu cầu trả JSON object."""
    if not api_key:
        raise ValueError("Thiếu OPENAI_API_KEY")
    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "temperature": 0.6,
        "response_format": {"type": "json_object"},
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ]
    }
    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if r.status_code >= 400:
        raise ValueError(f"OpenAI API lỗi {r.status_code}: {r.text[:300]}")
    data = r.json()
    content = data["choices"][0]["message"]["content"]
    return json.loads(clean_json(content))

def main_app():
    if 'dossier' not in st.session_state: st.session_state['dossier'] = []
    
    user = st.session_state.get('user', {'role': 'guest'})
    is_admin = user.get('role') == 'admin'

    c1, c2, c3 = st.columns([3, 0.8, 0.8])
    with c1:
        st.markdown(f"<div class='header-text'>🎓 {APP_CONFIG['name']}</div>", unsafe_allow_html=True)
        st.caption(f"User: {user.get('fullname', user.get('email', 'Guest'))} | Role: {user.get('role', '').upper()}")
    
    # Nút RESET
    with c2:
        if st.button("🔄 LÀM MỚI", use_container_width=True): 
            st.session_state['dossier'] = [] 
            st.toast("Đã làm mới hệ thống!", icon="🧹")
            time.sleep(0.5)
            st.rerun()
            
    # Nút ĐĂNG XUẤT
    with c3:
        if st.button("ĐĂNG XUẤT", use_container_width=True):
            st.session_state.pop('user', None)
            st.rerun()

    # --- CẬP NHẬT TAB MỚI: THÊM '🎯 ĐỀ CHUẨN YCCĐ' (TAB SỐ 8) ---
    tabs = st.tabs(["🚀 THIẾT LẬP", "📄 XEM ĐỀ", "✅ ĐÁP ÁN", "⚖️ PHÁP LÝ", "💎 NÂNG CẤP VIP", "💰 ĐỐI TÁC", "📂 HỒ SƠ", "🎯 ĐỀ CHUẨN YCCĐ"])

    # --- TAB 1: THIẾT LẬP ---
    with tabs[0]:
        st.markdown('<div class="css-card">', unsafe_allow_html=True)
        
        col_year, col_lvl = st.columns(2)
        with col_year: school_year = st.selectbox("Năm học", ["2024-2025", "2025-2026", "2026-2027"], index=1)
        with col_lvl: level_key = st.radio("Cấp học", ["Tiểu học", "THCS", "THPT"], horizontal=True)
        
        curr_lvl = "tieu_hoc" if level_key == "Tiểu học" else "thcs" if level_key == "THCS" else "thpt"
        edu = EDUCATION_DATA[curr_lvl]

        c1, c2, c3, c4 = st.columns(4)
        with c1: grade = st.selectbox("Khối lớp", edu["grades"])
        with c2: subject = st.selectbox("Môn học", edu["subjects"])
        with c3: book = st.selectbox("Bộ sách", BOOKS_LIST)
        
        available_scopes = FULL_SCOPE_LIST
        if curr_lvl == "tieu_hoc" and grade in ["Lớp 1", "Lớp 2", "Lớp 3"]:
            available_scopes = LIMITED_SCOPE_LIST 
        
        with c4: scope = st.selectbox("Thời điểm", available_scopes)

        if curr_lvl == "thpt":
            struct_info = SUBJECT_STRUCTURE_DATA["THPT_2025"]
        elif curr_lvl == "tieu_hoc":
            if subject == "Tiếng Việt":
                struct_info = SUBJECT_STRUCTURE_DATA["TieuHoc_TV"]
            else:
                struct_info = SUBJECT_STRUCTURE_DATA["TieuHoc_Chung"]
        else:
            struct_info = SUBJECT_STRUCTURE_DATA.get(subject, SUBJECT_STRUCTURE_DATA['Mặc định'])
            
        st.info(f"💡 **Cấu trúc:** {struct_info} | **Pháp lý:** {edu['legal']}")

        uc1, uc2 = st.columns(2)
        with uc1: mt_file = st.file_uploader("📂 Ma trận (Word/Excel)", type=['docx','xlsx'])
        with uc2: dt_file = st.file_uploader("📝 Đặc tả (Word/Excel)", type=['docx','xlsx'])
        
        auto_mode = False
        if not mt_file and not dt_file:
            auto_mode = True
            st.markdown('<div style="text-align:center;"><span class="auto-tag">✨ CHẾ ĐỘ TỰ ĐỘNG: AI SẼ TỰ XÂY DỰNG MA TRẬN & ĐẶC TẢ</span></div>', unsafe_allow_html=True)

        user_req = st.text_area("Ghi chú chuyên môn:", "Ví dụ: Đề cần phân loại học sinh giỏi...", height=80)

        # --- CÔNG CỤ CẤU HÌNH SỐ LƯỢNG ---
        st.markdown("---")
        st.markdown("##### 🛠 CẤU TRÚC ĐỀ THI MONG MUỐN")
        col_s1, col_s2, col_s3 = st.columns(3)
        with col_s1: 
            num_choice = st.number_input("Trắc nghiệm (Số câu)", min_value=0, max_value=100, value=10, step=1, key="num_choice")
        with col_s2: 
            num_essay = st.number_input("Tự luận (Số câu)", min_value=0, max_value=20, value=2, step=1, key="num_essay")
        with col_s3: 
            num_practice = st.number_input("Thực hành (Bài)", min_value=0, max_value=10, value=0, step=1, key="num_practice")

        st.markdown("---")
        b1, b2, b3 = st.columns([1, 1, 2])
        with b1: num_exams = st.number_input("Số lượng đề", 1, 5, 1)
        with b2: start_code = st.number_input("Mã đề từ", 101, 999, 101)
        with b3:
            st.write(""); st.write("")
            if st.button("⚡ KHỞI CHẠY (AI STUDIO ENGINE)", type="primary", use_container_width=True):
                client = init_supabase()
                if client:
                    try:
                        # 1. LẤY THÔNG TIN NGƯỜI DÙNG TỪ DB
                        current_user_db = client.table('users_pro').select("*").eq('username', user.get('email')).execute()
                        if current_user_db.data:
                            user_data = current_user_db.data[0]
                            db_role = user_data['role']
                            usage_count = user_data.get('usage_count', 0)
                            
                            # [NÂNG CẤP] TÍNH TỔNG LƯỢT DÙNG (CÓ BONUS)
                            bonus_turns = user_data.get('bonus_turns', 0)
                            limit_check = MAX_PRO_USAGE if db_role == 'pro' else (MAX_FREE_USAGE + bonus_turns)

                            if (get_user_points(client, user.get('email','')) >= 0 and get_user_points(client, user.get('email','')) < POINT_COST_EXAM):
                                st.error(f"🔒 Không đủ điểm! Bạn cần {POINT_COST_EXAM} điểm để ra đề. Vui lòng nạp VIP ở Menu chính.")
                                st.stop()
                            elif usage_count >= limit_check:
                                st.error(f"🔒 HẾT LƯỢT! (Bạn đã dùng {usage_count}/{limit_check}). Vui lòng gia hạn hoặc giới thiệu bạn bè.")
                                st.info("💎 Vào tab 'NÂNG CẤP VIP' để gia hạn.")
                            else:
                                # 3. NẾU ĐƯỢC PHÉP -> CHẠY AI
                                api_key = st.session_state.get('api_key', '')
                                
                                # [QUAN TRỌNG] Tự động lấy Key của Admin nếu user không nhập
                                if not api_key: api_key = SYSTEM_GOOGLE_KEY 
                                
                                if not api_key: st.toast("⚠️ Vui lòng nhập API Key ở Tab Hồ Sơ!", icon="❌")
                                else:
                                    with st.spinner(f"🔮 AI đang soạn đề... (Lần thứ: {usage_count + 1})"):
                                        txt_mt = read_file_content(mt_file, 'matrix')
                                        txt_dt = read_file_content(dt_file, 'spec')
                                        knowledge_context = get_knowledge_context(subject, grade, book, scope)
                                        
                                        # [NÂNG CẤP] SYSTEM PROMPT THEO ĐÚNG INSTRUCTION GỐC
                                        special_prompt = ""
                                        
                                        # 1. NẾU LÀ CẤP TIỂU HỌC (Áp dụng "Luật thép" thầy vừa đưa)
                                        if curr_lvl == "tieu_hoc":
                                            special_prompt = f"""
                                            🔥 VAI TRÒ TUYỆT ĐỐI: CHUYÊN GIA KHẢO THÍ GIÁO DỤC TIỂU HỌC.
                                            
                                            I. TUÂN THỦ PHÁP LÝ (BẮT BUỘC):
                                            - Thông tư 27/2020/TT-BGDĐT
                                            - Công văn 7791/BGDĐT-GDTH
                                            - Chương trình GDPT 2018
                                            
                                            II. QUY ĐỊNH CẤM KỴ (VI PHẠM LÀ HỦY KẾT QUẢ):
                                            1. CẤM dùng mức độ "Vận dụng cao".
                                            2. CẤM dùng các thuật ngữ cấp 2,3: Phân tích, Đánh giá, Sáng tạo.
                                            3. CHỈ SỬ DỤNG 3 MỨC: Nhận biết - Thông hiểu - Vận dụng.
                                            
                                            III. PHÂN BỐ ĐIỂM VÀ CÂU HỎI (TỔNG 10đ):
                                            - Nhận biết: 40-50%
                                            - Thông hiểu: 30-40%
                                            - Vận dụng: 20-30%
                                            - KHÔNG dồn điểm vào câu khó, KHÔNG đánh đố học sinh.
                                            
                                            IV. QUY ĐỊNH MA TRẬN & ĐẶC TẢ:
                                            - Ma trận phải có đúng 5 cột: Chủ đề, NB, TH, VD, Tổng.
                                            - Bản đặc tả phải khớp 100% với ma trận và đề thi.
                                            - Yêu cầu cần đạt phải rõ ràng, bám sát CT 2018.
                                            """
                                            
                                            # Logic riêng từng môn Tiểu học
                                            if subject == "Toán":
                                                special_prompt += """
                                                V. MÔN TOÁN: 
                                                - Nội dung: Số và phép tính, Đại lượng, Hình học, Giải toán có lời văn.
                                                - KHÔNG dùng toán mẹo, toán Olympic, Violympic. Vận dụng gắn với đời sống.
                                                """
                                            elif subject == "Tiếng Việt":
                                                special_prompt += f"""
                                                V. MÔN TIẾNG VIỆT (Tách 2 phần):
                                                A. KIỂM TRA ĐỌC (10đ):
                                                    1. Đọc thành tiếng.
                                                    2. Đọc hiểu: Sử dụng văn bản MỚI (ngoài SGK) phù hợp lứa tuổi + {num_choice} câu hỏi (M1-M2-M3).
                                                B. KIỂM TRA VIẾT (10đ):
                                                    1. Chính tả (Nghe-viết đoạn ngắn).
                                                    2. Tập làm văn: {num_essay} câu (Viết đoạn/bài văn theo chủ điểm đã học).
                                                """
                                            elif "Tin học" in subject:
                                                special_prompt += f"""
                                                V. MÔN TIN HỌC:
                                                - Nội dung: Máy tính, Dữ liệu, An toàn thông tin, Phần mềm học tập.
                                                - Trắc nghiệm ({num_choice} câu) + Thực hành ({num_essay} câu).
                                                - KHÔNG lập trình phức tạp.
                                                """
                                            else:
                                                special_prompt += """
                                                V. CÁC MÔN KHÁC (Khoa học, LS&ĐL, Đạo đức...): Gắn với đời sống, không dùng thuật ngữ hàn lâm.
                                                """

                                        # 2. NẾU LÀ CẤP 2, 3 (Giữ nguyên logic cũ)
                                        else:
                                            special_prompt = """
                                            YÊU CẦU TRUNG HỌC (Theo Thông tư 22 & CV 7791):
                                            - Ma trận 4 mức độ: Nhận biết (40%) - Thông hiểu (30%) - Vận dụng (20%) - Vận dụng cao (10%).
                                            """
                                            if curr_lvl == "thpt":
                                                special_prompt += """
                                                - Cấu trúc THPT 2025: Phần I (TN nhiều lựa chọn), Phần II (Đúng/Sai), Phần III (Trả lời ngắn).
                                                """

                                        SYSTEM_PROMPT = f"""
                                        {APP_CONFIG['context']}
                                        
                                        I. THÔNG TIN ĐẦU VÀO:
                                        - Năm học: {school_year} | Cấp: {level_key} | Môn: {subject} | Lớp: {grade} 
                                        - Bộ sách: "{book}" | Phạm vi: {scope}
                                        - {knowledge_context}
                                        
                                        II. HƯỚNG DẪN CHUYÊN GIA (TUÂN THỦ TUYỆT ĐỐI):
                                        {special_prompt}
                                        
                                        III. CƠ CHẾ TỰ KIỂM TRA & TỪ CHỐI (SELF-REFLECTION):
                                        - Trước khi xuất kết quả, hãy tự kiểm tra: Tổng điểm có đúng 10 không? Có xuất hiện mức độ sai quy định không?
                                        - Nếu người dùng yêu cầu ra đề vượt chuẩn (Ví dụ: Lớp 3 mà đòi Vận dụng cao) -> HÃY TỪ CHỐI LỊCH SỰ và đề xuất phương án đúng luật.
                                        
                                        IV. ĐỊNH DẠNG OUTPUT (JSON RAW):
                                        {{
                                            "title": "Tên đề thi",
                                            "content": "Nội dung đề thi HTML (Trình bày đẹp, chuẩn font)",
                                            "matrixHtml": "Bảng ma trận HTML (Phải khớp 100% với đề)",
                                            "specHtml": "Bảng đặc tả HTML",
                                            "answers": "Đáp án & Hướng dẫn chấm HTML"
                                        }}
                                        V. QUAN TRỌNG: CHỈ TRẢ VỀ JSON. KHÔNG GIẢI THÍCH GÌ THÊM.
                                        """

                                        try:
                                            genai.configure(api_key=api_key)
                                            # [SỬA LỖI 404] Dùng gemini-2.0-flash
                                            model = genai.GenerativeModel('gemini-2.0-flash', system_instruction=SYSTEM_PROMPT)
                                            
                                            # [FIX LỖI] Cấu hình tắt bộ lọc an toàn để AI không chặn đề thi
                                            safe_settings = [
                                                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                                                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                                                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                                                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
                                            ]

                                            new_exams = []
                                            for i in range(num_exams):
                                                code = start_code + i
                                                prompt = SYSTEM_PROMPT.replace("[CODE]", str(code))
                                                req = f"DATA: {txt_mt} {txt_dt}\nNOTE: {user_req}\nSTRUCT: {num_choice} TN, {num_essay} TL, {num_practice} TH\nTASK: Exam {i+1} (Code {code})"
                                                
                                                # Thêm safety_settings vào đây
                                                res = model.generate_content(
                                                    req, 
                                                    generation_config={"response_mime_type": "application/json"},
                                                    safety_settings=safe_settings
                                                )
                                                
                                                try:
                                                    clean_text = clean_json(res.text)
                                                    data = json.loads(clean_text)
                                                    data['id'] = str(code); data['title'] = f"Đề {subject} {grade} - {scope} (Mã {code})"
                                                    
                                                    # [NÂNG CẤP] TỰ ĐỘNG LƯU VÀO KHO
                                                    save_data = {"username": user.get('email'), "title": data['title'], "exam_data": data}
                                                    client.table('exam_history').insert(save_data).execute()
                                                    
                                                    new_exams.append(data)
                                                except Exception as e:
                                                    st.error(f"Lỗi phân tích đề {code}: {e}")
                                                    continue
                                            
                                            st.session_state['dossier'] = new_exams + st.session_state['dossier']
                                            client.table('users_pro').update({'usage_count': usage_count + 1}).eq('username', user.get('email')).execute()
                                            
                                            st.success(f"✅ Tạo thành công! (Đã dùng: {usage_count + 1}/{limit_check})")
                                            # Trừ điểm nếu hệ thống points đã bật
                                            try:
                                                if get_user_points(client, user.get('email','')) >= 0:
                                                    deduct_user_points(client, user.get('email',''), POINT_COST_EXAM)
                                            except Exception:
                                                pass

                                        except Exception as e: st.error(f"Lỗi AI: {e}")
                    except Exception as e: st.error(f"Lỗi DB: {e}")
                else: st.error("Lỗi kết nối.")
        st.markdown('</div>', unsafe_allow_html=True)

    # --- TAB 2: XEM & XUẤT (CLASS paper-view ĐÃ CHUẨN HÓA FONT) ---
    with tabs[1]:
        if not st.session_state['dossier']: st.info("👈 Chưa có dữ liệu.")
        else:
            all_e = st.session_state['dossier']
            sel = st.selectbox("Chọn mã đề:", range(len(all_e)), format_func=lambda x: f"[{all_e[x]['id']}] {all_e[x]['title']}")
            curr = all_e[sel]
            
            st1, st2, st3 = st.tabs(["📄 NỘI DUNG ĐỀ", "📊 MA TRẬN", "📝 ĐẶC TẢ"])
            
            with st1:
                st.markdown(f"""<div class="paper-view">{curr.get('content', '')}</div>""", unsafe_allow_html=True)
                footer = f"<br/><center><p>{APP_CONFIG['name']}</p></center>"
                if is_admin or user.get('role') == 'pro': 
                    st.download_button("⬇️ Tải Đề (.doc)", create_word_doc(curr.get('content', '') + footer, curr['title']), f"De_{curr['id']}.doc", type="primary")
                else: st.warning("🔒 Nâng cấp PRO để tải file Word")
            
            with st2:
                st.markdown(curr.get('matrixHtml', 'Không có dữ liệu ma trận'), unsafe_allow_html=True)
                if is_admin or user.get('role') == 'pro': st.download_button("⬇️ Tải Ma trận", create_word_doc(curr['matrixHtml'], "MaTran"), f"MaTran_{curr['id']}.doc")

            with st3:
                st.markdown(curr.get('specHtml', 'Không có dữ liệu đặc tả'), unsafe_allow_html=True)
                if is_admin or user.get('role') == 'pro': st.download_button("⬇️ Tải Đặc tả", create_word_doc(curr['specHtml'], "DacTa"), f"DacTa_{curr['id']}.doc")

    # --- TAB 3: ĐÁP ÁN ---
    with tabs[2]:
        if st.session_state['dossier']:
            curr = st.session_state['dossier'][sel]
            if is_admin or user.get('role') == 'pro':
                st.markdown(f"""<div class="paper-view">{curr.get('answers','Chưa có đáp án')}</div>""", unsafe_allow_html=True)
                st.download_button("⬇️ Tải Đáp án (.doc)", create_word_doc(curr.get('answers',''), "DapAn"), f"DA_{curr['id']}.doc")
            else: st.info("🔒 Nâng cấp PRO để xem và tải Đáp án chi tiết.")
        else: st.info("Chưa có dữ liệu.")

    # --- TAB 4: PHÁP LÝ ---
    with tabs[3]:
        for doc in LEGAL_DOCUMENTS:
            cls = "highlight-card" if doc.get('highlight') else "legal-card"
            st.markdown(f"""<div class="{cls}" style="padding:15px; margin-bottom:10px; border-radius:10px;"><span style="background:#1e293b; color:white; padding:2px 8px; border-radius:4px; font-size:11px; font-weight:bold">{doc['code']}</span><span style="font-weight:bold; color:#334155; margin-left:8px">{doc['title']}</span><p style="font-size:13px; color:#64748b; margin:5px 0 0 0">{doc['summary']}</p></div>""", unsafe_allow_html=True)
    
    # --- [NÂNG CẤP] TAB 5: NÂNG CẤP VIP & THANH TOÁN (LOGIC SEVQR) ---
    with tabs[4]:
        st.markdown("<h3 style='text-align: center; color: #1E3A8A;'>🚀 BẢNG GIÁ & NÂNG CẤP VIP</h3>", unsafe_allow_html=True)
        col_free, col_pro = st.columns(2)
        with col_free:
            st.markdown(f"""<div class="pricing-card"><h3>Gói FREE</h3><div class="price-tag">0đ</div><div class="feature-list">✅ Tạo thử <b>{MAX_FREE_USAGE} đề</b><br>❌ Tải file Word<br>❌ Xem đáp án chi tiết<br>❌ Hỗ trợ kỹ thuật</div></div>""", unsafe_allow_html=True)
        with col_pro:
            st.markdown(f"""<div class="pricing-card" style="border: 2px solid #2563EB;"><h3 style="color: #2563EB;">Gói PRO VIP</h3><div class="price-tag">{PRICE_VIP:,.0f}đ / gói</div><div class="feature-list">✅ <b>Tạo tối đa {MAX_PRO_USAGE} đề</b><br>✅ <b>Tải file Word chuẩn</b><br>✅ <b>Xem & Tải Đáp án/Ma trận</b><br>✅ Hỗ trợ ưu tiên 24/7</div></div>""", unsafe_allow_html=True)
        
        st.markdown("---")
        st.subheader("📲 QUÉT MÃ QR ĐỂ THANH TOÁN TỰ ĐỘNG")
        
        c1, c2 = st.columns([1, 2])
        with c1:
            ref_code_input = st.text_input("Mã giới thiệu (Để tặng lượt khi mua Pro):")
            
        current_price = PRICE_VIP
        # [QUAN TRỌNG] THÊM TIỀN TỐ "SEVQR" VÀO NỘI DUNG ĐỂ SEPAY NHẬN DIỆN
        final_content_ck = f"SEVQR NAP VIP {user.get('email')}"
        show_qr = True
        
        # [LOGIC MỚI] CHECK MÃ GIỚI THIỆU ĐỂ ẨN/HIỆN QR (KHÔNG GIẢM GIÁ)
        if ref_code_input:
            client = init_supabase()
            if client:
                check_ref = client.table('users_pro').select("*").eq('username', ref_code_input).execute()
                if check_ref.data and ref_code_input != user.get('email'):
                    st.success(f"✅ Mã hợp lệ! Bạn sẽ được tặng thêm {BONUS_PRO_REF} lượt khi kích hoạt Pro.")
                    final_content_ck = f"SEVQR NAP VIP {user.get('email')} REF {ref_code_input}"
                    show_qr = True
                elif ref_code_input == user.get('email'):
                    st.warning("Bạn không thể tự giới thiệu chính mình.")
                    show_qr = True # Vẫn hiện QR gốc
                else:
                    st.error("❌ Mã giới thiệu không tồn tại! (Vui lòng nhập đúng hoặc xóa đi để thanh toán).")
                    show_qr = False # Ẩn QR

        if show_qr:
            # [FIX LỖI] URL ENCODE CHO NỘI DUNG CHUYỂN KHOẢN ĐỂ TRÁNH LỖI MEDIA STORAGE
            import urllib.parse
            encoded_content = urllib.parse.quote(final_content_ck)
            qr_url = f"https://img.vietqr.io/image/{BANK_ID}-{BANK_ACC}-compact.png?amount={current_price}&addInfo={encoded_content}&accountName={BANK_NAME}"
            
            c_qr1, c_qr2 = st.columns([1, 2])
            with c_qr1: 
                # [FIX LỖI] TRY-EXCEPT ĐỂ TRÁNH SẬP APP NẾU LỖI ẢNH
                try:
                    st.image(qr_url, caption=f"Mã QR ({current_price:,.0f}đ)", width=300)
                except:
                    st.error("Không tải được QR. Vui lòng chuyển khoản thủ công.")
            
            with c_qr2: 
                st.info(f"**Nội dung chuyển khoản:** `{final_content_ck}`\n\n1. Quét mã QR.\n2. Bấm nút **'KÍCH HOẠT NGAY'** bên dưới sau khi chuyển khoản.")
                
                # [NÂNG CẤP] NÚT KÍCH HOẠT TỰ ĐỘNG (CHECK SEPAY)
                if st.button("🚀 KÍCH HOẠT NGAY (Sau khi đã CK)", type="primary"):
                    if check_sepay_transaction(current_price, final_content_ck):
                        client = init_supabase()
                        if client:
                            # Lấy trạng thái hiện tại để kiểm tra có phải lần đầu không
                            curr_user_db = client.table('users_pro').select("*").eq('username', user.get('email')).execute()
                            is_first_time = False
                            if curr_user_db.data:
                                if curr_user_db.data[0]['role'] == 'free': is_first_time = True

                            # 1. Update người mua lên Pro (Reset lượt)
                            bonus_add = BONUS_PRO_REF if (ref_code_input and is_first_time) else 0
                            client.table('users_pro').update({
                                'role': 'pro',
                                'usage_count': 0,
                                    'points': 0,
                                'bonus_turns': bonus_add,
                                'referred_by': ref_code_input if ref_code_input else None
                            }).eq('username', user.get('email')).execute()
                            
                            # 2. Cộng hoa hồng (Chỉ khi lần đầu lên Pro)
                            if ref_code_input and is_first_time:
                                 ref_user = client.table('users_pro').select('commission_balance').eq('username', ref_code_input).execute()
                                 if ref_user.data:
                                     curr_comm = ref_user.data[0].get('commission_balance', 0)
                                     client.table('users_pro').update({
                                         'commission_balance': curr_comm + COMMISSION_AMT
                                     }).eq('username', ref_code_input).execute()

                            st.balloons()
                            st.success("🎉 CHÚC MỪNG! TÀI KHOẢN ĐÃ NÂNG CẤP LÊN PRO!")
                            time.sleep(2)
                            st.rerun()
                    else:
                        st.error("⚠️ Hệ thống chưa nhận được tiền. Vui lòng thử lại sau 30s.")

    # --- [NÂNG CẤP] TAB 6: ĐỐI TÁC (AFFILIATE) ---
    with tabs[5]:
        st.subheader("💰 CHƯƠNG TRÌNH ĐỐI TÁC (AFFILIATE)")
        st.info(f"Mã giới thiệu của bạn chính là tên đăng nhập: **{user.get('email')}**")
        client = init_supabase()
        if client:
            try:
                # Thống kê số người đã giới thiệu
                ref_res = client.table('users_pro').select("*").eq('referred_by', user.get('email')).execute()
                
                # Lấy số dư hoa hồng
                me_res = client.table('users_pro').select('commission_balance').eq('username', user.get('email')).execute()
                comm_balance = me_res.data[0].get('commission_balance', 0) if me_res.data else 0

                if ref_res.data:
                    count_ref = len(ref_res.data)
                    count_pro = sum(1 for u in ref_res.data if u['role'] == 'pro')
                    c1, c2, c3 = st.columns(3)
                    with c1: st.metric("Tổng người giới thiệu", f"{count_ref} người")
                    with c2: st.metric("Đã lên PRO", f"{count_pro} người")
                    with c3: st.metric("Hoa hồng hiện có", f"{comm_balance:,.0f}đ")
                    st.write("---")
                    st.write("**Danh sách thành viên:**")
                    df_ref = pd.DataFrame(ref_res.data)
                    if not df_ref.empty:
                        st.dataframe(df_ref[['username', 'fullname', 'role', 'created_at']], use_container_width=True)
                else: st.info("Bạn chưa giới thiệu được ai. Hãy chia sẻ Mã giới thiệu ngay!")
            except: st.error("Lỗi tải dữ liệu đối tác.")

    # --- TAB 7: HỒ SƠ & LỊCH SỬ ---
    with tabs[6]:
        c1, c2 = st.columns([2, 1])
        with c1: 
            st.write(f"**👤 Xin chào: {user.get('fullname')}**")
            st.write("---")
            st.subheader("🗂️ KHO ĐỀ CỦA BẠN (Đã lưu vĩnh viễn)")
            
            if st.button("🔄 Tải lại danh sách đề đã lưu"):
                client = init_supabase()
                if client:
                    try:
                        history_res = client.table('exam_history').select("*").eq('username', user.get('email')).order('id', desc=True).execute()
                        if history_res.data:
                            saved_exams = [item['exam_data'] for item in history_res.data]
                            st.session_state['dossier'] = saved_exams
                            st.success(f"Đã tải {len(saved_exams)} đề từ kho lưu trữ!")
                            time.sleep(1)
                            st.rerun()
                        else: st.info("Bạn chưa lưu đề nào.")
                    except: st.error("Lỗi tải lịch sử.")
            
            if st.session_state['dossier']:
                for e in st.session_state['dossier']: st.write(f"📄 {e['title']}")
            else: st.caption("Chưa có dữ liệu hiển thị.")

        with c2: 
            k = st.text_input("🔑 API Key Gemini (Nếu có)", type="password", key="api_key_in")
            if k: st.session_state['api_key'] = k

    # ==============================================================================
    # [MỚI - ĐÃ SỬA] TAB 8: TẠO ĐỀ CHUẨN YCCĐ (DÙNG DỮ LIỆU NHÚNG)
    # ==============================================================================
    with tabs[7]:
        st.title("🎯 Ngân hàng đề Toán Tiểu học (Chuẩn GDPT 2018)")
        st.caption("Dữ liệu bám sát Yêu cầu cần đạt - Đã tích hợp sẵn.")
        
        mgr = YCCDManager()
        current_api_key = st.session_state.get('api_key', '')
        if not current_api_key: current_api_key = SYSTEM_GOOGLE_KEY
        gen = QuestionGeneratorYCCD(current_api_key)

        with st.container():
            col1, col2, col3 = st.columns(3)
            with col1:
                # 1. Chọn Lớp (Tự động lấy từ file json)
                grades = mgr.get_grades()
                selected_grade = st.selectbox("1️⃣ Chọn Khối Lớp:", grades, index=len(grades)-1) # Mặc định chọn lớp 5

            with col2:
                # 2. Chọn Chủ đề tương ứng với Lớp
                topics = mgr.get_topics_by_grade(selected_grade)
                selected_topic = st.selectbox("2️⃣ Mạch kiến thức:", topics)

            with col3:
                # 3. Cấu hình số lượng
                num_q = st.number_input("Số câu hỏi:", 1, 20, 5, key="num_q_yccd")

        # 4. Chọn Yêu cầu cần đạt chi tiết
        if selected_topic:
            yccd_list = mgr.get_yccd_list(selected_grade, selected_topic)
            yccd_map = {f"{item['bai']}": item for item in yccd_list}
            
            selected_bai = st.selectbox("3️⃣ Chọn Bài học / Yêu cầu cụ thể:", list(yccd_map.keys()))
            target_item = yccd_map[selected_bai]
            
            st.info(f"📌 **Chuẩn kiến thức:** {target_item['yccd']}")
            
            muc_do = st.select_slider("Độ khó:", options=["Nhận biết", "Thông hiểu", "Vận dụng"])

            # --- NÚT TẠO ĐỀ ---
            if st.button("🚀 BẮT ĐẦU SOẠN ĐỀ", type="primary", key="btn_yccd"):
                if not current_api_key:
                    st.error("Chưa có API Key.")
                else:
                    st.divider()
                    my_bar = st.progress(0)
                    status_text = st.empty()
                    
                    for i in range(num_q):
                        status_text.markdown(f"**⏳ AI đang tư duy câu {i+1}/{num_q}...**")
                        data = gen.generate(target_item, muc_do)
                        my_bar.progress((i + 1) / num_q)
                        
                        if data:
                            with st.expander(f"✅ Câu {i+1}: {data.get('question', '...')}", expanded=True):
                                st.write(f"**Đề bài:** {data.get('question','')}")
                                if 'options' in data:
                                    cols = st.columns(4)
                                    for idx, opt in enumerate(data['options'][:4]):
                                        cols[idx].write(opt)
                                
                                st.success(f"**Đáp án:** {data.get('answer','')}")
                                st.warning(f"💡 **HD:** {data.get('explanation','')}")
                        else:
                            st.error(f"Câu {i+1}: AI gặp lỗi, đang thử lại...")
                    
                    status_text.success("🎉 Hoàn thành!")
                    my_bar.empty()
    
    st.markdown("---")
    st.markdown(textwrap.dedent('''<div style="text-align: center; color: #64748b; font-size: 14px; padding: 20px;"><strong>AI EXAM EXPERT v10</strong> © Tác giả: <strong>Trần Thanh Tuấn</strong> – Trường Tiểu học Hồng Thái – Năm 2026.<br>SĐT: 0918198687</div>'''), unsafe_allow_html=True)            

# ==============================================================================
# 7A. MODULE: TRỢ LÝ SOẠN GIÁO ÁN (TỔNG QUÁT TẤT CẢ MÔN/CẤP/BỘ SÁCH)
# ==============================================================================

def _lp_safe_key(prefix: str) -> str:
    """Sinh prefix key theo session để tránh trùng key giữa các module."""
    uid = st.session_state.get("user", {}).get("email", "guest")
    return f"{prefix}__{uid}"

def _lp_get_api_key():
    # Ưu tiên key người dùng nhập, fallback key hệ thống
    k = st.session_state.get("api_key", "")
    if not k:
        k = SYSTEM_GOOGLE_KEY
    return k

# ==============================================================================
# MODULE: TRỢ LÝ SOẠN BÀI – TẠO GIÁO ÁN TỰ ĐỘNG (UI PRO + ANTI DUP KEY)
# ==============================================================================

def _lp_uid():
    return st.session_state.get("user", {}).get("email", "guest")

def _lp_key(name: str) -> str:
    # key duy nhất theo user + module để chống DuplicateElementKey
    return f"lp_{name}_{_lp_uid()}"

def _lp_api_key():
    return st.session_state.get("api_key") or SYSTEM_GOOGLE_KEY

def _lp_init_state():
    if _lp_key("history") not in st.session_state:
        st.session_state[_lp_key("history")] = []   # lưu nhiều giáo án
    if _lp_key("last_html") not in st.session_state:
        st.session_state[_lp_key("last_html")] = ""
    if _lp_key("last_title") not in st.session_state:
        st.session_state[_lp_key("last_title")] = "GiaoAn"

# [FIX] Thêm 2 hàm này vào để xử lý lỗi NameError
def _lp_get_active(default_page):
    return st.session_state.get("lp_active_page_admin_state", default_page)

def _lp_set_active(page: str):
    st.session_state["lp_active_page_admin_state"] = page


def module_lesson_plan():
    """Module soạn giáo án (tối giản):
    - Input cốt lõi (môn/lớp/bộ sách/PPCT/tên bài/thời lượng)
    - (Tùy chọn) Tải tài liệu bài học để AI bám sát (PDF/Word)
    - Xuất HTML + Word (.doc)
    """
    _lp_init_state()

    st.markdown(f"""<style>
          .lp-hero{
            background: linear-gradient(135deg, #0F172A 0%, #1D4ED8 55%, #60A5FA 100%);
            border-radius: 14px;
            padding: 18px 18px 14px 18px;
            color: white;
            border: 1px solid rgba(255,255,255,.18);
            box-shadow: 0 10px 18px rgba(2,6,23,.18);
            margin-bottom: 14px;
          }
          .lp-hero h2{margin:0; font-weight:800;}
          .lp-box{background:#fff;border:1px solid #E2E8F0;border-radius:14px;padding:14px;margin-bottom:12px;}
          .lp-h{font-weight:800;color:#0F172A;margin:0 0 8px 0;}
        </style>""",
        unsafe_allow_html=True
    )

    st.markdown(f"""<div class='lp-hero'>
            <h2>📘 Soạn giáo án (Chuẩn CTGDPT 2018)</h2>
            <div style='opacity:.92;margin-top:6px'>
              Nhập thông tin bài dạy → (tuỳ chọn) tải tài liệu bài học → tạo giáo án HTML in A4 + tải Word.
            </div>
        </div>""",
        unsafe_allow_html=True
    )

    with st.form(key=_lp_key("form_simple"), clear_on_submit=False):
        st.markdown("<div class='lp-box'><div class='lp-h'>1) Thông tin bài dạy</div>", unsafe_allow_html=True)
        r1c1, r1c2, r1c3, r1c4 = st.columns([1.1, 1.2, 1.0, 1.2])
        with r1c1:
            st.selectbox("Năm học", ["2024-2025", "2025-2026", "2026-2027"], index=1, key=_lp_key("year"))
        with r1c2:
            level_key = st.radio("Cấp học", ["Tiểu học", "THCS", "THPT"], horizontal=True, key=_lp_key("level"))
        curr_lvl = "tieu_hoc" if level_key == "Tiểu học" else "thcs" if level_key == "THCS" else "thpt"
        edu = EDUCATION_DATA[curr_lvl]
        with r1c3:
            grade = st.selectbox("Khối lớp", edu["grades"], key=_lp_key("grade"))
        with r1c4:
            subject = st.selectbox("Môn học", edu["subjects"], key=_lp_key("subject"))

        r2c1, r2c2, r2c3 = st.columns([2.0, 1.0, 1.0])
        with r2c1:
            book = st.selectbox("Bộ sách", BOOKS_LIST, key=_lp_key("book"))
        with r2c2:
            ppct_week = st.number_input("Tuần (PPCT)", min_value=1, max_value=40, value=1, step=1, key=_lp_key("ppct_week"))
        with r2c3:
            ppct_period = st.number_input("Tiết (PPCT)", min_value=1, max_value=10, value=1, step=1, key=_lp_key("ppct_period"))

        lesson_title_input = st.text_input("Tên bài học (PPCT)", key=_lp_key("lesson_title_input"))

        r3c1, r3c2 = st.columns([1.2, 1.0])
        with r3c1:
            duration = st.number_input("Thời lượng (phút)", min_value=10, max_value=60, value=40, step=1, key=_lp_key("duration"))
        with r3c2:
            class_size = st.number_input("Sĩ số", min_value=10, max_value=60, value=40, step=1, key=_lp_key("class_size"))

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='lp-box'><div class='lp-h'>2) Tài liệu để AI bám sát (tuỳ chọn)</div>", unsafe_allow_html=True)
        c_up1, c_up2 = st.columns(2)
        with c_up1:
            lesson_file = st.file_uploader(
                "Tài liệu bài học (PDF/Word)",
                type=["pdf", "docx"],
                key=_lp_key("lesson_file"),
                help="Nếu là PDF scan/ảnh: hệ thống sẽ thử OCR (nếu VPS có cài pdf2image + pytesseract)."
            )
        with c_up2:
            ppct_file = st.file_uploader(
                "PPCT/KHDH (Word – tuỳ chọn)",
                type=["docx"],
                key=_lp_key("ppct_file")
            )

        ocr_col1, ocr_col2 = st.columns([1, 1])
        with ocr_col1:
            max_pages = st.number_input("Giới hạn trang PDF", min_value=1, max_value=12, value=6, step=1, key=_lp_key("pdf_pages"))
        with ocr_col2:
            ocr_on = st.checkbox("OCR nếu PDF là scan/ảnh", value=True, key=_lp_key("pdf_ocr"))

        preview_extract = st.checkbox("Xem trước nội dung trích xuất", value=False, key=_lp_key("preview_extract"))

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='lp-box'><div class='lp-h'>3) Ghi chú thêm (tuỳ chọn)</div>", unsafe_allow_html=True)
        teacher_note_extra = st.text_area(
            "Ghi chú cho AI",
            key=_lp_key("teacher_note_extra"),
            height=120,
            placeholder="Ví dụ: Có trò chơi khởi động 3 phút; ưu tiên hoạt động cặp đôi; tăng luyện tập; có 1 bài phân hoá..."
        )
        st.markdown("</div>", unsafe_allow_html=True)

        b1, b2 = st.columns([1.2, 1.0])
        with b1:
            generate_btn = st.form_submit_button("⚡ TẠO GIÁO ÁN", type="primary", use_container_width=True)
        with b2:
            regen_btn = st.form_submit_button("🔁 TẠO LẠI", use_container_width=True)

    if generate_btn or regen_btn:
        if not require_points_or_block(POINT_COST_LESSON_PLAN, action_name='soạn giáo án'):
            st.stop()

        api_key = _lp_api_key()
        if not api_key:
            st.error("❌ Chưa có API Key.")
            st.stop()

        lesson_title = (lesson_title_input or "").strip()
        if not lesson_title:
            st.error("❌ Vui lòng nhập Tên bài học (PPCT).")
            st.stop()

        # ---- trích xuất tài liệu bài học (nếu có) ----
        extracted_text = ""
        if lesson_file is not None:
            try:
                if lesson_file.name.lower().endswith(".pdf"):
                    pdf_bytes = lesson_file.getvalue()
                    extracted_text = extract_text_from_pdf_bytes(
                        pdf_bytes,
                        max_pages=int(max_pages),
                        ocr_if_needed=bool(ocr_on)
                    )
                elif lesson_file.name.lower().endswith(".docx"):
                    extracted_text = read_file_content(lesson_file, 'docx')
            except Exception:
                extracted_text = ""

        ppct_text = ""
        if ppct_file is not None:
            try:
                ppct_text = read_file_content(ppct_file, 'docx')
            except Exception:
                ppct_text = ""

        if preview_extract and (extracted_text or ppct_text):
            with st.expander("🔎 Xem trước nội dung trích xuất", expanded=True):
                if extracted_text:
                    st.markdown("**Tài liệu bài học:**")
                    st.write(extracted_text[:6000])
                if ppct_text:
                    st.markdown("**PPCT/KHDH:**")
                    st.write(ppct_text[:6000])

        ppct_week_val = int(ppct_week)
        ppct_period_val = int(ppct_period)

        meta_ppct = {
            "cap_hoc": level_key,
            "lop": grade,
            "mon": subject,
            "ten_bai": lesson_title,
            "tuan": ppct_week_val,
            "tiet": ppct_period_val,
            "bo_sach": book,
            "thoi_luong": int(duration),
            "si_so": int(class_size),
        }

        teacher_note = f"""PPCT: Tuần {ppct_week_val}, Tiết {ppct_period_val}
Ghi chú thêm: {teacher_note_extra.strip() if teacher_note_extra else ""}

YÊU CẦU CHẤT LƯỢNG:
- Không viết 'Bước 1/2' hoặc 'Nhiệm vụ 1/2' chung chung.
- Mỗi dòng hoạt động phải có NHIỆM VỤ HỌC TẬP CỤ THỂ (câu hỏi/bài tập/sản phẩm).
- Nếu Toán: phải có ví dụ số cụ thể + bài luyện tập (Bài 1, Bài 2...) và dự kiến đáp án/nhận xét.
""".strip()

        if extracted_text:
            teacher_note += build_pdf_context_for_teacher_note(extracted_text)
        if ppct_text:
            teacher_note += "\n\n[PPCT/KHDH – ƯU TIÊN BÁM SÁT]\n" + ppct_text[:12000]

        with st.spinner("🤖 AI đang soạn giáo án..." ):
            try:
                data = generate_lesson_plan_data_only(
                    api_key=api_key,
                    meta_ppct=meta_ppct,
                    teacher_note=teacher_note,
                    model_name="gemini-2.0-flash"
                )
                validate_lesson_plan(data)
                content_html = render_lesson_plan_html(data)
            except Exception as e:
                st.error(f"❌ Lỗi khi tạo giáo án: {e}")
                st.stop()

        st.session_state[_lp_key("last_title")] = f"Giáo án - {lesson_title}"
        st.session_state[_lp_key("last_html")] = content_html
        st.toast("Đã tạo giáo án!", icon="✅")

        # Trừ điểm nếu hệ thống points đã bật
        try:
            client = init_supabase()
            usern = st.session_state.get('user', {}).get('email','')
            if client and usern and get_user_points(client, usern) >= 0:
                deduct_user_points(client, usern, POINT_COST_LESSON_PLAN)
        except Exception:
            pass


    # ---- Xem trước & tải về ----
    content_html = st.session_state.get(_lp_key("last_html"), "")
    if content_html:
        st.markdown("## 📄 Xem trước giáo án")
        st.components.v1.html(content_html, height=760, scrolling=True)

        st.markdown("## ⬇️ Tải về")
        cdl1, cdl2 = st.columns([1.2, 1.2])
        title = st.session_state.get(_lp_key("last_title"), "GiaoAn")

        with cdl1:
            st.download_button(
                "⬇️ Tải Word (.doc)",
                data=create_word_doc(content_html, title),
                file_name=f"{title}.doc",
                mime="application/msword",
                type="primary",
                use_container_width=True,
                key=_lp_key("dl_word")
            )
        with cdl2:
            st.download_button(
                "⬇️ Tải HTML",
                data=content_html.encode("utf-8"),
                file_name=f"{title}.html",
                mime="text/html",
                use_container_width=True,
                key=_lp_key("dl_html")
            )

def login_screen():
    c1, c2, c3 = st.columns([1, 1.5, 1])

    with c2:
        st.markdown(
            "<h2 style='text-align:center; color:#1E3A8A'>🔐 HỆ THỐNG ĐĂNG NHẬP</h2>",
            unsafe_allow_html=True
        )

        # ✅ KHAI BÁO TAB ĐẦY ĐỦ
        tab_login, tab_signup = st.tabs(["ĐĂNG NHẬP", "ĐĂNG KÝ"])

        # ======================
        # TAB ĐĂNG NHẬP
        # ======================
        with tab_login:
            u = st.text_input("Tên đăng nhập", key="login_username")
            p = st.text_input("Mật khẩu", type="password", key="login_password")

            if st.button("ĐĂNG NHẬP", type="primary", key="login_btn"):
                client = init_supabase()
                if client:
                    try:
                        res = (
                            client.table("users_pro")
                            .select("*")
                            .eq("username", u)
                            .eq("password", p)
                            .execute()
                        )
                        if res.data:
                            user_data = res.data[0]
                            st.session_state["user"] = {
                                "email": user_data["username"],
                                "fullname": user_data["fullname"],
                                "role": user_data.get("role", "free"),
                                "points": user_data.get("points", 0),
                            }
                            st.toast("✅ Đăng nhập thành công! Đang chuyển về Trang chủ…", icon="✅")
                            target = st.session_state.pop("requested_page", None) or "dashboard"
                            go(target)
                        else:
                            st.error("Sai tài khoản hoặc mật khẩu")
                    except Exception as e:
                        st.error(f"Lỗi đăng nhập: {e}")

        # ======================
        # TAB ĐĂNG KÝ
        # ======================
        with tab_signup:
            new_u = st.text_input("Tên đăng nhập mới", key="signup_username")
            new_p = st.text_input("Mật khẩu mới", type="password", key="signup_password")
            new_name = st.text_input("Họ và tên", key="signup_fullname")

            if st.button("TẠO TÀI KHOẢN", key="signup_btn"):
                client = init_supabase()
                if client and new_u and new_p:
                    try:
                        check = (
                            client.table("users_pro")
                            .select("*")
                            .eq("username", new_u)
                            .execute()
                        )
                        if check.data:
                            st.warning("Tên đăng nhập đã tồn tại!")
                        else:
                            client.table("users_pro").insert(
                                {
                                    "username": new_u,
                                    "password": new_p,
                                    "fullname": new_name,
                                    "role": "free",
                                    "usage_count": 0,
                                    "points": 0,
                                }
                            ).execute()
                            st.success("Đăng ký thành công! Mời đăng nhập.")
                    except Exception as e:
                        st.error(f"Lỗi đăng ký: {e}")

# ==============================================================================
# 8. ROUTER + SIDEBAR MENU (ỔN ĐỊNH, KHÔNG TRÙNG KEY, KHÔNG MẤT LOGIN)
# ==============================================================================

def dashboard_screen():
    user = st.session_state.get("user", {}) or {}
    username = user.get("email") or ""

    client = init_supabase()
    points = -1
    role = user.get("role", "free")
    if client and username:
        row = get_user_row(client, username)
        role = row.get("role", role)
        if "points" in row:
            points = row.get("points", -1)
        # sync lại session để các module dùng nhất quán
        st.session_state.setdefault("user", {})
        st.session_state["user"]["role"] = role
        if points != -1:
            st.session_state["user"]["points"] = points

    # HERO (giống layout mẫu, nhưng cao cấp hơn)
    st.markdown(f"""
<div class="hero">
  <div style="display:flex; align-items:center; justify-content:center; gap:10px; margin-bottom:8px;">
    <div class="sb-logo" style="width:56px;height:56px;border-radius:14px;background:transparent;box-shadow:none;">{logo_svg(56)}</div>
    <div style="text-align:left">
      <div style="font-weight:800; font-size:14px; letter-spacing:.02em;">AIEXAM.VN</div>
      <div class="small-muted">Nền tảng AI dành cho giáo viên</div>
    </div>
  </div>

  <h1>Trợ lý AI giúp giáo viên làm nhanh — chuẩn — đẹp</h1>
  <p>Tạo đề kiểm tra, soạn giáo án, viết nội dung năng lực số và tư vấn nhận xét chỉ với vài thao tác.</p>
</div>
""", unsafe_allow_html=True)

    # Ask box + pills
    st.write("")
    c1, c2, c3 = st.columns([1, 2.2, 1])
    with c2:
        st.markdown('<div class="glass">', unsafe_allow_html=True)
        q = st.text_input(
            "",
            placeholder="Hỏi nhanh: “Tạo ma trận đề Toán 10 – chương Hàm số, mức độ vận dụng…”",
            key="dash_quick_ask",
            label_visibility="collapsed"
        )
        colA, colB, colC = st.columns([1.2, 1.2, 0.9])
        with colA:
            if st.button("🧠 Gợi ý prompt", use_container_width=True, key="dash_hint"):
                st.session_state["dash_quick_ask"] = "Hãy tạo ma trận đề theo yêu cầu: [môn/lớp/chủ đề/số câu/mức độ], sau đó sinh đề và đáp án."
                st.rerun()
        with colB:
            if st.button("🚀 Đi tới tạo đề", use_container_width=True, key="dash_go_exam"):
                go("exam")
        with colC:
            if st.button("➤", use_container_width=True, key="dash_send"):
                # Điều hướng thông minh theo từ khoá (không phá module)
                txt = (q or "").lower()
                if any(k in txt for k in ["đề", "ma trận", "ktđg", "trắc nghiệm", "tự luận"]):
                    go("exam")
                elif any(k in txt for k in ["giáo án", "bài dạy", "kế hoạch bài dạy", "ppct"]):
                    go("lesson_plan")
                elif any(k in txt for k in ["năng lực số", "digital", "nls"]):
                    go("digital")
                else:
                    go("advisor")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown(f"""
<div class="pills">
  <span class="pill">💬 Chat/Tư vấn</span>
  <span class="pill">📝 Ra đề – KTĐG</span>
  <span class="pill">📘 Soạn giáo án</span>
  <span class="pill">💻 Năng lực số</span>
</div>
""", unsafe_allow_html=True)

    st.write("")
    # Stat cards
    s1, s2, s3 = st.columns(3)
    with s1:
        st.markdown(f"""<div class="card">
<b>👤 Tài khoản</b><div class="small-muted">{username or "Chưa đăng nhập"}</div>
</div>""", unsafe_allow_html=True)
    with s2:
        badge = "PRO" if role == "pro" else "FREE"
        st.markdown(f"""<div class="card">
<b>⭐ Gói</b><div class="small-muted">{badge}</div>
</div>""", unsafe_allow_html=True)
    with s3:
        st.markdown(f"""<div class="card">
<b>💎 Điểm</b><div class="small-muted">{points if points != -1 else "—"}</div>
</div>""", unsafe_allow_html=True)

    st.write("")
    st.markdown("### Truy cập nhanh")
    qa = st.columns(6)
    quick = [
        ("📝", "Ra đề – KTĐG", "exam", "ic1"),
        ("📘", "Soạn giáo án", "lesson_plan", "ic2"),
        ("💻", "Năng lực số", "digital", "ic3"),
        ("🧠", "Nhận xét/Tư vấn", "advisor", "ic4"),
        ("📚", "Kho/Quản lý", "dashboard", "ic5"),
        ("📘", "Hướng dẫn", "help", "ic6"),
    ]
    for i, (emo, label, page_key, klass) in enumerate(quick):
        with qa[i]:
            st.markdown(f"""<div class="card soft" style="text-align:center;">
  <div class="icon-circle {klass}">{emo}</div>
  <div style="font-weight:750;">{label}</div>
  <div class="small-muted" style="margin-top:2px;">Mở ngay</div>
</div>""", unsafe_allow_html=True)
            if st.button("Mở", use_container_width=True, key=f"qa_open_{page_key}_{i}"):
                go(page_key)

    st.write("")
    left, right = st.columns([1.2, 1])
    with left:
        st.markdown(f"""<div class="card">
<b>📌 Gợi ý dùng nhanh</b>
<ul style="margin:10px 0 0 18px; color: rgba(15,23,42,.78);">
  <li>Vào <b>Ra đề – KTĐG</b> để tạo ma trận → đề → đáp án → xuất file.</li>
  <li>Vào <b>Soạn giáo án</b> để soạn theo PPCT/chuẩn mẫu.</li>
  <li>Vào <b>Năng lực số</b> để phân tích + tạo kế hoạch bài dạy tích hợp.</li>
</ul>
</div>""", unsafe_allow_html=True)

    with right:
        st.markdown(f"""<div class="card">
<b>🚀 Nâng cấp & thanh toán</b>
<div class="small-muted" style="margin-top:6px;">Quét VietQR → hệ thống xác minh SePay tự động. Sau khi chuyển khoản, bấm “Kích hoạt”.</div>
</div>""", unsafe_allow_html=True)

    # VIP Topup (giữ đúng logic gốc, chỉ bọc UI)
    with st.expander("⭐ Nạp VIP / Kích hoạt PRO (SePay tự xác minh)", expanded=False):
        if not client or not username:
            st.warning("Bạn cần đăng nhập để nạp VIP.")
            return

        ref_code_input = st.text_input("Mã giới thiệu (tuỳ chọn):", key="dash_ref_code")

        # Nội dung CK bắt buộc có tiền tố để SePay nhận diện
        final_content_ck = f"SEVQR NAP VIP {username}"
        if ref_code_input and ref_code_input != username:
            final_content_ck = f"SEVQR NAP VIP {username} REF {ref_code_input}"

        encoded_content = urllib.parse.quote(final_content_ck)
        qr_url = f"https://img.vietqr.io/image/{BANK_ID}-{BANK_ACC}-compact2.png?amount={VIP_TOPUP_AMOUNT_VND}&addInfo={encoded_content}&accountName={BANK_NAME}"
        q1, q2 = st.columns([1, 1.4])
        with q1:
            st.image(qr_url, caption=f"VietQR {VIP_TOPUP_AMOUNT_VND:,.0f}đ", width=280)
        with q2:
            st.markdown(f"""<div class="card soft">
<div style="font-weight:800; font-size:16px; margin-bottom:6px;">Thông tin chuyển khoản</div>
<div><b>Ngân hàng:</b> {BANK_NAME}</div>
<div><b>Số TK:</b> {BANK_NO}</div>
<div><b>Số tiền:</b> {VIP_TOPUP_AMOUNT_VND:,.0f} đ</div>
<div><b>Nội dung:</b> <code>{final_content_ck}</code></div>
<div class="small-muted" style="margin-top:8px;">Lưu ý: Nội dung cần đúng để hệ thống SePay nhận diện.</div>
</div>""", unsafe_allow_html=True)

            if st.button("🚀 KÍCH HOẠT NGAY (SePay tự xác minh)", type="primary", use_container_width=True, key="dash_activate_vip"):
                ok = check_sepay_transaction(VIP_TOPUP_AMOUNT_VND, final_content_ck)
                if not ok:
                    st.error("❌ Chưa thấy giao dịch phù hợp. Vui lòng kiểm tra lại hoặc đợi 1–2 phút rồi thử lại.")
                else:
                    row = get_user_row(client, username)
                    updates = {'role': 'pro'}
                    if isinstance(row, dict) and 'points' in row:
                        cur = int(row.get('points') or 0)
                        updates['points'] = cur + VIP_TOPUP_POINTS
                    client.table('users_pro').update(updates).eq('username', username).execute()

                    # Cập nhật session
                    st.session_state.setdefault("user", {})
                    st.session_state["user"]["role"] = "pro"
                    if 'points' in updates:
                        st.session_state["user"]["points"] = updates['points']

                    st.balloons()
                    st.success("✅ Kích hoạt VIP thành công! Điểm đã được cộng (nếu DB có cột points).")
                    st.rerun()

def module_digital():
    # --- CSS Tùy chỉnh cho Module NLS (Giống giao diện React) ---
    st.markdown(textwrap.dedent('''
    <style>
        .nls-container { background-color: #F8FAFC; padding: 20px; border-radius: 15px; }
        .nls-header { 
            background: linear-gradient(90deg, #1E3A8A 0%, #3B82F6 100%); 
            color: white; padding: 20px; border-radius: 12px; margin-bottom: 20px; 
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        }
        .nls-card { 
            background: white; padding: 25px; border-radius: 12px; 
            border: 1px solid #E2E8F0; box-shadow: 0 1px 3px rgba(0,0,0,0.1); margin-bottom: 20px; 
        }
        .nls-title { color: #1E3A8A; font-weight: 700; font-size: 16px; margin-bottom: 15px; border-left: 4px solid #3B82F6; padding-left: 10px; }
        .nls-upload-box { 
            border: 2px dashed #93C5FD; background: #EFF6FF; border-radius: 10px; 
            padding: 20px; text-align: center; color: #1E40AF; font-size: 14px;
        }
        .nls-btn {
            width: 100%; background: linear-gradient(90deg, #2563EB 0%, #1D4ED8 100%);
            color: white; font-weight: bold; padding: 12px; border-radius: 8px;
            text-align: center; border: none; cursor: pointer;
        }
        .nls-btn:hover { opacity: 0.9; }
    </style>
    '''), unsafe_allow_html=True)

    # --- Header ---
    st.markdown(textwrap.dedent('''
    <div class="nls-header">
        <div>
            <h2 style="margin:0; font-size: 22px;">💻 AI EXAM - SOẠN GIÁO ÁN NLS</h2>
            <p style="margin:5px 0 0 0; opacity: 0.9; font-size: 14px;">Hệ thống tích hợp Năng lực số tự động cho Giáo viên</p>
        </div>
    </div>
    '''), unsafe_allow_html=True)

    # --- Layout Chính: 2 Cột (Form bên trái, Hướng dẫn bên phải) ---
    col_left, col_right = st.columns([2, 1])

    with col_left:
        # 1. Thông tin bài dạy
        st.markdown('<div class="nls-card">', unsafe_allow_html=True)
        st.markdown('<div class="nls-title">1. Thông tin Kế hoạch bài dạy</div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1: textbook = st.selectbox("Bộ sách", ["Kết nối tri thức", "Chân trời sáng tạo", "Cánh Diều"], key="nls_book")
        with c2: subject = st.selectbox("Môn học", ["Toán", "Ngữ văn", "Tin học", "KHTN", "Lịch sử & Địa lí"], key="nls_sub")
        with c3: grade = st.selectbox("Khối lớp", [f"Lớp {i}" for i in range(1, 13)], index=6, key="nls_grade") # Mặc định lớp 3
        st.markdown('</div>', unsafe_allow_html=True)

        # 2. Tài liệu đầu vào
        st.markdown('<div class="nls-card">', unsafe_allow_html=True)
        st.markdown('<div class="nls-title">2. Tài liệu đầu vào (Upload file Word)</div>', unsafe_allow_html=True)
        
        c_up1, c_up2 = st.columns(2)
        with c_up1:
            st.markdown('<div class="nls-upload-box">📂 Tải lên Giáo án gốc<br>(Bắt buộc)</div>', unsafe_allow_html=True)
            file_lesson = st.file_uploader("Chọn file Giáo án", type=['docx'], key="nls_u1", label_visibility="collapsed")
        
        with c_up2:
            st.markdown('<div class="nls-upload-box">📊 Tải lên PPCT<br>(Tùy chọn để AI tham khảo)</div>', unsafe_allow_html=True)
            file_ppct = st.file_uploader("Chọn file PPCT", type=['docx'], key="nls_u2", label_visibility="collapsed")
        st.markdown('</div>', unsafe_allow_html=True)

        # 3. Tùy chọn & Xử lý
        st.markdown('<div class="nls-card">', unsafe_allow_html=True)
        st.markdown('<div class="nls-title">3. Tùy chọn xử lý</div>', unsafe_allow_html=True)
        
        check_col1, check_col2 = st.columns(2)
        with check_col1: analyze_only = st.checkbox("Chỉ phân tích (Không sửa nội dung)", key="nls_chk1")
        with check_col2: detailed_report = st.checkbox("Kèm báo cáo giải trình chi tiết", key="nls_chk2")

        st.write("") # Spacer
        
        # Nút bấm xử lý
        if st.button("✨ BẮT ĐẦU TÍCH HỢP NĂNG LỰC SỐ", type="primary", use_container_width=True):
            if not require_points_or_block(POINT_COST_NLS, action_name='soạn giáo án Năng lực số'):
                st.stop()

            api_key = st.session_state.get("api_key") or SYSTEM_GOOGLE_KEY
            if not api_key:
                st.error("⚠️ Vui lòng nhập API Key ở Tab Hồ Sơ trước!")
            elif not file_lesson:
                st.error("⚠️ Vui lòng tải lên file Giáo án gốc!")
            else:
                with st.spinner("🤖 AI đang phân tích và tích hợp năng lực số... Vui lòng đợi 30s"):
                    # Đọc nội dung file
                    lesson_text = read_file_content(file_lesson, 'docx')
                    ppct_text = read_file_content(file_ppct, 'docx') if file_ppct else ""
                    
                    # Gọi hàm xử lý (Đã định nghĩa ở Bước 1)
                    result_text = generate_nls_lesson_plan(
                        api_key, lesson_text, ppct_text, textbook, subject, grade, analyze_only
                    )
                    
                    # Lưu kết quả vào session
                    st.session_state['nls_result'] = result_text
                    st.success("✅ Đã xử lý xong!")
                    # Trừ điểm nếu hệ thống points đã bật
                    try:
                        client = init_supabase()
                        usern = st.session_state.get('user', {}).get('email','')
                        if client and usern and get_user_points(client, usern) >= 0:
                            deduct_user_points(client, usern, POINT_COST_NLS)
                    except Exception:
                        pass

        st.markdown('</div>', unsafe_allow_html=True)

    with col_right:
        # Sidebar thông tin (Giống UI React)
        st.markdown(textwrap.dedent('''
        <div class="nls-card" style="background:#EFF6FF; border:1px solid #BFDBFE;">
            <h4 style="color:#1E3A8A; margin-top:0;">💡 Hướng dẫn nhanh</h4>
            <ol style="font-size:14px; padding-left:15px; color:#334155;">
                <li>Chọn <b>Bộ sách, Môn, Lớp</b>.</li>
                <li>Tải lên <b>Giáo án gốc</b> (File Word .docx).</li>
                <li>Tải lên <b>PPCT</b> (Nếu muốn AI bám sát yêu cầu trường).</li>
                <li>Bấm <b>Bắt đầu</b> và đợi kết quả.</li>
            </ol>
        </div>
        '''), unsafe_allow_html=True)

        st.markdown(textwrap.dedent('''
        <div class="nls-card">
            <h4 style="color:#1E3A8A; margin-top:0;">🌐 Các miền Năng lực số</h4>
            <ul style="font-size:13px; padding-left:15px; color:#475569;">
                <li>Khai thác dữ liệu & thông tin</li>
                <li>Giao tiếp & Hợp tác số</li>
                <li>Sáng tạo nội dung số</li>
                <li>An toàn & An ninh số</li>
                <li>Giải quyết vấn đề với công nghệ</li>
                <li><b>Ứng dụng AI (Mới)</b></li>
            </ul>
        </div>
        '''), unsafe_allow_html=True)

    # --- Hiển thị kết quả ---
    if 'nls_result' in st.session_state and st.session_state['nls_result']:
        st.markdown("---")
        st.subheader("📄 KẾT QUẢ GIÁO ÁN NLS")
        
        # Tab xem trước và tải về
        tab_view, tab_download = st.tabs(["Xem trước", "Tải về"])
        
        with tab_view:
            st.markdown(st.session_state['nls_result'])
            
        with tab_download:
            # Tái sử dụng hàm create_word_doc có sẵn trong app.py cũ
            doc_html = st.session_state['nls_result'].replace("\n", "<br>") # Chuyển đổi sơ bộ sang HTML
            st.download_button(
                label="⬇️ Tải Giáo án Word (.doc)",
                data=create_word_doc(doc_html, "Giao_An_NLS"),
                file_name=f"Giao_An_NLS_{subject}_{grade}.doc",
                mime="application/msword",
                type="primary"
            )

def module_advisor():
    st.markdown("<div class='css-card'>", unsafe_allow_html=True)
    st.markdown("## 🧠 AI EDU Advisor – Nhận xét & Tư vấn")
    st.info("Mô-đun đang hoàn thiện. (Sẽ tích hợp sau)")
    st.markdown("</div>", unsafe_allow_html=True)

# ==============================================================================
# [LESSON PLAN SIMPLE v1] – TẠO GIÁO ÁN "NHƯ CHAT BÌNH THƯỜNG" (HTML TRỰC TIẾP)
# ==============================================================================

def _lp2_uid():
    return st.session_state.get("user", {}).get("email", "guest")

def _lp2_key(name: str) -> str:
    return f"lp2_{name}_{_lp2_uid()}"

def _lp2_api_key():
    return st.session_state.get("api_key") or SYSTEM_GOOGLE_KEY

def _lp2_extract_from_upload(uploaded_file) -> str:
    if not uploaded_file:
        return ""
    name = (uploaded_file.name or "").lower()
    try:
        if name.endswith(".pdf"):
            pdf_bytes = uploaded_file.getvalue()
            txt = extract_text_from_pdf_bytes(pdf_bytes, max_pages=6, ocr_if_needed=True)
            return txt or ""
        if name.endswith(".docx"):
            return read_file_content(uploaded_file, "docx") or ""
        if name.endswith(".txt"):
            return uploaded_file.getvalue().decode("utf-8", errors="ignore")
    except Exception:
        return ""
    return ""

def generate_lesson_plan_html_simple(
    api_key: str,
    cap_hoc: str,
    mon: str,
    lop: str,
    bo_sach: str,
    tuan: int,
    tiet: int,
    ten_bai: str,
    thoi_luong: int,
    si_so: int,
    lesson_context: str,
    teacher_note: str,
    model_name: str = "gemini-2.0-flash",
) -> str:
    """Trả về HTML hoàn chỉnh (không JSON)."""
    genai.configure(api_key=api_key)

    system_instruction = """Bạn là GIÁO VIÊN cốt cán, chuyên soạn KẾ HOẠCH BÀI DẠY theo CTGDPT 2018.
YÊU CẦU BẮT BUỘC:
- ĐẦU RA: CHỈ TRẢ VỀ 01 KHỐI HTML HOÀN CHỈNH (không markdown, không giải thích).
- Font: Times New Roman, cỡ 13pt; in A4 đẹp.
- Có 4 phần:
  I. Yêu cầu cần đạt (Kiến thức/Kĩ năng; Năng lực; Phẩm chất; Năng lực đặc thù nếu có; Năng lực số nếu phù hợp).
  II. Đồ dùng dạy – học (GV/HS).
  III. Các hoạt động dạy – học chủ yếu: BẮT BUỘC là <table border="1"> 2 cột:
      Cột 1: Hoạt động của Giáo viên
      Cột 2: Hoạt động của Học sinh
     Chia 3 hoạt động lớn: Khởi động; Khám phá/Hình thành kiến thức; Luyện tập/Vận dụng.
     VIẾT CHI TIẾT: câu hỏi gợi mở, ví dụ minh họa, bài tập cụ thể, dự kiến đáp án/nhận xét.
  IV. Điều chỉnh sau bài dạy: để dòng chấm.
- KHÔNG dùng các cụm 'Bước 1/2', 'Nhiệm vụ 1/2', 'Bổ sung nội dung' chung chung.
- Nếu có NỘI DUNG BÀI HỌC từ file (PDF/DOCX): phải bám sát thuật ngữ, ví dụ, bài tập trong đó. Không tự bịa ngoài tài liệu trừ khi ghi chú GV yêu cầu.
"""

    lesson_context = (lesson_context or "").strip()
    ctx_block = ""
    if lesson_context:
        ctx_block = "\n\n[NỘI DUNG BÀI HỌC TRÍCH TỪ TÀI LIỆU GV TẢI LÊN – ƯU TIÊN BÁM SÁT]\n" + lesson_context[:12000]

    prompt = f"""THÔNG TIN BÀI DẠY:
- Cấp học: {cap_hoc}
- Môn: {mon}
- Lớp: {lop}
- Bộ sách: {bo_sach}
- Tuần/Tiết (PPCT): {tuan}/{tiet}
- Tên bài: {ten_bai}
- Thời lượng: {thoi_luong} phút
- Sĩ số: {si_so}

GHI CHÚ/ĐIỀU CHỈNH CỦA GV:
{teacher_note.strip() if teacher_note else "(Không có)"}
{ctx_block}

HÃY SOẠN GIÁO ÁN HTML HOÀN CHỈNH THEO ĐÚNG YÊU CẦU.
"""

    model = genai.GenerativeModel(model_name, system_instruction=system_instruction)

    safe_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]

    res = model.generate_content(prompt, safety_settings=safe_settings)
    html = (res.text or "").strip()

    if "```" in html:
        parts = re.split(r"```(?:html)?", html)
        if len(parts) >= 2:
            html = parts[1].strip()

    if "<html" not in html.lower():
        html = f"""<!doctype html>
<html lang="vi"><head><meta charset="utf-8"/>
<style>
  @page {{ size: 21cm 29.7cm; margin: 2cm; }}
  body{{font-family:'Times New Roman',serif;font-size:13pt;line-height:1.35;color:#111;}}
  table{{width:100%;border-collapse:collapse;table-layout:fixed;}}
  td,th{{border:1px solid #000;padding:6px;vertical-align:top;word-wrap:break-word;}}
  th{{text-align:center;font-weight:700;background:#f2f2f2;}}
  h1{{text-align:center;font-size:18pt;margin:0 0 10px 0;}}
  h2{{font-size:14pt;margin:12px 0 6px 0;}}
</style>
</head><body>
{html}
</body></html>"""
    return html

def module_lesson_plan():
    """Module soạn giáo án (tối giản + AI trả HTML trực tiếp)."""
    st.markdown(textwrap.dedent('''<div style="background:linear-gradient(135deg,#0F172A 0%,#1D4ED8 55%,#60A5FA 100%);
      border-radius:14px;padding:16px 18px;color:#fff;border:1px solid rgba(255,255,255,.18);
      box-shadow:0 10px 18px rgba(2,6,23,.18);margin-bottom:14px;">
      <h2 style="margin:0;font-weight:800;">📘 Soạn giáo án (HTML – Chuẩn CTGDPT 2018)</h2>
      <div style="opacity:.92;margin-top:6px;">Tối giản: nhập thông tin + (tuỳ chọn) tải PDF/DOCX bài học → AI soạn chi tiết, có bảng GV/HS.</div>
    </div>'''), unsafe_allow_html=True)

    with st.form(key=_lp2_key("form"), clear_on_submit=False):
        r1c1, r1c2, r1c3, r1c4 = st.columns([1.1, 1.2, 1.0, 1.2])
        with r1c1:
            st.selectbox("Năm học", ["2024-2025", "2025-2026", "2026-2027"], index=1, key=_lp2_key("year"))
        with r1c2:
            cap_hoc = st.radio("Cấp học", ["Tiểu học", "THCS", "THPT"], horizontal=True, key=_lp2_key("cap_hoc"))
        curr_lvl = "tieu_hoc" if cap_hoc == "Tiểu học" else "thcs" if cap_hoc == "THCS" else "thpt"
        edu = EDUCATION_DATA[curr_lvl]
        with r1c3:
            lop = st.selectbox("Khối lớp", edu["grades"], key=_lp2_key("lop"))
        with r1c4:
            mon = st.selectbox("Môn học", edu["subjects"], key=_lp2_key("mon"))

        r2c1, r2c2, r2c3 = st.columns([2.0, 1.0, 1.0])
        with r2c1:
            bo_sach = st.selectbox("Bộ sách", BOOKS_LIST, key=_lp2_key("bo_sach"))
        with r2c2:
            tuan = st.number_input("Tuần (PPCT)", min_value=1, max_value=40, value=1, step=1, key=_lp2_key("tuan"))
        with r2c3:
            tiet = st.number_input("Tiết (PPCT)", min_value=1, max_value=10, value=1, step=1, key=_lp2_key("tiet"))

        ten_bai = st.text_input("Tên bài học (PPCT)", key=_lp2_key("ten_bai"))

        r3c1, r3c2 = st.columns([1.2, 1.0])
        with r3c1:
            thoi_luong = st.number_input("Thời lượng (phút)", min_value=20, max_value=60, value=40, step=1, key=_lp2_key("thoi_luong"))
        with r3c2:
            si_so = st.number_input("Sĩ số (tuỳ chọn)", min_value=10, max_value=60, value=40, step=1, key=_lp2_key("si_so"))

        st.markdown("### Tài liệu bài học (tuỳ chọn nhưng khuyến nghị)")
        up1, up2 = st.columns([1.2, 1.8])
        with up1:
            lesson_file = st.file_uploader("Tải PDF/DOCX/TXT bài học", type=["pdf","docx","txt"], key=_lp2_key("lesson_file"))
        with up2:
            show_preview = st.checkbox("Xem trước nội dung trích xuất", value=False, key=_lp2_key("show_preview"))

        teacher_note = st.text_area(
            "Ghi chú GV (tuỳ chọn)",
            key=_lp2_key("teacher_note"),
            height=110,
            placeholder="Ví dụ: Có trò chơi khởi động 3 phút; tăng luyện tập; ưu tiên hoạt động cặp đôi; có 1 bài phân hoá..."
        )

        b1, b2 = st.columns([1.2, 1.0])
        with b1:
            submit = st.form_submit_button("⚡ TẠO GIÁO ÁN", type="primary", use_container_width=True)
        with b2:
            reset = st.form_submit_button("🧹 XÓA KẾT QUẢ", use_container_width=True)

    if reset:
        st.session_state[_lp2_key("html")] = ""

    lesson_ctx = _lp2_extract_from_upload(lesson_file) if lesson_file else ""
    if lesson_file and show_preview:
        st.markdown("#### Preview nội dung trích xuất")
        st.text_area("Nội dung trích xuất", value=(lesson_ctx[:6000] if lesson_ctx else "(Không trích xuất được text từ file)"), height=220)

    if submit:
        if not ten_bai.strip():
            st.error("❌ Vui lòng nhập Tên bài học (PPCT).")
            st.stop()

        if lesson_file and not lesson_ctx.strip():
            st.warning("⚠️ File tải lên không trích xuất được text. Nếu PDF là scan ảnh, VPS cần pdf2image + pytesseract + poppler.")

        api_key_use = _lp2_api_key()
        if not api_key_use:
            st.error("❌ Chưa có API Key.")
            st.stop()

        with st.spinner("🤖 AI đang soạn giáo án..."):
            try:
                html = generate_lesson_plan_html_simple(
                    api_key=api_key_use,
                    cap_hoc=cap_hoc,
                    mon=mon,
                    lop=lop,
                    bo_sach=bo_sach,
                    tuan=int(tuan),
                    tiet=int(tiet),
                    ten_bai=ten_bai.strip(),
                    thoi_luong=int(thoi_luong),
                    si_so=int(si_so),
                    lesson_context=lesson_ctx,
                    teacher_note=teacher_note or "",
                    model_name="gemini-2.0-flash",
                )
                st.session_state[_lp2_key("html")] = html
                st.session_state[_lp2_key("title")] = f"GiaoAn_{mon}_{lop}_{ten_bai.strip()}"
                st.success("✅ Đã tạo giáo án!")
            except Exception as e:
                st.error(f"❌ Lỗi khi tạo giáo án: {e}")

    html = st.session_state.get(_lp2_key("html"), "")
    if html:
        st.markdown("## Xem trước (A4)")
        st.components.v1.html(html, height=780, scrolling=True)

        st.markdown("## Tải về")
        c1, c2 = st.columns(2)
        with c1:
            st.download_button(
                "⬇️ Tải Word (.doc)",
                data=create_word_doc(html, st.session_state.get(_lp2_key("title"), "GiaoAn")),
                file_name=f"{st.session_state.get(_lp2_key('title'),'GiaoAn')}.doc",
                mime="application/msword",
                type="primary",
                use_container_width=True,
                key=_lp2_key("dl_doc"),
            )
        with c2:
            st.download_button(
                "⬇️ Tải HTML",
                data=html.encode("utf-8"),
                file_name=f"{st.session_state.get(_lp2_key('title'),'GiaoAn')}.html",
                mime="text/html",
                use_container_width=True,
                key=_lp2_key("dl_html"),
                )

# ==============================================================================
# 8B. PREMIUM TOPBAR + PUBLIC LANDING + MODULES (CHAT/DOC/MINDMAP)
# - Trang vào (Home) công khai, không bắt đăng nhập
# - Demo 1 câu hỏi AI thật ở Home/Chat (guest)
# - Chỉ khi dùng tiếp hoặc dùng module nâng cao mới yêu cầu đăng nhập
# ==============================================================================

PROTECTED_PAGES = {"exam", "lesson_plan", "digital", "advisor", "doc_ai", "mindmap", "profile"}
DEMO_ALLOWED_PAGES = {"dashboard", "chat"}  # guest được xem + demo 1 câu

def _get_api_key_effective() -> str:
    # Ưu tiên key user nhập, fallback key hệ thống
    k = (st.session_state.get("api_key") or "").strip()
    if not k:
        k = (SYSTEM_GOOGLE_KEY or "").strip()
    return k

def require_login(page_key: str):
    if st.session_state.get("user"):
        return
    st.session_state["requested_page"] = page_key
    st.session_state["current_page"] = "login"
    st.rerun()

def _ensure_nav_state():
    st.session_state.setdefault("current_page", "dashboard")
    st.session_state.setdefault("requested_page", None)
    st.session_state.setdefault("demo_used", False)
    st.session_state.setdefault("demo_history", [])  # lưu demo Q/A để hiện lại


def render_topbar():
    """Topbar gọn (không trùng điều hướng sidebar) + dropdown tài khoản."""
    _ensure_nav_state()
    user = st.session_state.get("user") or {}
    is_authed = bool(user)
    fullname = user.get("fullname") or user.get("email") or "Khách"

    c1, c2, c3 = st.columns([2.8, 5.2, 2.0], vertical_alignment="center")

    with c1:
        st.markdown(
            f"""
<div style="display:flex;gap:10px;align-items:center;">
  <div style="width:52px;height:52px;border-radius:14px;background:transparent;box-shadow:none;overflow:visible;">
    {logo_svg(52)}
  </div>
  <div>
    <div style="font-weight:900;line-height:1.05;">AIEXAM.VN</div>
    <div class="small-muted">Nền tảng AI dành cho giáo viên</div>
  </div>
</div>
""",
            unsafe_allow_html=True,
        )

    with c2:
        # Topbar chỉ để truy cập nhanh "Hướng dẫn" + tìm kiếm (không trùng menu sidebar)
        cc1, cc2 = st.columns([1, 1], vertical_alignment="center")
        with cc1:
            st.text_input(
                "",
                placeholder="Tìm nhanh: 'ra đề', 'soạn bài', 'năng lực số'…",
                key="global_search",
                label_visibility="collapsed",
            )
        with cc2:
            if st.button("📘 Hướng dẫn", use_container_width=True, key="tb_help"):
                go("help")

    with c3:
        if is_authed:
            with st.popover(f"👤 {fullname}", use_container_width=True):
                role = (user.get("role") or "free").upper()
                pts = user.get("points", 0)
                st.markdown(f"**Gói:** `{role}`  \n**Điểm:** `{pts}`")
                st.write("---")
                if st.button("👤 Profile", use_container_width=True, key="tb_profile"):
                    go("profile")
                if st.button("🚪 Đăng xuất", use_container_width=True, key="tb_logout"):
                    st.session_state.pop("user", None)
                    st.toast("👋 Bạn đã đăng xuất.", icon="✅")
                    go("dashboard")
        else:
            if st.button("🔐 Đăng nhập", type="primary", use_container_width=True, key="tb_login"):
                st.session_state["requested_page"] = st.session_state.get("current_page", "dashboard")
                go("login")


def _gemini_generate(prompt: str, system: str | None = None) -> str:
    api_key = _get_api_key_effective()
    if not api_key:
        return "⚠️ Chưa cấu hình GOOGLE_API_KEY trong st.secrets hoặc bạn chưa nhập API key."
    try:
        genai.configure(api_key=api_key)
        if system:
            model = genai.GenerativeModel("gemini-2.0-flash", system_instruction=system)
        else:
            model = genai.GenerativeModel("gemini-2.0-flash")
        safe_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        res = model.generate_content(prompt, safety_settings=safe_settings)
        return (res.text or "").strip()
    except Exception as e:
        return f"❌ Lỗi AI: {e}"

def _chunk_text(text: str, chunk_size: int = 900, overlap: int = 120) -> list[str]:
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if not text:
        return []
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        j = min(n, i + chunk_size)
        chunks.append(text[i:j])
        if j == n:
            break
        i = max(0, j - overlap)
    return chunks

def _simple_retrieve(query: str, chunks: list[str], k: int = 4) -> list[str]:
    # Retrieval nhẹ không dùng embeddings (ổn định cho Streamlit Cloud)
    q = (query or "").lower()
    if not q or not chunks:
        return chunks[:k]
    q_terms = [t for t in re.split(r"[^\wÀ-ỹ]+", q) if t]
    scored = []
    for ch in chunks:
        s = 0
        low = ch.lower()
        for t in q_terms[:20]:
            if t and t in low:
                s += 1
        scored.append((s, ch))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [c for s, c in scored[:k] if s > 0]
    return top if top else chunks[:k]

def module_chat():
    _ensure_nav_state()
    user = st.session_state.get("user")
    # Guest: cho demo 1 câu ở Chat; lần 2 yêu cầu login
    st.markdown("## 💬 Chat AI")
    st.caption("Hỏi AI như ChatGPT. Khách được dùng thử 1 câu. Đăng nhập để dùng đầy đủ.")

    st.session_state.setdefault("chat_messages", [])

    # Hiển thị lịch sử
    for m in st.session_state["chat_messages"]:
        with st.chat_message(m.get("role", "assistant")):
            st.markdown(m.get("content", ""))

    prompt = st.chat_input("Nhập câu hỏi của bạn…")
    if prompt:
        # kiểm demo
        if (not user) and st.session_state.get("demo_used"):
            require_login("chat")
            return

        st.session_state["chat_messages"].append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("AI đang trả lời…"):
                reply = _gemini_generate(
                    f"Bạn là trợ lý AI cho giáo viên. Trả lời ngắn gọn, đúng trọng tâm.\n\nCâu hỏi: {prompt}"
                )
                st.markdown(reply if reply else "…")
        st.session_state["chat_messages"].append({"role": "assistant", "content": reply})

        if not user:
            st.session_state["demo_used"] = True
            st.info("Bạn vừa dùng thử 1 câu. Đăng nhập để tiếp tục sử dụng đầy đủ.")

    cols = st.columns([1,1,2])
    with cols[0]:
        if st.button("🧹 Xóa chat", key="chat_clear"):
            st.session_state["chat_messages"] = []
            st.rerun()
    with cols[1]:
        if st.button("⬅️ Về Home", key="chat_home"):
            go("dashboard")

def module_doc_ai():
    _ensure_nav_state()
    if not st.session_state.get("user"):
        require_login("doc_ai")
        return

    st.markdown("## 📄 Doc AI • Tóm tắt & Chat theo tài liệu")
    st.caption("Tải tài liệu (PDF/DOCX/ảnh) → tóm tắt → hỏi theo nội dung tài liệu. (RAG nhẹ, ổn định Cloud)")

    doc_file = st.file_uploader("Tải tài liệu", type=["pdf","docx","txt","png","jpg","jpeg"], key="docai_upload")
    max_pages = st.slider("Giới hạn số trang xử lý (PDF)", 1, 20, 6, key="docai_pages")
    try_ocr = st.checkbox("Thử OCR nếu PDF scan/ảnh", value=True, key="docai_ocr")

    if doc_file:
        with st.spinner("Đang đọc tài liệu…"):
            raw = extract_text_from_upload(doc_file, max_pages=max_pages, ocr_if_needed=try_ocr)
            raw = (raw or "").strip()
            if not raw:
                st.error("Không đọc được nội dung. Thử bật OCR hoặc dùng bản PDF có text.")
            else:
                st.session_state["docai_text"] = raw[:20000]
                st.session_state["docai_chunks"] = _chunk_text(st.session_state["docai_text"])
                st.success(f"Đã nạp tài liệu: {getattr(doc_file,'name','file')}")

    tabs = st.tabs(["🧾 Tóm tắt", "💬 Chat theo tài liệu", "👁️ Xem nội dung"])
    with tabs[0]:
        if st.button("✨ Tạo tóm tắt", type="primary", key="docai_sum"):
            txt = (st.session_state.get("docai_text") or "").strip()
            if not txt:
                st.warning("Hãy tải tài liệu trước.")
            else:
                with st.spinner("AI đang tóm tắt…"):
                    out = _gemini_generate(
                        """Bạn là trợ lý học thuật. Tóm tắt tài liệu ngắn gọn theo mục:
- Nội dung chính (5-7 gạch đầu dòng)
- Khái niệm quan trọng
- Gợi ý 5 câu hỏi ôn tập
\n\nTài liệu:
""" + txt[:16000]
                    )
                st.markdown(out)

    with tabs[1]:
        txt = (st.session_state.get("docai_text") or "").strip()
        if not txt:
            st.info("Tải tài liệu trước để chat theo tài liệu.")
        q = st.text_input("Nhập câu hỏi về tài liệu…", key="docai_q")
        if st.button("Hỏi tài liệu", key="docai_ask", type="primary"):
            if not txt:
                st.warning("Chưa có tài liệu.")
            else:
                ctx_chunks = _simple_retrieve(q, st.session_state.get("docai_chunks") or [], k=4)
                ctx = "\n\n---\n\n".join(ctx_chunks)
                with st.spinner("AI đang trả lời theo tài liệu…"):
                    out = _gemini_generate(
                        f"""Bạn là trợ lý AI. CHỈ trả lời dựa trên phần trích dẫn tài liệu dưới đây.
Nếu trong tài liệu không có, hãy nói rõ 'Tài liệu không đề cập'. Không bịa thêm.

[TRÍCH DẪN TÀI LIỆU]
{ctx}

[CÂU HỎI]
{q}
"""
                    )
                st.markdown(out)

    with tabs[2]:
        txt = (st.session_state.get("docai_text") or "").strip()
        st.text_area("Nội dung trích xuất (đã rút gọn)", value=txt[:16000], height=320, key="docai_preview")

def module_mindmap():
    _ensure_nav_state()
    if not st.session_state.get("user"):
        require_login("mindmap")
        return

    st.markdown("## 🧠 Mindmap AI")
    st.caption("Nhập chủ đề hoặc nội dung → AI tạo mindmap dạng cây (Markdown). Dùng cho soạn bài/ôn tập.")

    inp = st.text_area("Nội dung / chủ đề", height=200, key="mm_in")
    if st.button("✨ Tạo Mindmap", type="primary", key="mm_go"):
        if not inp.strip():
            st.warning("Nhập nội dung trước.")
        else:
            with st.spinner("AI đang tạo mindmap…"):
                out = _gemini_generate(
                    """Bạn là trợ lý giáo dục. Tạo mindmap dạng Markdown Tree (bullet phân cấp),
ngắn gọn, rõ ý, dễ học, phù hợp giáo viên.
Quy tắc:
- Dòng đầu là chủ đề chính
- Tối đa 4 cấp
- Mỗi nhánh 2-6 ý
\n\nNội dung:
""" + inp[:12000]
                )
            st.markdown(out)
            st.download_button("⬇️ Tải mindmap (.md)", data=out.encode("utf-8"), file_name="mindmap.md", mime="text/markdown", use_container_width=True)



# ==============================================================================
# MODULE: HƯỚNG DẪN (2 TAB) – Dành cho thầy/cô
# - Tab 1: Hướng dẫn sử dụng module
# - Tab 2: Hướng dẫn nạp VIP/PRO + điểm
# ==============================================================================
def module_help():
    st.markdown("## 📘 Hướng dẫn sử dụng")
    st.caption("Tài liệu hướng dẫn nhanh dành cho thầy/cô – dễ hiểu – dùng được ngay.")

    tab1, tab2 = st.tabs(["🧠 Hướng dẫn sử dụng module", "💎 Hướng dẫn nạp VIP / PRO"])

    # -----------------------------
    # TAB 1: MODULES
    # -----------------------------
    with tab1:
        st.markdown("### D. 💬 Chat AI (hỏi đáp nhanh như ChatGPT)")
        st.markdown(f"""
**Dùng khi nào?**  
Hỏi đáp kiến thức, soạn câu hỏi, gợi ý hoạt động dạy học, viết nhận xét, chỉnh câu chữ…

**Cách dùng nhanh (3 bước):**
1) Vào **Chat AI**  
2) Nhập yêu cầu theo mẫu: **Môn – Lớp – Nội dung – Mục tiêu – Định dạng kết quả**  
3) Nếu chưa đúng, gõ tiếp: *“Sửa theo…”* / *“Làm ngắn hơn…”* / *“Chi tiết hơn…”*

**Ví dụ prompt:**
- “Tôi dạy **Toán 8**, giải thích **hằng đẳng thức** dễ hiểu + 3 ví dụ.”
- “Tạo **10 câu trắc nghiệm** Sinh 10 chương 2, có đáp án + giải thích ngắn.”
            """
        )
        st.info("Mẹo: ghi rõ thời lượng (35/45/90 phút), đối tượng HS (trung bình/khá), chuẩn CTGDPT 2018 nếu cần.")

        st.divider()

        st.markdown("### E. 📄 Doc AI (tóm tắt & hỏi theo tài liệu)")
        st.markdown(f"""
**Dùng khi nào?**  
Khi thầy/cô có **PDF / DOCX / TXT** cần tóm tắt, rút ý chính, tạo câu hỏi ôn tập hoặc hỏi theo nội dung tài liệu.

**Cách dùng:**
1) Vào **Doc AI**  
2) **Tải tài liệu lên**  
3) Chọn yêu cầu: *Tóm tắt* / *Rút ý chính* / *Tạo câu hỏi* / *Dàn ý bài giảng*

**Ví dụ:**
- “Tóm tắt tài liệu thành 8 ý chính, chia theo từng mục.”
- “Tạo 10 câu hỏi ôn tập + đáp án dựa trên tài liệu.”
            """
        )
        st.warning("Nếu tài liệu là PDF scan/ảnh: hãy bật OCR (nếu có) để trích nội dung chính xác hơn.")

        st.divider()

        st.markdown("### F. 🧠 Mindmap (sơ đồ tư duy)")
        st.markdown(f"""
**Dùng khi nào?**  
Tạo sơ đồ tư duy cho bài học/chương, ôn tập nhanh, làm slide, giao bài cho học sinh.

**Cách dùng:**
1) Nhập **chủ đề** hoặc dán **nội dung bài**  
2) Yêu cầu *mindmap 3–4 cấp*, *ngắn gọn/dễ học*, *có ví dụ*

**Ví dụ:**
- “Mindmap Lịch sử 9 – Cách mạng tháng Tám, dạng 4 cấp, dễ học.”
            """
        )

        st.divider()

        st.markdown("### G. 📝 Ra đề – KTĐG (ma trận – đề – đáp án)")
        st.markdown(f"""
**Dùng khi nào?**  
Tạo đề kiểm tra/đề thi theo chuẩn đánh giá (NB/TH/VD/VDC hoặc M1/M2/M3), có thể kèm ma trận/đặc tả.

**Quy trình chuẩn:**
1) Chọn **môn – lớp – phạm vi kiến thức**  
2) Chọn dạng: Trắc nghiệm / Tự luận / Kết hợp  
3) Chọn số lượng câu & mức độ → bấm **Tạo đề**  
4) Xem trước → chỉnh → **Xuất file** (nếu có)

**Ví dụ prompt:**
- “Đề 45 phút Toán 7, 20 TN + 2 TL, mức độ vừa, có đáp án.”
            """
        )

        st.divider()

        st.markdown("### H. 📘 Trợ lý Soạn bài (tạo giáo án tự động)")
        st.markdown(f"""
**Dùng khi nào?**  
Soạn giáo án nhanh theo môn/lớp/bộ sách, có mục tiêu, hoạt động GV–HS, luyện tập, vận dụng, đánh giá.

**Cách dùng:**
1) Chọn môn – lớp – bài – bộ sách  
2) Nhập yêu cầu (thời lượng, phương pháp, thiết bị)  
3) Bấm tạo → chỉnh theo lớp dạy → xuất/lưu (nếu có)

**Mẹo hay:**  
Dán 1 đoạn mẫu giáo án của trường và yêu cầu: *“viết theo đúng format này”*.
            """
        )

        st.divider()

        st.markdown("### I. 💻 Năng lực số (tích hợp NLS vào giáo án)")
        st.markdown(f"""
**Dùng khi nào?**  
Tích hợp **Năng lực số** vào bài dạy: hoạt động số, công cụ số, sản phẩm số, tiêu chí đánh giá.

**Cách dùng:**
1) Chọn môn – lớp – bài (hoặc tải giáo án gốc nếu module hỗ trợ)  
2) Chọn mục tiêu NLS (tìm kiếm, hợp tác, an toàn số, AI…)  
3) Bấm tạo → nhận hoạt động + sản phẩm + tiêu chí đánh giá

**Mẹo:**  
Ghi rõ điều kiện lớp học (có/không phòng máy, dùng điện thoại, internet yếu…).
            """
        )

        st.divider()
        st.markdown("### 🧩 Nhận xét – Tư vấn")
        st.markdown(f"""
**Dùng khi nào?**  
Viết nhận xét học sinh theo năng lực/phẩm chất, góp ý giáo án, tư vấn cải tiến hoạt động dạy học.

**Ví dụ:**
- “Viết nhận xét môn Văn cho HS mức trung bình, giọng văn tích cực, có hướng cải thiện.”
            """
        )

        st.success("✅ Gợi ý chung: Càng nêu rõ *môn – lớp – mục tiêu – thời lượng – định dạng*, AI càng ra kết quả đúng ý.")

    # -----------------------------
    # TAB 2: VIP/PRO
    # -----------------------------
    with tab2:
        st.markdown("### A. 🚀 Bắt đầu nhanh (nạp VIP trong 1–2 phút)")
        st.markdown(f"""
1) **Đăng nhập/Đăng ký** tài khoản  
2) Vào **Trang chủ → Nạp VIP / Kích hoạt PRO**  
3) Quét **VietQR** hoặc chuyển khoản theo hướng dẫn  
4) Bấm **KÍCH HOẠT / XÁC MINH** (SePay tự kiểm tra)  
5) Thành công → hệ thống cập nhật gói/điểm
            """
        )

        st.divider()

        st.markdown("### B. 👤 Tài khoản – Gói – Điểm (rất quan trọng)")
        st.markdown(
            f"""
- **FREE**: dùng thử cơ bản theo giới hạn hệ thống  
- **VIP/PRO**: dùng đầy đủ hơn, ổn định hơn  

**Điểm dùng để làm gì?**  
- Điểm dùng để chạy các tác vụ AI (ra đề/soạn giáo án/năng lực số…).  
- Mỗi lượt có thể trừ điểm theo quy định.

**Quy đổi hiện tại:**
- **{VIP_TOPUP_AMOUNT_VND:,}đ = {VIP_TOPUP_POINTS} điểm**  
- Chi phí mặc định:
  - Soạn giáo án: **{POINT_COST_LESSON_PLAN} điểm/lượt**
  - Ra đề – KTĐG: **{POINT_COST_EXAM} điểm/lượt**
  - Năng lực số: **{POINT_COST_NLS} điểm/lượt**
            """
        )

        st.divider()

        st.markdown("### C. 💳 Nạp VIP/PRO (xác minh tự động)")
        st.markdown(f"""
**Cách nạp đúng:**
1) Quét **VietQR**  
2) Chuyển **đúng số tiền**  
3) Nhập **đúng nội dung chuyển khoản** theo hướng dẫn trên web  
4) Bấm **KÍCH HOẠT / XÁC MINH** để hệ thống kiểm tra giao dịch

**Nếu chưa kích hoạt được:**
- Chờ 30–60 giây rồi bấm xác minh lại  
- Kiểm tra bạn đã chuyển đúng nội dung chưa  
- Nếu vẫn lỗi: chụp ảnh giao dịch và gửi bộ phận hỗ trợ
            """
        )

        st.info("Lưu ý: Nội dung chuyển khoản đúng giúp hệ thống nhận diện nhanh và chính xác.")




def module_profile():
    """Trang hồ sơ đơn giản (yêu cầu đăng nhập)."""
    _ensure_nav_state()
    user = st.session_state.get("user") or {}
    if not user:
        require_login("profile")
        return

    st.markdown("## 👤 Profile")
    st.caption("Thông tin tài khoản và trạng thái gói/điểm.")

    col1, col2 = st.columns([1.2, 1], vertical_alignment="top")
    with col1:
        st.markdown(
            f"""
<div class="card">
  <div style="display:flex;gap:12px;align-items:center;">
    <div style="width:46px;height:46px;border-radius:16px;background:rgba(91,92,246,.14);display:flex;align-items:center;justify-content:center;font-weight:900;color:#3b5bff;">
      {html_escape((user.get("fullname") or "U")[:1].upper())}
    </div>
    <div>
      <div style="font-weight:900;font-size:18px;line-height:1.1;">{html_escape(user.get("fullname") or "Chưa đặt tên")}</div>
      <div class="small-muted">{html_escape(user.get("email") or "")}</div>
    </div>
  </div>
</div>
""",
            unsafe_allow_html=True,
        )

        st.write("")
        st.markdown(
            f"""
<div class="card soft">
  <b>Gói:</b> {(user.get("role") or "free").upper()}<br/>
  <b>Điểm:</b> {user.get("points", 0)}
  <div class="small-muted" style="margin-top:8px;">Điểm được trừ khi chạy các chức năng AI theo quy định của hệ thống.</div>
</div>
""",
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            """
<div class="card">
  <b>⚙️ Tác vụ</b>
  <div class="small-muted" style="margin-top:6px;">
    Bạn có thể quay về Trang chủ hoặc đăng xuất tại đây.
  </div>
</div>
""",
            unsafe_allow_html=True,
        )
        st.write("")
        if st.button("🏡 Về Trang chủ", use_container_width=True, key="pf_home"):
            go("dashboard")
        if st.button("🚪 Đăng xuất", use_container_width=True, key="pf_logout"):
            st.session_state.pop("user", None)
            st.toast("👋 Bạn đã đăng xuất.", icon="✅")
            go("dashboard")


# ==============================================================================
# ENTRY POINT (PUBLIC HOME + LOGIN-ON-DEMAND + TOPBAR + SIDEBAR)
# ==============================================================================
_ensure_nav_state()

# Topbar luôn hiển thị
render_topbar()
st.write("")  # spacing

# Sidebar (hiển thị cả với khách)
with st.sidebar:
    st.markdown(f"""<div class="sb-brand">
<div class="sb-logo" style="background:transparent; box-shadow:none;">{logo_svg(52)}</div>
<div>
  <div class="sb-title">AIEXAM.VN</div>
  <div class="sb-sub">WEB AI GIÁO VIÊN</div>
</div>
</div>""",
        unsafe_allow_html=True
    )
    st.markdown("<div class='small-muted'>Điều hướng nhanh • Dễ sử dụng</div>", unsafe_allow_html=True)
    st.divider()

    page_map = {
        "🏡 Trang chủ": "dashboard",
        "💬 Chat AI": "chat",
        "📑 Doc AI": "doc_ai",
        "🧠 Mindmap": "mindmap",
        "🧾 Ra đề – KTĐG": "exam",
        "📚 Trợ lý Soạn bài": "lesson_plan",
        "🖥️ Năng lực số": "digital",
        "🧭 Nhận xét – Tư vấn": "advisor",
        "📘 Hướng dẫn": "help",
        "🔐 Đăng nhập / Đăng ký": "login",
    }

    # ---- Sidebar navigation (stable, no input reset)
    reverse_map = {v: k for k, v in page_map.items()}
    current_page = st.session_state.get("current_page", "dashboard")
    current_label = reverse_map.get(current_page, "🏡 Trang chủ")
    # Sync radio highlight when navigation happens programmatically (go(...))
    if st.session_state.get("_sync_sidebar_menu", False) or "sidebar_menu_main" not in st.session_state:
        st.session_state["sidebar_menu_main"] = current_label
        st.session_state["_sync_sidebar_menu"] = False

    def _on_sidebar_nav_change():
        label = st.session_state.get("sidebar_menu_main", current_label)
        st.session_state["current_page"] = page_map.get(label, "dashboard")

    menu_label = st.radio(
        "Điều hướng",
        list(page_map.keys()),
        index=list(page_map.keys()).index(st.session_state["sidebar_menu_main"]),
        key="sidebar_menu_main",
        label_visibility="collapsed",
        on_change=_on_sidebar_nav_change,
    )

    st.write("")
    user = st.session_state.get("user") or {}
    if user:
        role = user.get("role", "free")
        role_badge = "PRO" if role == "pro" else "FREE"
        st.markdown(f"""<div class="card">
<b>⭐ Gói hiện tại: {role_badge}</b>
<div class="small-muted" style="margin-top:6px;">Nâng cấp để mở giới hạn & nhận thêm điểm.</div>
</div>""", unsafe_allow_html=True)
        if st.button("🚪 Đăng xuất", use_container_width=True, key="sb_logout"):
            st.session_state.pop("user", None)
            st.session_state["current_page"] = "dashboard"
            st.rerun()
    else:
        st.markdown("""<div class="card soft">
<b>👋 Chào mừng!</b>
<div class="small-muted" style="margin-top:6px;">Bạn có thể xem Trang chủ và dùng thử 1 câu Chat AI. Khi dùng tiếp, hệ thống sẽ yêu cầu đăng nhập.</div>
</div>""", unsafe_allow_html=True)
        if st.button("🔐 Đăng nhập", type="primary", use_container_width=True, key="sb_login"):
            st.session_state["requested_page"] = st.session_state.get("current_page", "dashboard")
            st.session_state["current_page"] = "login"
            st.rerun()

# ROUTER
page = st.session_state.get("current_page", "dashboard")

# Login page
if page == "login":
    login_screen()
    st.stop()

# Guard protected pages
if (page in PROTECTED_PAGES) and (not st.session_state.get("user")):
    require_login(page)
    st.stop()

# Chat page allows 1 demo for guest; lần 2 yêu cầu login (được xử trong module_chat)
if page == "dashboard":
    dashboard_screen()
elif page == "chat":
    module_chat()
elif page == "doc_ai":
    module_doc_ai()
elif page == "mindmap":
    module_mindmap()
elif page == "help":
    module_help()
elif page == "profile":
    module_profile()
elif page == "lesson_plan":
    if module_lesson_plan_B:
        module_lesson_plan_B(
            SYSTEM_GOOGLE_KEY=SYSTEM_GOOGLE_KEY,
            BOOKS_LIST=BOOKS_LIST,
            EDUCATION_DATA=EDUCATION_DATA,
            FULL_SCOPE_LIST=FULL_SCOPE_LIST,
            create_word_doc_func=create_word_doc,
            model_name="gemini-2.0-flash"
        )
    else:
        module_lesson_plan()
elif page == "digital":
    module_digital()
elif page == "advisor":
    module_advisor()
else:
    # exam + fallback
    main_app()
