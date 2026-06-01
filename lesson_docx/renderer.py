# -*- coding: utf-8 -*-
"""
Renderer — CODE 100% quyết định layout DOCX. Render JSON → DOCX khớp PDF mẫu.

AI KHÔNG được động vào layout. Tất cả định nghĩa font/lề/bảng/thứ tự mục
được hardcode tại đây để đảm bảo file giáo viên tải về CHÍNH XÁC theo
mẫu chuẩn 2026-2027 mà tổ chuyên môn yêu cầu.

Spec layout (khớp PDF "Giao_an_Toan_4_Bai_1_Chuan_hoa.pdf"):
- Khổ A4, lề T2-B2-L3-R2 cm, TNR 13pt, giãn dòng 1.15
- Header trang: trái Phòng/Trường + phải Quốc hiệu/Tiêu ngữ (có gạch ngang)
- Footer trang: "Kế hoạch bài dạy [tên bài]" italic centered
- 7 mục La Mã (I-VII), bảng hoạt động 3 cột, bảng đánh giá 4 cột
- Phần ký: Tổ trưởng/Chuyên môn | Giáo viên
"""

from __future__ import annotations

import io
from typing import Any, Dict, List

from docx import Document
from docx.document import Document as _DocClass
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.table import _Cell, Table

from .schema import SAFE_PLACEHOLDER


# ============================================================================
# Constants — hardcode theo PDF mẫu
# ============================================================================

FONT_NAME = "Times New Roman"
SIZE_NORMAL = 13
SIZE_TITLE_MAIN = 16        # "KẾ HOẠCH BÀI DẠY"
SIZE_TITLE_SUB = 14         # "Môn:...", "Bài X. ...", "I. ...", "II. ..."
SIZE_TABLE = 12             # nội dung bảng
SIZE_HEADER_FOOTER = 11     # header/footer trang
LINE_SPACING = 1.15
DOTTED_LINE = "." * 100     # dòng kẻ chấm trong mục VII


# ============================================================================
# Helpers cấp thấp — thao tác run/paragraph/cell/table
# ============================================================================

def _set_run(run, size: int = SIZE_NORMAL, bold: bool = False, italic: bool = False, underline: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    # Set East Asian font để chữ tiếng Việt render đúng trong Word
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    except Exception:
        pass


def _add_para(doc, text: str, *, size: int = SIZE_NORMAL, bold: bool = False, italic: bool = False,
              align=None, space_after: int = 4, keep_with_next: bool = False, indent_cm: float = 0.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic)
    return p


def _add_labeled_para(doc, label: str, value: str, *, size: int = SIZE_NORMAL, indent_cm: float = 0.0):
    """Đoạn dạng 'Label: value' — label in đậm, value thường."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(4)
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    r1 = p.add_run(label)
    _set_run(r1, size=size, bold=True)
    r2 = p.add_run(value)
    _set_run(r2, size=size)
    return p


def _add_bullet(doc, text: str, *, size: int = SIZE_NORMAL, indent_cm: float = 0.5):
    """Bullet kiểu '– text' (gạch ngang en-dash) như mẫu PDF."""
    return _add_para(doc, f"– {text}", size=size, indent_cm=indent_cm, space_after=2)


def _set_cell_borders(cell: _Cell) -> None:
    """Set border 4 cạnh cho 1 cell (single line)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "auto")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_shading(cell: _Cell, fill_hex: str) -> None:
    """Tô màu nền cell (fill_hex không có '#', vd 'EEF4FF')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _set_row_as_header(row) -> None:
    """Row sẽ lặp lại trên đầu khi bảng tràn sang trang mới."""
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _write_cell(cell: _Cell, text: str, *, bold: bool = False, italic: bool = False,
                size: int = SIZE_TABLE, align=None) -> None:
    """Ghi text vào cell với format chuẩn. Hỗ trợ multi-line qua \\n."""
    # Xoá paragraph mặc định, tự thêm
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lines = (text or "").split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        if align is not None:
            p.alignment = align
        run = p.add_run(line)
        _set_run(run, size=size, bold=bold, italic=italic)


def _make_borderless_table(doc, rows: int, cols: int) -> Table:
    """Bảng không viền, dùng cho header trang và phần ký."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Tắt border tất cả cell
    for row in tbl.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tc_borders.append(b)
            tc_pr.append(tc_borders)
    return tbl


def _make_bordered_table(doc, rows: int, cols: int, header_row: bool = True) -> Table:
    """Bảng có viền full, header row màu nhạt + lặp khi tràn trang, cantSplit từng row."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, row in enumerate(tbl.rows):
        _set_row_cant_split(row)
        if header_row and i == 0:
            _set_row_as_header(row)
            for cell in row.cells:
                _set_cell_shading(cell, "EEF4FF")
        for cell in row.cells:
            _set_cell_borders(cell)
    return tbl


# ============================================================================
# Page setup, Header & Footer
# ============================================================================

def _setup_page(doc: _DocClass) -> None:
    """A4, lề T2-B2-L3-R2 cm. Style Normal: TNR 13pt, line spacing 1.15."""
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME  # type: ignore[attr-defined]
    normal.font.size = Pt(SIZE_NORMAL)  # type: ignore[attr-defined]
    normal.paragraph_format.line_spacing = LINE_SPACING  # type: ignore[attr-defined]
    normal.paragraph_format.space_after = Pt(4)  # type: ignore[attr-defined]
    normal.paragraph_format.space_before = Pt(0)  # type: ignore[attr-defined]
    try:
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)  # type: ignore[attr-defined]
    except Exception:
        pass


def _render_header(doc: _DocClass, department: str, school: str) -> None:
    """Header trang: trái UBND + TRƯỜNG, phải Quốc hiệu + Tiêu ngữ.

    CHỈ hiển thị trên TRANG 1 (Word feature: different_first_page_header_footer).
    Các trang sau không có header — chỉ có footer "Kế hoạch bài dạy [tên bài]".
    """
    sec = doc.sections[0]

    # Bật cờ "first page header khác" → trang 1 có header, trang 2+ trống.
    sec.different_first_page_header_footer = True

    # `first_page_header` chỉ tồn tại sau khi flag được bật ở trên.
    first_header = sec.first_page_header

    # Xoá paragraph mặc định của first-page header
    for p in list(first_header.paragraphs):
        p._element.getparent().remove(p._element)

    # Dùng table 2 cột không viền để có 2 cột song song
    tbl = first_header.add_table(rows=1, cols=2, width=Cm(16))  # type: ignore[attr-defined]
    tbl.autofit = False
    # Tắt border tất cả cell
    for cell in tbl.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tc_borders.append(b)
        tc_pr.append(tc_borders)

    left, right = tbl.rows[0].cells

    # CỘT TRÁI: UBND ... / TRƯỜNG ... (in hoa, đậm)
    # Format mới theo yêu cầu: cấp UBND huyện/xã thay vì Phòng GD&ĐT;
    # tên trường không bị ép thêm tiền tố "TIỂU HỌC".
    left.text = ""
    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(f"UBND {department}")
    _set_run(r1, size=SIZE_HEADER_FOOTER, bold=True)
    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(0)
    # Nếu school đã bắt đầu bằng "TRƯỜNG" thì dùng nguyên; không thì thêm "TRƯỜNG " phía trước.
    school_clean = school.strip().upper() if school else ""
    if school_clean.startswith("TRƯỜNG"):
        line2 = school_clean
    else:
        line2 = f"TRƯỜNG {school_clean}" if school_clean else "TRƯỜNG"
    r2 = p2.add_run(line2)
    _set_run(r2, size=SIZE_HEADER_FOOTER, bold=True)

    # CỘT PHẢI: Quốc hiệu (in hoa đậm) / Tiêu ngữ (đậm, gạch chân)
    right.text = ""
    rp1 = right.paragraphs[0]
    rp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp1.paragraph_format.space_after = Pt(0)
    rr1 = rp1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    _set_run(rr1, size=SIZE_HEADER_FOOTER, bold=True)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp2.paragraph_format.space_after = Pt(0)
    rr2 = rp2.add_run("Độc lập - Tự do - Hạnh phúc")
    _set_run(rr2, size=SIZE_HEADER_FOOTER, bold=True, underline=True)


def _render_footer(doc: _DocClass, lesson_title: str) -> None:
    """Footer trang: italic centered '[tên bài]' — lặp MỌI trang (cả trang 1).

    Vì `different_first_page_header_footer = True` đã bật trong _render_header,
    cần set CẢ first_page_footer LẪN footer thường, không thì trang 1 sẽ trống footer.
    """
    sec = doc.sections[0]
    text = f"Kế hoạch bài dạy {lesson_title}"

    def _write_footer(footer_obj) -> None:
        # Xoá paragraph mặc định nếu có
        for p in list(footer_obj.paragraphs):
            if p._element.getparent() is not None:
                p._element.getparent().remove(p._element)
        p = footer_obj.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        _set_run(run, size=SIZE_HEADER_FOOTER - 1, italic=True)

    _write_footer(sec.first_page_footer)
    _write_footer(sec.footer)


# ============================================================================
# Body sections
# ============================================================================

def _render_title_block(doc: _DocClass, lesson_info: Dict[str, str]) -> None:
    """Khối tiêu đề: KẾ HOẠCH BÀI DẠY / Môn-Lớp / Bài X. TÊN / Bộ sách-Thời lượng."""
    _add_para(doc, "KẾ HOẠCH BÀI DẠY",
              size=SIZE_TITLE_MAIN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=4, keep_with_next=True)
    _add_para(doc, f"Môn: {lesson_info['subject']} - Lớp {lesson_info['grade']}",
              size=SIZE_TITLE_SUB, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=2, keep_with_next=True)
    lesson_line = lesson_info["lessonTitle"]
    if lesson_info.get("lessonNumber"):
        lesson_line = f"{lesson_info['lessonNumber']}. {lesson_line}"
    _add_para(doc, lesson_line,
              size=SIZE_TITLE_SUB, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=4, keep_with_next=True)
    _add_para(doc, f"Bộ sách: {lesson_info['textbookSeries']} - Thời lượng: {lesson_info['duration']}",
              size=SIZE_NORMAL, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=8)


def _render_info_table(doc: _DocClass, lesson_info: Dict[str, str]) -> None:
    """Bảng thông tin 2x3 (3 dòng × 4 cột), labels đậm."""
    pairs = [
        ("Tuần/Tiết PPCT", f"{lesson_info.get('week', '...')}/{lesson_info.get('period', '...')}",
         "Ngày dạy", lesson_info.get("teachingDate", SAFE_PLACEHOLDER)),
        ("Giáo viên", lesson_info.get("teacherName", SAFE_PLACEHOLDER),
         "Lớp", lesson_info.get("grade", "")),
        ("Năm học", lesson_info.get("schoolYear", SAFE_PLACEHOLDER) if "schoolYear" in lesson_info else "...",
         "Trang SGK/Học liệu", lesson_info.get("sgkPages", SAFE_PLACEHOLDER)),
    ]
    tbl = _make_bordered_table(doc, rows=3, cols=4, header_row=False)
    tbl.autofit = False
    for i, (l1, v1, l2, v2) in enumerate(pairs):
        cells = tbl.rows[i].cells
        _write_cell(cells[0], l1, bold=True, size=SIZE_TABLE)
        _write_cell(cells[1], v1, size=SIZE_TABLE)
        _write_cell(cells[2], l2, bold=True, size=SIZE_TABLE)
        _write_cell(cells[3], v2, size=SIZE_TABLE)
    _add_para(doc, "", space_after=4)


def _render_heading(doc: _DocClass, text: str) -> None:
    """Tiêu đề mục La Mã (I, II, III...) — 14pt in hoa đậm."""
    _add_para(doc, text, size=SIZE_TITLE_SUB, bold=True,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, keep_with_next=True)


def _render_subheading(doc: _DocClass, text: str) -> None:
    """Tiêu đề con (1., 2., 3...) — 13pt đậm."""
    _add_para(doc, text, size=SIZE_NORMAL, bold=True,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, keep_with_next=True)


def _render_objectives(doc: _DocClass, obj: Dict[str, Any], subject: str) -> None:
    """Mục I — YÊU CẦU CẦN ĐẠT."""
    _render_heading(doc, "I. YÊU CẦU CẦN ĐẠT")
    _add_para(doc, obj.get("intro", "Sau bài học, học sinh đạt được các yêu cầu sau:"),
              italic=False, space_after=4)

    _render_subheading(doc, f"1. Năng lực đặc thù môn {subject}")
    for item in obj.get("specificCompetencies", []):
        _add_bullet(doc, item)

    _render_subheading(doc, "2. Năng lực chung")
    for item in obj.get("generalCompetencies", []):
        _add_bullet(doc, item)

    _render_subheading(doc, "3. Phẩm chất")
    for item in obj.get("qualities", []):
        _add_bullet(doc, item)


def _render_teaching_materials(doc: _DocClass, mat: Dict[str, List[str]]) -> None:
    """Mục II — ĐỒ DÙNG DẠY HỌC."""
    _render_heading(doc, "II. ĐỒ DÙNG DẠY HỌC")
    _render_subheading(doc, "1. Giáo viên")
    for item in mat.get("teacher", []):
        _add_bullet(doc, item)
    _render_subheading(doc, "2. Học sinh")
    for item in mat.get("students", []):
        _add_bullet(doc, item)


def _render_teaching_process(doc: _DocClass, activities: List[Dict[str, Any]]) -> None:
    """Mục III — bảng 3 cột (GV / HS / Sản phẩm-Đánh giá) cho mỗi hoạt động."""
    _render_heading(doc, "III. CÁC HOẠT ĐỘNG DẠY HỌC CHỦ YẾU")
    for idx, act in enumerate(activities, 1):
        # Tiêu đề hoạt động: "X. Tên hoạt động (X phút)"
        title = act.get("title", "").strip()
        duration = act.get("duration", "").strip()
        header_text = f"{idx}. {title}"
        if duration:
            header_text += f" ({duration})"
        _add_para(doc, header_text, size=SIZE_NORMAL, bold=True,
                  space_after=2, keep_with_next=True)

        # Mục tiêu (label đậm + nội dung thường)
        objective = act.get("objective", "").strip()
        if objective:
            _add_labeled_para(doc, "Mục tiêu: ", objective)

        # Bảng 3 cột
        rows = act.get("rows", [])
        if not rows:
            continue
        tbl = _make_bordered_table(doc, rows=1 + len(rows), cols=3, header_row=True)
        tbl.autofit = False
        # Header row
        hcells = tbl.rows[0].cells
        _write_cell(hcells[0], "Hoạt động của giáo viên", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(hcells[1], "Hoạt động của học sinh", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(hcells[2], "Sản phẩm/đánh giá", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        # Data rows
        for i, row in enumerate(rows, 1):
            cells = tbl.rows[i].cells
            _write_cell(cells[0], row.get("teacherActivities", ""))
            _write_cell(cells[1], row.get("studentActivities", ""))
            _write_cell(cells[2], row.get("productAndAssessment", ""))

        _add_para(doc, "", space_after=4)


def _render_differentiation(doc: _DocClass, diff: Dict[str, str]) -> None:
    """Mục IV — ĐIỀU CHỈNH, HỖ TRỢ VÀ PHÂN HÓA."""
    _render_heading(doc, "IV. ĐIỀU CHỈNH, HỖ TRỢ VÀ PHÂN HÓA")
    items = [
        ("Học sinh còn hạn chế: ", diff.get("weakerStudents", SAFE_PLACEHOLDER)),
        ("Học sinh hoàn thành tốt: ", diff.get("advancedStudents", SAFE_PLACEHOLDER)),
        ("Lớp thiếu thiết bị: ", diff.get("limitedResources", SAFE_PLACEHOLDER)),
    ]
    for label, value in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        r0 = p.add_run("– ")
        _set_run(r0, size=SIZE_NORMAL)
        r1 = p.add_run(label)
        _set_run(r1, size=SIZE_NORMAL, bold=True)
        r2 = p.add_run(value)
        _set_run(r2, size=SIZE_NORMAL)


def _render_assessment(doc: _DocClass, assess: List[Dict[str, str]]) -> None:
    """Mục V — ĐÁNH GIÁ THƯỜNG XUYÊN — bảng 4 cột."""
    _render_heading(doc, "V. ĐÁNH GIÁ THƯỜNG XUYÊN TRONG BÀI HỌC")
    if not assess:
        _add_para(doc, SAFE_PLACEHOLDER, italic=True)
        return
    tbl = _make_bordered_table(doc, rows=1 + len(assess), cols=4, header_row=True)
    tbl.autofit = False
    # Header
    h = tbl.rows[0].cells
    _write_cell(h[0], "Nội dung đánh giá", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(h[1], "Mức cần đạt", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(h[2], "Phương pháp", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(h[3], "Công cụ/minh chứng", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Data
    for i, row in enumerate(assess, 1):
        cells = tbl.rows[i].cells
        _write_cell(cells[0], row.get("content", ""))
        _write_cell(cells[1], row.get("criteria", ""))
        _write_cell(cells[2], row.get("method", ""))
        _write_cell(cells[3], row.get("tool", ""))
    _add_para(doc, "", space_after=4)


def _render_worksheet(doc: _DocClass, ws: List[Dict[str, str]]) -> None:
    """Mục VI — PHIẾU HỌC TẬP THAM KHẢO."""
    if not ws:
        return  # mục VI là optional
    _render_heading(doc, "VI. PHIẾU HỌC TẬP THAM KHẢO")
    for item in ws:
        title = item.get("title", "").strip()
        body = item.get("body", "").strip()
        if title:
            _add_para(doc, title, bold=True, space_after=2, keep_with_next=True)
        if body:
            for line in body.split("\n"):
                _add_bullet(doc, line) if line.strip() else None


def _render_post_lesson_adjustment(doc: _DocClass) -> None:
    """Mục VII — ĐIỀU CHỈNH SAU BÀI DẠY: 3 mục con, mỗi mục có 2 dòng kẻ chấm."""
    _render_heading(doc, "VII. ĐIỀU CHỈNH SAU BÀI DẠY")
    sections = [
        "1. Những điểm thực hiện tốt:",
        "2. Nội dung cần điều chỉnh/bổ sung:",
        "3. Học sinh cần hỗ trợ thêm:",
    ]
    for label in sections:
        _add_para(doc, label, bold=True, space_after=2)
        _add_para(doc, DOTTED_LINE, space_after=0)
        _add_para(doc, DOTTED_LINE, space_after=6)


def _render_signature(doc: _DocClass) -> None:
    """Phần ký cuối: Tổ trưởng/Chuyên môn | Giáo viên (bảng 2 cột không viền)."""
    _add_para(doc, "", space_after=12)
    tbl = _make_borderless_table(doc, rows=3, cols=2)
    tbl.autofit = False
    # Hàng 1: chức danh in hoa đậm
    _write_cell(tbl.rows[0].cells[0], "TỔ TRƯỞNG/CHUYÊN MÔN",
                bold=True, size=SIZE_NORMAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tbl.rows[0].cells[1], "GIÁO VIÊN",
                bold=True, size=SIZE_NORMAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Hàng 2: (Ký, ghi rõ họ tên)
    _write_cell(tbl.rows[1].cells[0], "(Ký, ghi rõ họ tên)",
                italic=True, size=SIZE_NORMAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tbl.rows[1].cells[1], "(Ký, ghi rõ họ tên)",
                italic=True, size=SIZE_NORMAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Hàng 3: chừa chỗ ký tay (khoảng trắng)
    _write_cell(tbl.rows[2].cells[0], "\n\n\n", size=SIZE_NORMAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tbl.rows[2].cells[1], "\n\n\n", size=SIZE_NORMAL, align=WD_ALIGN_PARAGRAPH.CENTER)


# ============================================================================
# Public API
# ============================================================================

def render_lesson_docx(data: Dict[str, Any]) -> bytes:
    """Render JSON giáo án → DOCX bytes (khớp 100% PDF mẫu).

    Args:
        data: dict đã được normalizer làm sạch. Schema xem `lesson_docx/schema.py`.

    Returns:
        Bytes của file .docx, sẵn sàng cho st.download_button.

    Raises:
        KeyError nếu thiếu trường BẮT BUỘC (caller nên chạy validator trước).
    """
    doc = Document()
    doc.core_properties.title = data.get("lessonInfo", {}).get("lessonTitle", "Ke hoach bai day")

    _setup_page(doc)

    document_info = data["documentInfo"]
    lesson_info = data["lessonInfo"]

    _render_header(doc,
                   department=document_info.get("department", SAFE_PLACEHOLDER),
                   school=document_info.get("school", ""))
    _render_footer(doc, lesson_title=lesson_info.get("lessonTitle", ""))

    _render_title_block(doc, lesson_info)
    # Bảng thông tin cần `schoolYear` (nằm ở documentInfo, không phải lessonInfo).
    # Gộp tạm để hàm dùng:
    info_for_table = dict(lesson_info)
    info_for_table["schoolYear"] = document_info.get("schoolYear", SAFE_PLACEHOLDER)
    _render_info_table(doc, info_for_table)

    _render_objectives(doc, data["objectives"], subject=lesson_info["subject"])
    _render_teaching_materials(doc, data["teachingMaterials"])
    _render_teaching_process(doc, data["teachingProcess"])
    _render_differentiation(doc, data["differentiation"])
    _render_assessment(doc, data.get("assessment", []))
    _render_worksheet(doc, data.get("worksheet", []))
    _render_post_lesson_adjustment(doc)
    _render_signature(doc)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
