# -*- coding: utf-8 -*-
"""
Module: Soạn giáo án AI nâng cao (chuẩn CTGDPT 2018 + CV 2345/BGDĐT-GDTH)
Tích hợp vào AIEXAM.VN — gói PRO+

Tính năng chính:
- OCR ảnh trang SGK bằng Gemini Vision
- Đọc file PDF/DOCX/TXT giáo viên upload
- Tìm bài học khớp trong kho metadata SGK nội bộ
- Đánh giá độ tin cậy nguồn (4 mức)
- Sinh giáo án HTML đầy đủ 12 mục theo CV 2345
- Quality gate 15 tiêu chí
- Xuất Word .docx + HTML A4

Cách gọi từ app.py:
    from lesson_plan_advanced import module_lesson_plan_advanced
    module_lesson_plan_advanced(
        api_key=SYSTEM_GOOGLE_KEY,
        point_check=require_points_or_block,
        point_cost=35,
    )
"""

from __future__ import annotations

import datetime as dt
import difflib
import html
import io
import json
import os
import re
import textwrap
import unicodedata
import urllib.parse
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

# Dependencies optional — module tự degrade nếu thiếu.
# Lưu lý do import fail vào _IMPORT_ERRORS để hiển thị/log rõ ràng, tránh "nuốt" lỗi.
_IMPORT_ERRORS: dict[str, str] = {}

# streamlit là hard dependency — app không chạy được nếu thiếu nên không wrap try/except.
import streamlit as st
import streamlit.components.v1 as components

try:
    from bs4 import BeautifulSoup, NavigableString  # type: ignore[import-not-found]
except Exception as _e:
    BeautifulSoup = None
    NavigableString = None
    _IMPORT_ERRORS["beautifulsoup4"] = f"{type(_e).__name__}: {_e}"

try:
    from PIL import Image
except Exception as _e:
    Image = None
    _IMPORT_ERRORS["Pillow"] = f"{type(_e).__name__}: {_e}"

try:
    from pypdf import PdfReader  # type: ignore[import-not-found]
except Exception as _e:
    PdfReader = None
    _IMPORT_ERRORS["pypdf"] = f"{type(_e).__name__}: {_e}"

try:
    import mammoth  # type: ignore[import-not-found]
except Exception as _e:
    mammoth = None
    _IMPORT_ERRORS["mammoth"] = f"{type(_e).__name__}: {_e}"

try:
    from docx import Document as _RealDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _RealWDAlign
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT as _RealWDTableAlign,
        WD_CELL_VERTICAL_ALIGNMENT as _RealWDVAlign,
    )
    from docx.shared import Inches as _RealInches, Mm as _RealMm, Pt as _RealPt

    # Re-export về tên gốc với kiểu Any để Pylance không cảnh báo Optional
    # khi runtime guard ở _create_docx đã đảm bảo không None.
    Document: Any = _RealDocument
    WD_ALIGN_PARAGRAPH: Any = _RealWDAlign
    WD_TABLE_ALIGNMENT: Any = _RealWDTableAlign
    WD_CELL_VERTICAL_ALIGNMENT: Any = _RealWDVAlign
    Inches: Any = _RealInches
    Mm: Any = _RealMm
    Pt: Any = _RealPt
except Exception as _e:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    Inches = None
    Mm = None
    Pt = None
    _IMPORT_ERRORS["python-docx"] = f"{type(_e).__name__}: {_e}"


# =============================================================================
# 1. Cấu hình
# =============================================================================

APP_VERSION = "2026.05-pro-v2"
BOOK_SERIES_DEFAULT = "Kết nối tri thức với cuộc sống"
BOOK_SERIES_LOCKED = "Kết nối tri thức với cuộc sống"  # CHỈ DUY NHẤT bộ sách này
SCHOOL_YEAR_DEFAULT = "2026-2027"  # Áp dụng từ năm học 2026-2027 trở đi
SCHOOL_YEAR_MIN = 2026  # Năm bắt đầu (chặn AI tham chiếu chương trình cũ)
PUBLISHER = "Nhà xuất bản Giáo dục Việt Nam"
DEFAULT_MODEL = "gemini-3.1-flash"
APP_KEY = "lp_adv"

APPROVED_DOMAINS = [
    "hoclieuso.nxbgd.vn",
    "hanhtrangso.nxbgd.vn",
    "nxbgd.vn",
    "moet.gov.vn",
    "igiaoduc.vn",
    "taphuan.csdl.edu.vn",
]

BOOK_SERIES_OPTIONS = [
    "Kết nối tri thức với cuộc sống",  # CHỈ duy nhất — bộ sách áp dụng từ 2026-2027
]

SUBJECTS_BY_LEVEL: Dict[str, Dict[str, Any]] = {
    "Tiểu học": {
        "grades": ["1", "2", "3", "4", "5"],
        "subjects": [
            "Tiếng Việt", "Toán", "Đạo đức", "Tự nhiên và Xã hội", "Khoa học",
            "Lịch sử và Địa lí", "Tin học", "Công nghệ", "Hoạt động trải nghiệm",
            "Mĩ thuật", "Âm nhạc", "Giáo dục thể chất", "Tiếng Anh",
        ],
        "default_duration": 35,
    },
    "THCS": {
        "grades": ["6", "7", "8", "9"],
        "subjects": [
            "Ngữ văn", "Toán", "Tiếng Anh", "Khoa học tự nhiên",
            "Lịch sử và Địa lí", "Tin học", "Công nghệ", "Giáo dục công dân",
            "Hoạt động trải nghiệm, hướng nghiệp", "Mĩ thuật", "Âm nhạc",
            "Giáo dục thể chất",
        ],
        "default_duration": 45,
    },
    "THPT": {
        "grades": ["10", "11", "12"],
        "subjects": [
            "Ngữ văn", "Toán", "Tiếng Anh", "Vật lí", "Hóa học", "Sinh học",
            "Lịch sử", "Địa lí", "Giáo dục kinh tế và pháp luật", "Tin học",
            "Công nghệ", "Hoạt động trải nghiệm, hướng nghiệp",
            "Giáo dục thể chất", "Giáo dục quốc phòng và an ninh",
        ],
        "default_duration": 45,
    },
}

DIGITAL_COMPETENCIES = [
    "Khai thác dữ liệu, thông tin và nội dung số",
    "Giao tiếp và hợp tác trong môi trường số",
    "Sáng tạo nội dung số",
    "An toàn số",
    "Giải quyết vấn đề trong môi trường số",
    "Sử dụng công cụ AI/học liệu số có trách nhiệm",
]

# Năng lực đặc thù theo môn — chuẩn CTGDPT 2018 (Thông tư 32/2018/TT-BGDĐT)
# Dùng để chèn cụ thể vào mục V, tránh GV viết "phù hợp với chuẩn đầu ra" chung chung
SUBJECT_SPECIFIC_COMPETENCIES: Dict[str, List[str]] = {
    "Toán": [
        "Năng lực tư duy và lập luận toán học",
        "Năng lực mô hình hóa toán học",
        "Năng lực giải quyết vấn đề toán học",
        "Năng lực giao tiếp toán học",
        "Năng lực sử dụng công cụ, phương tiện học toán",
    ],
    "Tiếng Việt": [
        "Năng lực ngôn ngữ: đọc",
        "Năng lực ngôn ngữ: viết",
        "Năng lực ngôn ngữ: nói",
        "Năng lực ngôn ngữ: nghe",
        "Năng lực văn học (cảm thụ, thưởng thức văn bản)",
    ],
    "Ngữ văn": [
        "Năng lực ngôn ngữ (đọc, viết, nói và nghe)",
        "Năng lực văn học",
    ],
    "Tin học": [
        "Năng lực sử dụng và quản lí các phương tiện công nghệ thông tin và truyền thông (NLa)",
        "Năng lực ứng xử phù hợp trong môi trường số (NLb)",
        "Năng lực giải quyết vấn đề với sự hỗ trợ của công nghệ thông tin và truyền thông (NLc)",
        "Năng lực ứng dụng công nghệ thông tin và truyền thông trong học và tự học (NLd)",
        "Năng lực hợp tác trong môi trường số (NLe)",
    ],
    "Khoa học": [
        "Năng lực nhận thức khoa học tự nhiên",
        "Năng lực tìm hiểu môi trường tự nhiên xung quanh",
        "Năng lực vận dụng kiến thức, kĩ năng đã học",
    ],
    "Khoa học tự nhiên": [
        "Năng lực nhận thức khoa học tự nhiên",
        "Năng lực tìm hiểu tự nhiên",
        "Năng lực vận dụng kiến thức, kĩ năng đã học",
    ],
    "Tự nhiên và Xã hội": [
        "Năng lực nhận thức khoa học",
        "Năng lực tìm hiểu môi trường tự nhiên và xã hội xung quanh",
        "Năng lực vận dụng kiến thức, kĩ năng đã học vào thực tiễn",
    ],
    "Lịch sử và Địa lí": [
        "Năng lực nhận thức khoa học Lịch sử",
        "Năng lực nhận thức khoa học Địa lí",
        "Năng lực tìm hiểu Lịch sử / Địa lí",
        "Năng lực vận dụng kiến thức, kĩ năng đã học",
    ],
    "Đạo đức": [
        "Năng lực điều chỉnh hành vi đạo đức",
        "Năng lực phát triển bản thân",
        "Năng lực tìm hiểu và tham gia hoạt động kinh tế - xã hội",
    ],
    "Tiếng Anh": [
        "Năng lực giao tiếp bằng tiếng Anh: nghe",
        "Năng lực giao tiếp bằng tiếng Anh: nói",
        "Năng lực giao tiếp bằng tiếng Anh: đọc",
        "Năng lực giao tiếp bằng tiếng Anh: viết",
    ],
    "Mĩ thuật": [
        "Năng lực quan sát và nhận thức thẩm mĩ",
        "Năng lực sáng tạo và ứng dụng thẩm mĩ",
        "Năng lực phân tích và đánh giá thẩm mĩ",
    ],
    "Âm nhạc": [
        "Năng lực thể hiện âm nhạc",
        "Năng lực cảm thụ và hiểu biết âm nhạc",
        "Năng lực ứng dụng và sáng tạo âm nhạc",
    ],
    "Giáo dục thể chất": [
        "Năng lực chăm sóc sức khỏe",
        "Năng lực vận động cơ bản",
        "Năng lực hoạt động thể dục thể thao",
    ],
    "Công nghệ": [
        "Năng lực nhận thức công nghệ",
        "Năng lực giao tiếp công nghệ",
        "Năng lực sử dụng công nghệ",
        "Năng lực đánh giá công nghệ",
        "Năng lực thiết kế kĩ thuật",
    ],
}

# 4 mục La Mã CHUẨN MẪU "KẾ HOẠCH BÀI DẠY" — theo spec giáo viên tiểu học nộp tổ chuyên môn.
# (Đơn giản hoá từ CV 2345 — chỉ giữ 4 mục bắt buộc, gộp năng lực/phẩm chất vào mục I.)
REQUIRED_SECTIONS = [
    "yêu cầu cần đạt",           # I — gồm kiến thức kĩ năng, năng lực, phẩm chất
    "đồ dùng dạy học",            # II — giáo viên / học sinh
    "các hoạt động dạy học",     # III — bảng 4 cột tiến trình
    "điều chỉnh sau bài dạy",    # IV — dòng kẻ chấm
]

SAMPLE_LESSONS: List[Dict[str, Any]] = [
    {
        "lesson_id": "KNTT_TINHOC_5_TIM_KIEM_DEMO",
        "book_series": BOOK_SERIES_DEFAULT,
        "grade": "5",
        "subject": "Tin học",
        "unit_name": "Chủ đề C. Tổ chức lưu trữ, tìm kiếm và trao đổi thông tin",
        "lesson_title": "Tìm kiếm thông tin trên Internet",
        "alternate_titles": ["Tìm kiếm thông tin", "Từ khóa tìm kiếm"],
        "suggested_duration": 35,
        "keywords": ["Internet", "tìm kiếm", "từ khóa", "máy tìm kiếm", "an toàn"],
        "learning_objectives": [
            "Nêu được vai trò của từ khóa khi tìm kiếm thông tin trên Internet.",
            "Thực hiện được thao tác tìm kiếm thông tin đơn giản theo yêu cầu học tập.",
            "Biết lựa chọn kết quả phù hợp, đáng tin cậy và an toàn ở mức độ phù hợp với lứa tuổi.",
        ],
        "main_content_summary": "Học sinh nhận biết nhu cầu tìm kiếm thông tin, biết sử dụng từ khóa phù hợp, thực hành tìm kiếm thông tin học tập trên Internet và bước đầu đánh giá kết quả tìm kiếm.",
        "digital_competencies": [
            "Khai thác dữ liệu, thông tin và nội dung số",
            "An toàn số",
            "Giải quyết vấn đề trong môi trường số",
        ],
        "official_url": "https://hoclieuso.nxbgd.vn/",
        "source_status": "metadata_demo",
        "last_reviewed": "2026-05-18",
    },
    {
        "lesson_id": "KNTT_TOAN_5_ON_TAP_PHAN_SO_DEMO",
        "book_series": BOOK_SERIES_DEFAULT,
        "grade": "5",
        "subject": "Toán",
        "unit_name": "Ôn tập và bổ sung",
        "lesson_title": "Ôn tập về phân số",
        "alternate_titles": ["Phân số", "Ôn tập phân số"],
        "suggested_duration": 35,
        "keywords": ["phân số", "tử số", "mẫu số", "so sánh", "rút gọn"],
        "learning_objectives": [
            "Củng cố được khái niệm phân số, tử số, mẫu số.",
            "Thực hiện được một số thao tác cơ bản với phân số ở mức độ phù hợp.",
            "Vận dụng phân số để giải quyết tình huống đơn giản trong học tập và đời sống.",
        ],
        "main_content_summary": "Bài học giúp học sinh ôn tập kiến thức trọng tâm về phân số và luyện tập vận dụng.",
        "digital_competencies": [
            "Khai thác dữ liệu, thông tin và nội dung số",
            "Giải quyết vấn đề trong môi trường số",
        ],
        "official_url": "https://hoclieuso.nxbgd.vn/",
        "source_status": "metadata_demo",
        "last_reviewed": "2026-05-18",
    },
    {
        "lesson_id": "KNTT_TIENGVIET_5_DOC_MO_RONG_DEMO",
        "book_series": BOOK_SERIES_DEFAULT,
        "grade": "5",
        "subject": "Tiếng Việt",
        "unit_name": "Đọc - Viết - Nói và nghe",
        "lesson_title": "Đọc mở rộng",
        "alternate_titles": ["Bài đọc mở rộng", "Đọc sách báo"],
        "suggested_duration": 35,
        "keywords": ["đọc mở rộng", "chia sẻ bài đọc", "cảm nhận"],
        "learning_objectives": [
            "Đọc được văn bản phù hợp với lứa tuổi theo yêu cầu đọc mở rộng.",
            "Ghi lại được thông tin chính và cảm nhận ngắn về văn bản đã đọc.",
            "Chia sẻ được bài đọc với bạn bằng lời nói rõ ràng, tự tin.",
        ],
        "main_content_summary": "Học sinh đọc, ghi chép và chia sẻ văn bản đọc mở rộng theo định hướng của giáo viên.",
        "digital_competencies": [
            "Khai thác dữ liệu, thông tin và nội dung số",
            "Giao tiếp và hợp tác trong môi trường số",
            "An toàn số",
        ],
        "official_url": "https://hoclieuso.nxbgd.vn/",
        "source_status": "metadata_demo",
        "last_reviewed": "2026-05-18",
    },
]


@dataclass
class _SourceBundle:
    metadata_matches: List[Dict[str, Any]] = field(default_factory=list)
    uploaded_text: str = ""
    uploaded_text_notes: List[str] = field(default_factory=list)
    image_context: str = ""
    image_notes: List[str] = field(default_factory=list)
    link_context: str = ""
    link_notes: List[str] = field(default_factory=list)
    teacher_note: str = ""
    trust_level: str = "Trung bình"
    trust_explanation: str = "Chưa có nguồn SGK xác thực trực tiếp."


# =============================================================================
# 2. Tiện ích
# =============================================================================

def _k(name: str) -> str:
    return f"{APP_KEY}_{name}"


def _norm(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "").lower().strip()
    text = re.sub(r"[​‌‍﻿]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text


def _slugify(text: str, max_len: int = 90) -> str:
    text = _norm(text)
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    text = text.replace("đ", "d").replace("Đ", "d")
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    return text[:max_len] or "giao_an"


def _safe_json_dumps(data: Any, max_chars: int = 12000) -> str:
    raw = json.dumps(data, ensure_ascii=False, indent=2)
    if len(raw) > max_chars:
        return raw[:max_chars] + "\n... [đã rút gọn]"
    return raw


def _get_metadata_dir() -> Path:
    val = ""
    try:
        val = st.secrets.get("SGK_DATA_DIR", "")
    except Exception:
        val = ""
    return Path(os.getenv("SGK_DATA_DIR") or val or "data_sgk_ket_noi")


def _ensure_sample_metadata() -> None:
    d = _get_metadata_dir()
    d.mkdir(parents=True, exist_ok=True)
    sample = d / "lessons_sample_demo.json"
    if not sample.exists():
        sample.write_text(
            json.dumps(SAMPLE_LESSONS, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


def _load_metadata() -> List[Dict[str, Any]]:
    _ensure_sample_metadata()
    lessons: List[Dict[str, Any]] = []
    for file in _get_metadata_dir().rglob("*.json"):
        try:
            data = json.loads(file.read_text(encoding="utf-8"))
            items = data if isinstance(data, list) else (data.get("lessons") if isinstance(data, dict) else None)
            if items is None and isinstance(data, dict) and data.get("lesson_title"):
                items = [data]
            for item in items or []:
                if isinstance(item, dict) and item.get("lesson_title"):
                    item["_metadata_file"] = str(file)
                    lessons.append(item)
        except Exception:
            continue
    seen, unique = set(), []
    for item in lessons:
        key = item.get("lesson_id") or f"{item.get('grade')}|{item.get('subject')}|{item.get('lesson_title')}"
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)
    return unique


def _score_match(lesson: Dict[str, Any], grade: str, subject: str, title: str) -> float:
    score = 0.0
    t_title = _norm(title)
    l_grade = str(lesson.get("grade", "")).strip()
    l_subj = _norm(str(lesson.get("subject", "")))
    t_subj = _norm(subject)

    if l_grade == str(grade).strip():
        score += 25
    elif l_grade:
        score -= 15

    if l_subj == t_subj:
        score += 25
    elif t_subj in l_subj or l_subj in t_subj:
        score += 15
    elif l_subj:
        score -= 10

    titles = [str(lesson.get("lesson_title", ""))] + [str(x) for x in lesson.get("alternate_titles", [])]
    best = 0.0
    for nt in titles:
        nt_n = _norm(nt)
        if not nt_n:
            continue
        r = difflib.SequenceMatcher(None, t_title, nt_n).ratio() * 35
        if t_title and t_title in nt_n:
            r += 15
        if nt_n and nt_n in t_title:
            r += 10
        best = max(best, r)
    score += best

    kw_hits = sum(1 for kw in lesson.get("keywords", []) if _norm(str(kw)) and _norm(str(kw)) in t_title)
    score += min(kw_hits * 5, 15)
    return max(0, min(score, 100))


def _search_metadata(grade: str, subject: str, title: str, limit: int = 5) -> List[Dict[str, Any]]:
    out = []
    for lesson in _load_metadata():
        s = _score_match(lesson, grade, subject, title)
        if s >= 25:
            item = dict(lesson)
            item["_match_score"] = round(s, 1)
            out.append((s, item))
    out.sort(key=lambda x: x[0], reverse=True)
    return [item for _, item in out[:limit]]


# =============================================================================
# 3. Đọc nguồn xác thực
# =============================================================================

def _extract_text_from_uploads(files) -> Tuple[str, List[str]]:
    if not files:
        return "", []
    all_text, notes = [], []
    for f in files:
        name = getattr(f, "name", "file")
        ext = Path(name).suffix.lower().strip(".")
        raw = f.getvalue()
        text = ""
        try:
            if ext == "txt":
                text = raw.decode("utf-8", errors="ignore")
            elif ext == "pdf":
                if PdfReader is None:
                    notes.append(f"{name}: cần cài pypdf (pip install pypdf).")
                else:
                    reader = PdfReader(io.BytesIO(raw))
                    pages = []
                    for i, p in enumerate(reader.pages[:20], 1):
                        try:
                            pages.append(f"[Trang {i}]\n{p.extract_text() or ''}")
                        except Exception:
                            pages.append(f"[Trang {i}] Không trích được.")
                    text = "\n\n".join(pages)
            elif ext in {"docx", "doc"}:
                if mammoth is not None:
                    text = mammoth.extract_raw_text(io.BytesIO(raw)).value or ""
                elif Document is not None and ext == "docx":
                    doc = Document(io.BytesIO(raw))
                    text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    notes.append(f"{name}: cần cài mammoth hoặc python-docx.")
            else:
                notes.append(f"{name}: định dạng chưa hỗ trợ.")
        except Exception as e:
            notes.append(f"{name}: lỗi khi đọc - {e}")
        text = (text or "").strip()
        if text:
            all_text.append(f"\n--- FILE: {name} ---\n{text[:12000]}")
            notes.append(f"{name}: đã trích {len(text)} ký tự.")
        else:
            notes.append(f"{name}: không có text. Nếu là ảnh scan, hãy upload sang ô ảnh trang SGK.")
    return "\n\n".join(all_text)[:30000], notes


def _call_gemini_text(api_key: str, prompt: str, model_name: str = DEFAULT_MODEL, temperature: float = 0.25) -> str:
    try:
        import google.generativeai as genai
    except Exception as e:
        raise RuntimeError("Chưa cài google-generativeai. Chạy: pip install google-generativeai") from e
    genai.configure(api_key=api_key)  # type: ignore[attr-defined]
    model = genai.GenerativeModel(model_name)  # type: ignore[attr-defined]
    resp = model.generate_content(
        prompt,
        generation_config={"temperature": temperature, "top_p": 0.9, "max_output_tokens": 12000},  # type: ignore[arg-type]
    )
    return getattr(resp, "text", "") or ""


def _extract_image_context(image_files, api_key: str, model_name: str) -> Tuple[str, List[str]]:
    if not image_files:
        return "", []
    if not api_key:
        return "", ["Có ảnh nhưng chưa có API key nên chưa OCR được."]
    if Image is None:
        return "", ["Chưa cài Pillow (pip install pillow)."]
    try:
        import google.generativeai as genai
    except Exception:
        return "", ["Chưa cài google-generativeai."]

    notes, pil_imgs = [], []
    for f in image_files[:8]:
        try:
            img = Image.open(io.BytesIO(f.getvalue()))
            pil_imgs.append(img)
            notes.append(f"{getattr(f, 'name', 'ảnh')}: nhận ảnh {img.size[0]}x{img.size[1]}.")
        except Exception as e:
            notes.append(f"{getattr(f, 'name', 'ảnh')}: lỗi - {e}")
    if not pil_imgs:
        return "", notes

    prompt = textwrap.dedent(f"""
        Bạn là trợ lý giáo dục Việt Nam. Hãy đọc các ảnh trang SGK/học liệu giáo viên cung cấp.
        Yêu cầu:
        1. Trích xuất tên bài, mục/chủ đề, nội dung chính, câu hỏi/bài tập, yêu cầu cần đạt nếu có.
        2. Không bịa nội dung không nhìn thấy trong ảnh.
        3. Nếu ảnh mờ/thiếu trang, ghi rõ phần chưa chắc chắn.
        4. Tóm tắt ngắn gọn để dùng làm căn cứ soạn giáo án theo bộ sách {BOOK_SERIES_DEFAULT}.
        Trả lời bằng tiếng Việt, dạng gạch đầu dòng rõ ràng.
    """).strip()
    try:
        genai.configure(api_key=api_key)  # type: ignore[attr-defined]
        model = genai.GenerativeModel(model_name)  # type: ignore[attr-defined]
        resp = model.generate_content([prompt, *pil_imgs])
        text = getattr(resp, "text", "") or ""
        notes.append("Đã OCR ảnh bằng AI." if text.strip() else "AI không trả về nội dung.")
        return text[:20000], notes
    except Exception as e:
        notes.append(f"Lỗi khi OCR ảnh: {e}")
        return "", notes


def _is_approved_url(url: str) -> bool:
    u = (url or "").strip().lower()
    return any(d in u for d in APPROVED_DOMAINS)


def _fetch_official_links(links_text: str, timeout: int = 12) -> Tuple[str, List[str]]:
    links = re.findall(r"https?://[^\s,;]+", links_text or "")
    if not links:
        return "", []
    if BeautifulSoup is None:
        return "", ["Chưa cài beautifulsoup4 nên chưa đọc link được."]
    ctxs, notes = [], []
    for url in links[:5]:
        if not _is_approved_url(url):
            notes.append(f"Bỏ qua link ngoài danh sách duyệt: {url}")
            continue
        try:
            req = urllib.request.Request(
                url,
                headers={
                    "User-Agent": "Mozilla/5.0 aiexam-lesson-plan-advanced",
                    "Accept-Language": "vi,en;q=0.8",
                },
            )
            with urllib.request.urlopen(req, timeout=timeout) as r:
                ctype = r.headers.get("content-type", "")
                if "text/html" not in ctype and "application/xhtml" not in ctype:
                    notes.append(f"{url}: không phải HTML ({ctype}).")
                    continue
                raw = r.read().decode("utf-8", errors="ignore")
            soup = BeautifulSoup(raw, "html.parser")
            for tag in soup(["script", "style", "noscript", "svg"]):
                tag.decompose()
            text = re.sub(r"\n{3,}", "\n\n", soup.get_text("\n"))
            text = re.sub(r"[ \t]{2,}", " ", text).strip()
            if len(_norm(text)) < 120:
                notes.append(f"{url}: trang dạng động, chưa lấy được nội dung.")
            else:
                ctxs.append(f"\n--- LINK: {url} ---\n{text[:8000]}")
                notes.append(f"{url}: đã lấy {min(len(text), 8000)} ký tự.")
        except Exception as e:
            notes.append(f"{url}: lỗi - {e}")
    return "\n\n".join(ctxs)[:20000], notes


def _determine_trust(b: _SourceBundle) -> _SourceBundle:
    has_upload = bool(b.uploaded_text.strip() or b.image_context.strip())
    has_good_meta = bool(b.metadata_matches and b.metadata_matches[0].get("_match_score", 0) >= 70)
    has_meta = bool(b.metadata_matches)
    has_link = bool(b.link_context.strip())

    if has_upload and has_good_meta:
        b.trust_level = "Rất cao"
        b.trust_explanation = "Có ảnh/file giáo viên cung cấp và bài khớp tốt trong kho metadata."
    elif has_upload:
        b.trust_level = "Rất cao"
        b.trust_explanation = "Có ảnh/file bài học do giáo viên cung cấp."
    elif has_good_meta:
        b.trust_level = "Cao"
        b.trust_explanation = "Tìm thấy bài khớp tốt trong kho metadata SGK."
    elif has_link and has_meta:
        b.trust_level = "Cao"
        b.trust_explanation = "Có dữ liệu link chính thống và bài khớp metadata."
    elif has_meta:
        b.trust_level = "Trung bình"
        b.trust_explanation = "Có bài tương đối khớp trong metadata, cần giáo viên kiểm tra lại."
    elif has_link:
        b.trust_level = "Trung bình"
        b.trust_explanation = "Có dữ liệu từ link chính thống, chưa khớp metadata."
    else:
        b.trust_level = "Thấp"
        b.trust_explanation = "Chưa có nguồn SGK xác thực. AI tạo bản nháp theo khung, giáo viên cần rà soát."
    return b


# =============================================================================
# 4. Prompt + Sinh HTML
# =============================================================================

def _build_prompt(*, phong_gd_dt, school, department, teacher, school_year, teaching_date, location,
                  level, grade, class_name, subject, book_series, lesson_title, duration,
                  period_note, sgk_pages, digital_level, diff_level, bundle: _SourceBundle) -> str:
    meta_json = _safe_json_dumps(bundle.metadata_matches[:3], 15000)
    # ÉP bộ sách = KNTT, năm học >= 2026-2027
    enforced_book = BOOK_SERIES_LOCKED
    enforced_year = school_year if school_year and school_year >= SCHOOL_YEAR_DEFAULT else SCHOOL_YEAR_DEFAULT

    # Tín hiệu nguồn để AI biết phải bám gì
    has_meta = bool(bundle.metadata_matches)
    has_image = bool(bundle.image_context.strip())
    has_upload = bool(bundle.uploaded_text.strip())
    has_link = bool(bundle.link_context.strip())
    source_signal = []
    if has_image: source_signal.append("ẢNH SGK (ưu tiên cao nhất — đây là nội dung thật của bài)")
    if has_upload: source_signal.append("FILE GV upload (ưu tiên cao)")
    if has_meta: source_signal.append("METADATA KNTT (khung tham chiếu)")
    if has_link: source_signal.append("LINK CHÍNH THỐNG")
    source_priority = " > ".join(source_signal) if source_signal else "KHÔNG CÓ NGUỒN XÁC THỰC — phải ghi cảnh báo trong mục II"

    return textwrap.dedent(f"""
Bạn là chuyên gia giáo dục Việt Nam, chuyên thiết kế giáo án theo định hướng phát triển phẩm chất, năng lực, tích hợp năng lực số.
Hãy soạn GIÁO ÁN HOÀN CHỈNH bằng tiếng Việt, dùng ngay được cho giáo viên.

# NGUYÊN TẮC TUYỆT ĐỐI (KHÔNG VI PHẠM)
1. **BỘ SÁCH DUY NHẤT**: Chỉ tham chiếu "Kết nối tri thức với cuộc sống" (NXB Giáo dục Việt Nam).
   - NGHIÊM CẤM tham chiếu Cánh Diều, Chân trời sáng tạo, hoặc SGK chương trình cũ.
   - Nếu trong tài liệu nguồn có nội dung từ bộ sách khác, BỎ QUA và chỉ dùng KNTT.
2. **NĂM HỌC ÁP DỤNG**: Từ năm học {enforced_year} trở đi. Không tham chiếu chương trình tiền 2026.
3. **THỨ TỰ ƯU TIÊN NGUỒN**: {source_priority}.
   - Khi các nguồn mâu thuẫn: ẢNH SGK / FILE GV THẮNG metadata. Metadata chỉ là khung tham chiếu phụ.
   - Khi viết mục III (Yêu cầu cần đạt), V (Năng lực đặc thù), IX (Tiến trình): nếu nội dung lấy từ ảnh/file/metadata, thêm chú thích nguồn ngắn ở cuối câu, vd: "(theo ảnh SGK)", "(theo metadata KNTT)". Nếu suy luận sư phạm chung, không cần chú thích.
4. Không bịa nội dung SGK. Chỉ khẳng định khi có căn cứ từ metadata, ảnh/file GV hoặc link chính thống của KNTT.
   - Nếu không chắc chắn nội dung cụ thể của bài (vd tên nhân vật, số liệu, công thức), viết theo hướng tổng quát rồi đánh dấu [GV kiểm tra SGK] ngay tại chỗ.
5. Nếu thiếu nguồn xác thực, vẫn tạo giáo án theo khung chuẩn nhưng ghi rõ ở mục "Nguồn học liệu": cần GV kiểm tra lại SGK KNTT bản {enforced_year}.
6. Không trích nguyên văn SGK quá 50 từ liên tục. Chỉ tóm tắt nội dung cần thiết.
7. Giáo án phải thực tế, GV có thể tải Word và chỉnh sửa dùng ngay.
   - Câu lệnh GV phải cụ thể: "GV chiếu slide trang ..., hỏi: ..." chứ KHÔNG viết "GV giới thiệu bài học".
   - Hoạt động HS phải quan sát được: "HS đọc thầm đoạn 1, gạch chân từ khóa" chứ KHÔNG viết "HS lắng nghe".
8. Tích hợp năng lực số phải tự nhiên, đúng hoạt động, không hình thức.
9. Tổng thời gian các hoạt động phải bằng đúng {duration} phút (sai số ±2 phút). Bảng tiến trình IX phải có cột "Thời gian" với số phút cụ thể từng hoạt động.
10. Trả về HTML sạch, KHÔNG dùng markdown, KHÔNG bọc ```html.

# VÍ DỤ ĐỊNH DẠNG MẪU (chỉ tham khảo cấu trúc, không copy nội dung)
<section><h2>III. Yêu cầu cần đạt</h2>
  <p><strong>Về kiến thức:</strong></p>
  <ul><li>Nêu được vai trò của từ khóa khi tìm kiếm thông tin trên Internet (theo metadata KNTT).</li></ul>
  <p><strong>Về kĩ năng:</strong></p>
  <ul><li>Thực hiện được thao tác gõ từ khóa vào ô tìm kiếm và mở kết quả phù hợp.</li></ul>
</section>

Ví dụ 1 hàng trong bảng tiến trình (4 cột):
<tr>
  <td>GV chiếu 3 hình ảnh (sách, bản đồ, máy tính), hỏi: "Khi cần tìm thông tin về một loài chim, em sẽ làm gì?"</td>
  <td>HS giơ tay trả lời cá nhân, mỗi em nêu 1 cách tìm thông tin</td>
  <td>3-5 câu trả lời từ HS được nói trước lớp</td>
  <td>GV quan sát, nhận xét miệng, ghi nhận sự tham gia</td>
</tr>

# THÔNG TIN BÀI DẠY
- Phòng GD&ĐT: {phong_gd_dt or "GV bổ sung"}
- Trường: {school or "Chưa nhập"}
- Tổ / khối chuyên môn: {department or "Chưa nhập"}
- Địa danh (dùng cho chữ ký): {location or "GV bổ sung"}
- Họ tên giáo viên: {teacher or "Chưa nhập"}
- **Năm học: {enforced_year}** (ÁP DỤNG TỪ 2026-2027)
- Ngày dạy: {teaching_date.strftime('%d/%m/%Y')}
- Cấp học: {level}
- Lớp: {class_name or grade}
- Môn học / hoạt động giáo dục: {subject}
- **Bộ sách: {enforced_book}** (KHÔNG dùng bộ khác)
- NXB: {PUBLISHER}
- Tên bài học: {lesson_title}
- Thời lượng / số tiết: {duration} phút
- Tuần / tiết / PPCT: {period_note or "Không bắt buộc"}
- Trang SGK / nguồn học liệu: {sgk_pages or "GV bổ sung sau"}
- Mức tích hợp năng lực số: {digital_level}
- Mức phân hóa: {diff_level}
- Mức tin cậy nguồn: {bundle.trust_level} - {bundle.trust_explanation}

# NGUỒN 1 - METADATA SGK KNTT
{meta_json if meta_json.strip() else "Không có metadata khớp."}

# NGUỒN 2 - FILE GIÁO VIÊN
{bundle.uploaded_text[:18000] if bundle.uploaded_text.strip() else "Không có."}

# NGUỒN 3 - ẢNH TRANG SÁCH KNTT (AI ĐỌC)
{bundle.image_context[:15000] if bundle.image_context.strip() else "Không có."}

# NGUỒN 4 - LINK CHÍNH THỐNG (chỉ chấp nhận hoclieuso.nxbgd.vn, hanhtrangso.nxbgd.vn, moet.gov.vn)
{bundle.link_context[:12000] if bundle.link_context.strip() else "Không có."}

# GHI CHÚ GIÁO VIÊN
{bundle.teacher_note or "Không có."}

# KHUNG GIÁO ÁN CHUẨN — Mẫu "KẾ HOẠCH BÀI DẠY" giáo viên tiểu học nộp tổ chuyên môn
# (Đơn giản hoá theo spec: 4 mục La Mã I-IV; gộp năng lực/phẩm chất vào mục I.)
Tạo HTML có cấu trúc CHÍNH XÁC sau:
<article class="lesson-plan">
  <!-- A. PHẦN ĐẦU VĂN BẢN: bảng 2 cột không viền -->
  <!--    Bên trái:  PHÒNG GD&ĐT: {department}   /   TRƯỜNG: {school}   (in hoa, đậm) -->
  <!--    Bên phải:  CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM (đậm, căn giữa)
                    Độc lập – Tự do – Hạnh phúc (đậm, gạch chân, căn giữa) -->

  <!-- B. TIÊU ĐỀ CHÍNH -->
  <h1>KẾ HOẠCH BÀI DẠY</h1>

  <!-- C. THÔNG TIN BÀI DẠY (không phải <section>, viết bằng các <p>) -->
  <p><strong>Môn học/HĐGD:</strong> {subject}</p>
  <p><strong>Lớp:</strong> {class_name or grade}</p>
  <p><strong>Tên bài học:</strong> {lesson_title}</p>
  <p><strong>Thời lượng:</strong> {duration} phút</p>
  <p><strong>Tuần:</strong> ... &nbsp;&nbsp; <strong>Tiết PPCT:</strong> ...</p>
  <p><strong>Ngày dạy:</strong> {teaching_date.strftime('%d/%m/%Y')}</p>
  <p><strong>Giáo viên:</strong> {teacher}</p>
  <p><strong>Bộ sách:</strong> {enforced_book} — NXBGDVN</p>

  <!-- D. 4 MỤC NỘI DUNG BẮT BUỘC -->

  <section><h2>I. YÊU CẦU CẦN ĐẠT</h2>
    <h3>1. Kiến thức, kĩ năng</h3>
    <p>Liệt kê 2-4 yêu cầu kiến thức/kĩ năng cụ thể, dùng động từ Bloom đo được (nêu, đọc, viết, so sánh, vận dụng...). Mỗi yêu cầu 1 câu ngắn, gắn với nội dung bài học cụ thể (KHÔNG viết chung chung).</p>
    <ul><li>...</li><li>...</li></ul>

    <h3>2. Năng lực</h3>
    <p><strong>a) Năng lực chung:</strong> chọn 1-2 trong 3 NL chung CTGDPT 2018 (tự chủ và tự học; giao tiếp và hợp tác; giải quyết vấn đề và sáng tạo) có biểu hiện cụ thể qua bài.</p>
    <ul><li>...</li></ul>
    <p><strong>b) Năng lực đặc thù môn {subject}:</strong> liệt kê 1-2 NLĐT theo Thông tư 32/2018, gắn với hoạt động cụ thể.</p>
    <ul><li>...</li></ul>

    <h3>3. Phẩm chất</h3>
    <p>Chọn 1-2 phẩm chất trong: yêu nước, nhân ái, chăm chỉ, trung thực, trách nhiệm. Mỗi phẩm chất 1 câu nêu biểu hiện qua bài.</p>
    <ul><li>...</li></ul>
  </section>

  <section><h2>II. ĐỒ DÙNG DẠY HỌC</h2>
    <h3>1. Giáo viên</h3>
    <ul><li>SGK {enforced_book} môn {subject} lớp {class_name or grade}.</li><li>...</li></ul>
    <h3>2. Học sinh</h3>
    <ul><li>SGK, vở ghi, đồ dùng học tập theo bài.</li><li>...</li></ul>
  </section>

  <section><h2>III. CÁC HOẠT ĐỘNG DẠY HỌC CHỦ YẾU</h2>
    BẮT BUỘC đúng 3 hoạt động (KHÔNG được thêm pha thứ 4, KHÔNG được tách "Củng cố/Dặn dò" riêng):
      1. Hoạt động Mở đầu (Khởi động)
      2. Hoạt động Hình thành kiến thức mới / Luyện tập, thực hành (tuỳ kiểu bài)
      3. Hoạt động Vận dụng, trải nghiệm

    Với MỖI hoạt động, dùng đúng cấu trúc sau (KHÔNG đổi):
      <h3>1. Hoạt động Mở đầu (khoảng X phút)</h3>
      <p><strong>Mục tiêu:</strong> 1 câu, nói rõ HS đạt được gì sau hoạt động.</p>
      <p><strong>Cách tiến hành:</strong></p>
      <table>
        <tr>
          <th>Hoạt động của giáo viên</th>
          <th>Hoạt động của học sinh</th>
          <th>Sản phẩm học tập</th>
          <th>Đánh giá</th>
        </tr>
        <tr>
          <td>Lệnh GV cụ thể, có câu nói/câu hỏi mẫu trong dấu ngoặc kép.</td>
          <td>Hành động HS quan sát được (đọc, viết, trả lời, thảo luận nhóm...).</td>
          <td>Sản phẩm cụ thể HS tạo ra (bài làm trong vở, câu trả lời, phiếu HT...).</td>
          <td>GV quan sát/nhận xét miệng/kiểm tra vở/chấm bài tại lớp.</td>
        </tr>
        <!-- có thể thêm nhiều <tr> nếu hoạt động chia nhiều bước -->
      </table>
    Tổng thời gian 3 hoạt động = {duration} phút (±2 phút).
  </section>

  <section><h2>IV. ĐIỀU CHỈNH SAU BÀI DẠY</h2>
    <p>.....................................................................................................................</p>
    <p>.....................................................................................................................</p>
    <p>.....................................................................................................................</p>
  </section>

  <!-- E. PHẦN KÝ DUYỆT (đặt cuối article) -->
  <!-- <p style="text-align:right">{location}, ngày DD tháng MM năm YYYY</p> (lấy từ ngày dạy) -->
  <!-- Bảng 2 cột không viền: TỔ TRƯỞNG CHUYÊN MÔN | NGƯỜI SOẠN — dưới mỗi cột "(Kí, ghi rõ họ tên)" -->
</article>

# ĐỊNH DẠNG
- HTML rõ ràng, tiếng Việt chuẩn (đủ dấu, không lỗi chính tả).
- KHÔNG viết "tôi là AI", KHÔNG xin lỗi, KHÔNG markdown, KHÔNG bọc ```html.
- KHÔNG có nội dung ngoài thẻ <article>.
- **CHỈ ĐÚNG 4 MỤC LA MÃ I-IV, KHÔNG THÊM MỤC NÀO (không V, không VI...).**
- Trường thiếu dữ liệu: dùng dấu "..." (3 chấm), KHÔNG viết "GV bổ sung", "Chưa nhập", "theo ảnh SGK", "AIEXAM".
    """).strip()


def _clean_html(raw: str) -> str:
    raw = (raw or "").strip()
    raw = re.sub(r"^```(?:html)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    m = re.search(r"<article[\s\S]*?</article>", raw, flags=re.IGNORECASE)
    if m:
        raw = m.group(0)
    if "<article" not in raw.lower():
        raw = f"<article class='lesson-plan'><h1>GIÁO ÁN</h1><p>{html.escape(raw)}</p></article>"
    return raw.strip()


def _wrap_a4(content_html: str, title: str = "GiaoAn") -> str:
    css = """
    <style>
      * { box-sizing: border-box; }
      body { margin:0; background:#eef2f7; color:#111827; font-family:"Times New Roman",Times,serif; }
      .page { width:210mm; min-height:297mm; margin:18px auto; background:#fff; padding:18mm 16mm; box-shadow:0 10px 28px rgba(15,23,42,.16); }
      article.lesson-plan { font-size:13.5pt; line-height:1.35; }
      h1 { text-align:center; font-size:18pt; margin:0 0 12px; text-transform:uppercase; color:#0f172a; }
      h2 { font-size:14.5pt; margin:14px 0 8px; color:#0f172a; border-left:5px solid #174ea6; padding-left:8px; }
      h3 { font-size:13.8pt; margin:10px 0 6px; }
      p { margin:5px 0; }
      ul, ol { margin:6px 0 6px 24px; }
      table { width:100%; border-collapse:collapse; margin:8px 0 12px; }
      th, td { border:1px solid #4b5563; padding:6px 7px; vertical-align:top; }
      th { background:#eef4ff; text-align:center; font-weight:700; }
      .source-warning { background:#fff7ed; border:1px solid #fed7aa; padding:8px; border-radius:8px; }
      @media print { body { background:#fff; } .page { margin:0; box-shadow:none; width:auto; min-height:auto; } }
    </style>
    """
    return f"<!doctype html><html lang='vi'><head><meta charset='utf-8'><title>{html.escape(title)}</title>{css}</head><body><div class='page'>{content_html}</div></body></html>"


def _create_docx(full_html: str) -> bytes:
    if Document is None or BeautifulSoup is None or NavigableString is None:
        import sys
        missing = []
        if Document is None:
            missing.append("python-docx")
        if BeautifulSoup is None or NavigableString is None:
            missing.append("beautifulsoup4")
        raise RuntimeError(
            f"Thiếu thư viện: {', '.join(missing)}. "
            f"Cài bằng: \"{sys.executable}\" -m pip install {' '.join(missing)} — "
            f"rồi RESTART streamlit."
        )
    NavStr = NavigableString  # alias để type checker narrow ra non-None
    soup = BeautifulSoup(full_html, "html.parser")
    page = soup.find("article") or soup.find("div", class_="page") or soup.body or soup
    doc = Document()
    sec = doc.sections[0]
    # Khổ A4 + lề chuẩn Nghị định 30/2020/NĐ-CP (trái 30mm, phải 20mm, trên/dưới 20mm)
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(20)
    # Style mặc định: Times New Roman 13pt, giãn dòng 1.15 (theo thể thức KHBD), cách sau đoạn 6pt
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"  # type: ignore[attr-defined]
    normal.font.size = Pt(13)  # type: ignore[attr-defined]
    normal.paragraph_format.line_spacing = 1.15  # type: ignore[attr-defined]
    normal.paragraph_format.space_after = Pt(6)  # type: ignore[attr-defined]
    normal.paragraph_format.space_before = Pt(0)  # type: ignore[attr-defined]

    # ----- HEADER: để trống (không brand) — file giáo án nộp tổ chuyên môn không được mang chữ AIEXAM -----
    # Vẫn giữ paragraph để section header tồn tại; nội dung rỗng.
    _ = sec.header.paragraphs[0]

    # ----- FOOTER: chỉ "Trang X / Y" (page numbering) -----
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("Trang ")
    f_run.font.name = "Times New Roman"
    f_run.font.size = Pt(9)

    def _add_page_field(paragraph, field_code: str):
        """Chèn field Word (PAGE hoặc NUMPAGES) vào paragraph để Word tự tính số trang."""
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        fld_begin = _OxmlElement("w:fldChar")
        fld_begin.set(_qn("w:fldCharType"), "begin")
        instr = _OxmlElement("w:instrText")
        instr.set(_qn("xml:space"), "preserve")
        instr.text = f" {field_code} "
        fld_end = _OxmlElement("w:fldChar")
        fld_end.set(_qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    _add_page_field(footer_p, "PAGE")
    sep_run = footer_p.add_run(" / ")
    sep_run.font.name = "Times New Roman"
    sep_run.font.size = Pt(9)
    _add_page_field(footer_p, "NUMPAGES")

    def add_p(text: str, *, bold: bool = False, size: int = 13, align=None, indent_em: int = 0, keep_with_next: bool = False):
        text = re.sub(r"\s+", " ", text or "").strip()
        if not text:
            return
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if indent_em > 0:
            p.paragraph_format.left_indent = Inches(0.25 * indent_em)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        if keep_with_next:
            # Anti-orphan: tiêu đề không nằm cuối trang một mình, luôn đi cùng đoạn ngay sau
            p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)

    def li_direct_text(li) -> str:
        """Text của <li> nhưng bỏ qua <ul>/<ol> con (để không trùng với handle_list đệ quy)."""
        parts = []
        for ch in li.children:
            if isinstance(ch, NavStr):
                parts.append(str(ch))
            elif ch.name in {"ul", "ol"}:
                continue
            else:
                parts.append(ch.get_text(" ", strip=True))
        return re.sub(r"\s+", " ", " ".join(parts)).strip()

    def handle_list(node, depth: int = 0, ordered: bool = False):
        lis = node.find_all("li", recursive=False)
        for idx, li in enumerate(lis, 1):
            marker = f"{idx}. " if ordered else "• "
            add_p(marker + li_direct_text(li), indent_em=depth)
            for nested in li.find_all(["ul", "ol"], recursive=False):
                handle_list(nested, depth=depth + 1, ordered=(nested.name == "ol"))

    def handle_table(t):
        rows = t.find_all("tr")
        if not rows:
            return
        max_cols = max(len(r.find_all(["th", "td"])) for r in rows)
        if max_cols == 0:
            return
        tbl = doc.add_table(rows=len(rows), cols=max_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            # cantSplit: cấm 1 ô trong row bị tách sang trang khác → bảng không vỡ giữa chừng
            tr_pr = tbl.rows[i]._tr.get_or_add_trPr()
            cant_split = _OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
            # Row đầu tiên là header → tblHeader để Word lặp lại header khi bảng tràn trang
            is_header_row = all((cells[j].name == "th") for j in range(min(len(cells), max_cols)) if cells)
            if i == 0 and is_header_row:
                tbl_header = _OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            for j in range(max_cols):
                c = tbl.cell(i, j)
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                txt = cells[j].get_text(" ", strip=True) if j < len(cells) else ""
                c.text = re.sub(r"\s+", " ", txt).strip()
                is_header = j < len(cells) and cells[j].name == "th"
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        if is_header:
                            run.bold = True
        doc.add_paragraph()

    def walk(node):
        for child in node.children:
            if isinstance(child, NavStr):
                continue
            name = child.name
            if name == "h1":
                # Tiêu đề chính "KẾ HOẠCH BÀI DẠY": cỡ 16, in hoa-đậm, căn giữa (theo spec)
                add_p(child.get_text(" ", strip=True), bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
            elif name == "h2":
                # I/II/III/IV: cỡ 14, đậm + anti-orphan
                add_p(child.get_text(" ", strip=True), bold=True, size=14, keep_with_next=True)
            elif name in {"h3", "h4"}:
                add_p(child.get_text(" ", strip=True), bold=True, size=13, keep_with_next=True)
            elif name == "p":
                add_p(child.get_text(" ", strip=True))
            elif name in {"ul", "ol"}:
                handle_list(child, ordered=(name == "ol"))
            elif name == "table":
                handle_table(child)
            elif name in {"section", "article", "div"}:
                walk(child)
            # các tag inline khác (strong, em, span, br...) bỏ qua — đã được get_text gộp ở cha

    walk(page)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _create_doc_html_fallback(full_html: str, title: str) -> bytes:
    """Fallback khi thiếu python-docx/bs4: bọc HTML thành file .doc mà Word/WPS/LibreOffice mở được.

    Kỹ thuật: Word từ Office 97 trở đi nhận diện HTML có namespace `xmlns:w="urn:schemas-microsoft-com:office:word"`
    và mở dưới dạng document có thể chỉnh sửa. Đây là chuẩn 'Word HTML Document' chính thức của Microsoft.
    Lề A4 + Times New Roman 13pt được set qua CSS @page và body để khớp định dạng giáo án.
    """
    safe_title = (title or "GiaoAn").replace('"', "'")
    word_css = """
    <style>
      @page WordSection1 { size: 21cm 29.7cm; margin: 2cm 2cm 2cm 3cm; }
      div.WordSection1 { page: WordSection1; }
      body { font-family: 'Times New Roman', serif; font-size: 13pt; line-height: 1.5; }
      table { border-collapse: collapse; width: 100%; }
      th, td { border: 1px solid #000; padding: 6px; vertical-align: top; }
      th { background: #eef4ff; font-weight: bold; text-align: center; }
      h1 { font-size: 15pt; text-align: center; }
      h2 { font-size: 14pt; }
      h3, h4 { font-size: 13pt; }
    </style>
    """
    # Tách body ra khỏi full_html để nhúng vào WordSection1 đúng cách
    m = re.search(r"<body[^>]*>(.*)</body>", full_html, flags=re.DOTALL | re.IGNORECASE)
    body_inner = m.group(1) if m else full_html
    doc_html = (
        "<html xmlns:o='urn:schemas-microsoft-com:office:office' "
        "xmlns:w='urn:schemas-microsoft-com:office:word' "
        "xmlns='http://www.w3.org/TR/REC-html40'>"
        "<head><meta charset='utf-8'>"
        f"<title>{html.escape(safe_title)}</title>"
        + word_css +
        "</head><body><div class='WordSection1'>"
        + body_inner +
        "</div></body></html>"
    )
    return doc_html.encode("utf-8")


# =============================================================================
# Helpers cho Phase B — smart validation tên bài và phát hiện môn/lớp lệch
# =============================================================================

def smart_capitalize_title(s: str) -> str:
    """Chuẩn hoá tên bài: viết hoa chữ đầu của câu/phụ-câu, giữ phần còn lại nguyên.

    Quy tắc:
    - Chữ đầu tiên của chuỗi → viết hoa
    - Sau dấu chấm/hai chấm/chấm phẩy + khoảng trắng → viết hoa
    - Không động vào dấu tiếng Việt sẵn có (giữ nguyên ă/â/ơ/ư/đ...)

    Ví dụ:
        "bài 1: ôn tập các số đến 100 000" → "Bài 1: Ôn tập các số đến 100 000"
        "TIẾT 1. nhận biết a a"             → "TIẾT 1. Nhận biết a a"
    """
    s = (s or "").strip()
    if not s:
        return ""
    out: list[str] = []
    capitalize_next = True
    for ch in s:
        if capitalize_next and ch.isalpha():
            out.append(ch.upper())
            capitalize_next = False
        else:
            out.append(ch)
        if ch in ".:;":
            capitalize_next = True
        elif not ch.isspace() and ch not in ".:;":
            # Đã gặp ký tự non-space sau dấu → tắt cờ
            pass
    return "".join(out)


# Heuristic phát hiện môn/lớp lệch với tên bài.
# Key = (grade, subject), value = list keyword đặc trưng (lowercase).
# Nguyên tắc: chỉ thêm keyword RẤT đặc trưng cho cặp (grade, subject) đó,
# tránh từ chung như "kể chuyện" vì xuất hiện nhiều cấp. Match dùng word boundary.
_GRADE_SUBJECT_KEYWORDS: dict[tuple[str, str], list[str]] = {
    # ===== TIẾNG VIỆT =====
    ("1", "Tiếng Việt"): [
        "a a", "b b", "c c", "âm", "vần", "đánh vần",
        "chữ a", "chữ b", "chữ c", "nhận biết chữ",
    ],
    ("2", "Tiếng Việt"): ["mở rộng vốn từ", "từ chỉ sự vật", "từ chỉ hoạt động"],
    ("3", "Tiếng Việt"): ["so sánh", "nhân hoá", "câu kể", "câu hỏi", "ôn tập giữa học kì"],
    ("4", "Tiếng Việt"): ["danh từ", "động từ", "tính từ", "câu kể ai làm gì", "đoạn văn miêu tả"],
    ("5", "Tiếng Việt"): ["đại từ", "quan hệ từ", "bài văn tả người", "bài văn tả cảnh", "tả người thân"],

    # ===== TOÁN =====
    ("1", "Toán"): ["đến 10", "trong phạm vi 10", "phép cộng trong phạm vi 10"],
    ("2", "Toán"): ["đến 20", "đến 100", "trong phạm vi 100"],
    ("3", "Toán"): ["đến 1 000", "đến 10 000", "bảng nhân", "bảng chia"],
    ("4", "Toán"): ["100 000", "đến 100 000", "đến 1 000 000", "phân số", "hàng nghìn", "chục nghìn"],
    ("5", "Toán"): [
        "số thập phân", "tỉ số", "phần trăm", "đến 1 tỉ",
        "tỉ lệ bản đồ", "hình tam giác", "hình thang", "diện tích hình thang",
    ],

    # ===== ĐẠO ĐỨC =====
    ("1", "Đạo đức"): ["em là học sinh", "ngồi học đúng tư thế", "lễ phép với người lớn"],
    ("2", "Đạo đức"): ["quý trọng thời gian", "bảo vệ của công", "nhận lỗi và sửa lỗi"],
    ("3", "Đạo đức"): ["tự lập", "kính trọng người lao động", "tự hào về tổ quốc", "ham học hỏi"],
    ("4", "Đạo đức"): [
        "hiếu thảo với ông bà cha mẹ", "tôn trọng tài sản", "bảo vệ môi trường",
        "biết ơn người lao động",
    ],
    ("5", "Đạo đức"): ["em yêu tổ quốc", "hợp tác với những người xung quanh", "kính già yêu trẻ"],

    # ===== TỰ NHIÊN VÀ XÃ HỘI (lớp 1-3) =====
    ("1", "Tự nhiên và Xã hội"): ["gia đình em", "trường học của em", "lớp học của em", "an toàn khi ở nhà"],
    ("2", "Tự nhiên và Xã hội"): ["nghề nghiệp", "cộng đồng địa phương", "cây cối quanh em", "con vật quanh em"],
    ("3", "Tự nhiên và Xã hội"): ["lá cây", "rễ cây", "thân cây", "hoạt động của tim", "hoạt động hô hấp"],

    # ===== KHOA HỌC (lớp 4-5) =====
    ("4", "Khoa học"): [
        "trao đổi chất", "vật chất và năng lượng", "ánh sáng", "âm thanh",
        "nhiệt độ", "vai trò của nước", "ba thể của nước",
    ],
    ("5", "Khoa học"): [
        "sự sinh sản", "vệ sinh tuổi dậy thì", "tre mây song",
        "sự biến đổi hoá học", "cao su", "chất dẻo",
    ],

    # ===== LỊCH SỬ VÀ ĐỊA LÍ (lớp 4-5) =====
    ("4", "Lịch sử và Địa lí"): [
        "hai bà trưng", "ngô quyền", "đinh bộ lĩnh", "lê lợi",
        "sông hồng", "đồng bằng bắc bộ", "đồng bằng nam bộ", "vùng duyên hải",
    ],
    ("5", "Lịch sử và Địa lí"): [
        "cách mạng tháng tám", "chiến dịch điện biên phủ", "bác hồ ra đi tìm đường cứu nước",
        "khí hậu việt nam", "dân cư việt nam", "asean",
    ],

    # ===== TIN HỌC (lớp 3-5) =====
    ("3", "Tin học"): ["máy tính và em", "phần mềm paint", "vẽ trên máy tính"],
    ("4", "Tin học"): ["soạn thảo văn bản", "trình bày văn bản đơn giản"],
    ("5", "Tin học"): ["tìm kiếm thông tin trên internet", "từ khoá tìm kiếm", "trình chiếu", "thư điện tử"],

    # ===== TIẾNG ANH (lớp 3-5) =====
    ("3", "Tiếng Anh"): [
        "unit 1: hello", "what's your name", "how old are you",
        "my school", "my classroom", "colours",
    ],
    ("4", "Tiếng Anh"): [
        "my hobbies", "my birthday", "my day", "my new friends",
        "where are you from",
    ],
    ("5", "Tiếng Anh"): [
        "what did you do", "past simple", "school subjects",
        "future plans", "my future job",
    ],
}


def _phrase_match(needle: str, haystack: str) -> bool:
    """So khớp keyword với word boundary để 'đến 10' KHÔNG match 'đến 100 000'."""
    pattern = r"(?<!\w)" + re.escape(needle) + r"(?!\w)"
    return bool(re.search(pattern, haystack, flags=re.IGNORECASE))


def detect_grade_subject_mismatch(
    title: str, grade: str, subject: str
) -> Optional[str]:
    """Phát hiện môn/lớp đang chọn có khớp tên bài hay không.

    Returns:
        None nếu OK (không phát hiện lệch hoặc khớp với lựa chọn hiện tại).
        Chuỗi cảnh báo nếu phát hiện lệch.
    """
    if not title:
        return None
    t = title.lower()
    matches: list[tuple[str, str, str]] = []
    for (g, s), kws in _GRADE_SUBJECT_KEYWORDS.items():
        for kw in kws:
            if _phrase_match(kw, t):
                matches.append((g, s, kw))
                break
    if not matches:
        return None
    # Nếu tổ hợp (grade, subject) hiện tại có trong matches → OK
    current = (str(grade), subject)
    if current in [(g, s) for g, s, _ in matches]:
        return None
    g, s, kw = matches[0]
    return (
        f"Tên bài chứa từ khoá '{kw}' — thường thuộc **{s} lớp {g}**, "
        f"không khớp với lựa chọn hiện tại (**{subject} lớp {grade}**). "
        "Kiểm tra lại tên bài / lớp / môn?"
    )


# Phrase cấm trong file giáo án nộp tổ chuyên môn (theo spec).
# Các cụm này nếu AI sinh ra phải được phát hiện và thay thế trước khi xuất DOCX.
_BANNED_PHRASES = (
    "AIEXAM",
    "GV bổ sung",
    "bổ sung sau",
    "theo ảnh SGK",
    "Chưa nhập",
    "GV kiểm tra sau",
    "GV kiểm tra SGK",
    "[GV bổ sung]",
)

# Map placeholder → giá trị lịch sự thay thế (giúp giáo viên dễ điền tay sau).
_PLACEHOLDER_REPLACEMENTS = {
    "GV bổ sung": "...",
    "Chưa nhập": "...",
    "bổ sung sau": "...",
    "GV kiểm tra sau": "...",
    "GV kiểm tra SGK": "...",
    "[GV bổ sung]": "...",
    "[GV kiểm tra SGK]": "...",
    "(theo ảnh SGK)": "",
    "(theo metadata KNTT)": "",
}


def _strip_banned_phrases(content_html: str) -> tuple[str, list[str]]:
    """Thay phrase cấm bằng placeholder dấu chấm. Trả về (html_sạch, danh_sách_phrase_đã_thay)."""
    found: list[str] = []
    cleaned = content_html
    for bad, good in _PLACEHOLDER_REPLACEMENTS.items():
        if bad in cleaned:
            found.append(bad)
            cleaned = cleaned.replace(bad, good)
    # Quét thêm phrase cấm chưa có replacement → xoá luôn (an toàn)
    for bad in _BANNED_PHRASES:
        if bad in cleaned and bad not in _PLACEHOLDER_REPLACEMENTS:
            found.append(bad)
            cleaned = cleaned.replace(bad, "...")
    return cleaned, found


def _validate(content_html: str, lesson_title: str, trust_level: str) -> Dict[str, Any]:
    if BeautifulSoup:
        text = _norm(BeautifulSoup(content_html, "html.parser").get_text(" "))
    else:
        text = _norm(re.sub(r"<[^>]+>", " ", content_html))
    checks = []
    for s in REQUIRED_SECTIONS:
        ok = _norm(s) in text
        checks.append((s, ok, "Có" if ok else "Thiếu"))
    title_ok = _norm(lesson_title)[:20] in text if lesson_title else True
    checks.append(("Tên bài học xuất hiện", title_ok, "Có" if title_ok else "Cần kiểm tra"))
    warn_ok = trust_level != "Thấp" or ("cần giáo viên kiểm tra" in text or "chưa có nguồn" in text)
    checks.append(("Cảnh báo khi thiếu nguồn", warn_ok, "Có" if warn_ok else "Thiếu"))

    # Kiểm tra phrase cấm — bắt buộc PASS để giáo án "đủ chuẩn nộp"
    banned_found = [p for p in _BANNED_PHRASES if p.lower() in content_html.lower()]
    no_banned = len(banned_found) == 0
    checks.append((
        "Không chứa phrase cấm (AIEXAM/GV bổ sung/Chưa nhập/...)",
        no_banned,
        "Sạch" if no_banned else f"Có: {', '.join(banned_found[:3])}",
    ))

    ok_count = sum(1 for _, ok, _ in checks if ok)
    score = round(ok_count / len(checks) * 100)
    level = "Tốt" if score >= 90 else "Khá" if score >= 75 else "Cần rà soát" if score >= 60 else "Chưa đạt"
    return {"score": score, "level": level, "checks": checks, "banned_phrases_found": banned_found}


# =============================================================================
# 5. UI
# =============================================================================

def _render_source_status(b: _SourceBundle) -> None:
    color = {"Rất cao": "#16a34a", "Cao": "#2563eb", "Trung bình": "#d97706", "Thấp": "#dc2626"}.get(b.trust_level, "#6b7280")
    st.markdown(
        f"""<div style="border:1px solid #e5e7eb;border-radius:12px;padding:12px;background:#fff;margin:8px 0;">
        <div style="font-weight:800;color:{color};">Mức tin cậy nguồn: {html.escape(b.trust_level)}</div>
        <div style="margin-top:4px;color:#374151;">{html.escape(b.trust_explanation)}</div></div>""",
        unsafe_allow_html=True,
    )
    if b.metadata_matches:
        st.markdown("**Bài khớp trong kho metadata:**")
        for it in b.metadata_matches[:3]:
            st.write(f"- {it.get('subject')} lớp {it.get('grade')} — **{it.get('lesson_title')}** (khớp: {it.get('_match_score')}%)")
    else:
        st.info("Chưa tìm thấy bài khớp trong kho metadata.")
    with st.expander("Nhật ký nguồn dữ liệu", expanded=False):
        for n in b.uploaded_text_notes + b.image_notes + b.link_notes:
            st.write("- " + n)


def module_lesson_plan_advanced(
    *,
    api_key: str = "",
    point_check: Optional[Callable[[int, str], bool]] = None,
    point_cost: int = 35,
    model_name: str = DEFAULT_MODEL,
    docx_renderer: Optional[Callable[[str, str], bytes]] = None,
) -> None:
    """Module soạn giáo án AI nâng cao — chỉ Pro/Admin được vào.
    Caller phải tự kiểm tra role trước khi gọi.
    """
    # Thanh thông tin gọn 1 dòng — không trùng với header gọi từ app.py
    st.markdown(
        f"""<div style="display:flex;flex-wrap:wrap;align-items:center;gap:8px;
        padding:10px 14px;margin-bottom:12px;border:1px solid #e2e8f0;border-radius:10px;
        background:#f8fafc;font-size:13px;color:#334155;">
          <span style="background:#1d4ed8;color:#fff;font-weight:700;padding:2px 8px;border-radius:6px;font-size:12px;">KNTT</span>
          <span><b>Năm học:</b> từ {SCHOOL_YEAR_DEFAULT}</span>
          <span style="color:#94a3b8">·</span>
          <span><b>14 mục</b> chuẩn CTGDPT 2018 + CV 2345</span>
          <span style="color:#94a3b8">·</span>
          <span><b>Phí:</b> {point_cost} điểm/lượt</span>
          <span style="margin-left:auto;color:#64748b;font-size:12px;">v{APP_VERSION}</span>
        </div>""",
        unsafe_allow_html=True,
    )

    # Năm học giờ auto-suy từ ngày dạy → không cần list dropdown nữa.

    # Map lớp → cấp học (tự suy ra, không cần ô riêng)
    grade_to_level = {g: lv for lv, info in SUBJECTS_BY_LEVEL.items() for g in info["grades"]}
    all_grades_flat = [g for lv in SUBJECTS_BY_LEVEL for g in SUBJECTS_BY_LEVEL[lv]["grades"]]

    # =========================================================================
    # FORM RÚT GỌN — chỉ giữ field thực sự xuất hiện trong file KHBD chuẩn.
    # Đã bỏ 11 field thừa: Tổ/khối CM, Lớp cụ thể (4A), Địa danh, Năm học (auto),
    # Mức NL số, Mức phân hóa, Ghi chú riêng, Link chính thống, Checkbox metadata,
    # nút Demo và nút Xóa. Default an toàn cho các biến downstream còn dùng.
    # =========================================================================
    with st.form(key=_k("form"), clear_on_submit=False):

        # === Phần 1: Cốt lõi — luôn hiện (chỉ 4 dropdown + 1 ô) ===
        lesson_title = st.text_input(
            "🎯 Tên bài học",
            placeholder="VD: Ôn tập các số đến 100 000",
            key=_k("title"),
        )

        c1, c2, c3, c4 = st.columns([0.7, 1.8, 1.0, 1.1])
        with c1:
            grade = str(st.selectbox("Lớp", all_grades_flat, key=_k("grade")) or all_grades_flat[0])
            level = grade_to_level[grade]
        with c2:
            subject = str(
                st.selectbox("Môn học", SUBJECTS_BY_LEVEL[level]["subjects"], key=_k("subject"))
                or SUBJECTS_BY_LEVEL[level]["subjects"][0]
            )
        with c3:
            duration = st.number_input(
                "Thời lượng (phút)", min_value=20, max_value=120,
                value=int(SUBJECTS_BY_LEVEL[level]["default_duration"]),
                step=5, key=_k("duration"),
            )
        with c4:
            teaching_date = st.date_input("Ngày dạy", value=dt.date.today(), key=_k("date"))

        # === Phần 2: Thông tin in nộp tổ (mặc định ẩn — chỉ mở khi cần) ===
        with st.expander("📋 Thông tin in nộp tổ chuyên môn (tuỳ chọn)", expanded=False):
            c1, c2 = st.columns(2)
            with c1:
                school = st.text_input(
                    "Tên trường", placeholder="VD: Tiểu học Hồng Thái", key=_k("school"),
                )
                phong_gd_dt = st.text_input(
                    "UBND / Phòng GD&ĐT",
                    placeholder="VD: Huyện Hoài Đức",
                    key=_k("phong"),
                )
            with c2:
                teacher = st.text_input("Họ tên giáo viên", key=_k("teacher"))
                period_note = st.text_input(
                    "Tuần / Tiết PPCT", placeholder="VD: Tuần 1 - Tiết 1", key=_k("period"),
                )
            sgk_pages = st.text_input(
                "Trang SGK (nếu có)", placeholder="VD: Trang 6 - 9", key=_k("sgk_pages"),
            )

        # === Phần 3: Đính kèm tài liệu SGK (giúp AI bám sát nội dung) ===
        with st.expander("📚 Đính kèm tài liệu SGK (giúp AI bám sát, tuỳ chọn)", expanded=False):
            image_files = st.file_uploader(
                "Ảnh trang SGK (JPG/PNG/WEBP)",
                type=["jpg", "jpeg", "png", "webp"],
                accept_multiple_files=True, key=_k("images"),
            )
            uploaded_files = st.file_uploader(
                "File bài học (PDF/DOCX/TXT)",
                type=["pdf", "docx", "doc", "txt"],
                accept_multiple_files=True, key=_k("files"),
            )

        # === Default ngầm cho các biến downstream còn dùng (không hiện UI) ===
        # Năm học auto-suy từ ngày dạy: tháng 1-7 → "Y-1/Y", tháng 8-12 → "Y/Y+1".
        _y = teaching_date.year
        _school_year_start = _y - 1 if teaching_date.month <= 7 else _y
        # Tôn trọng ràng buộc SCHOOL_YEAR_MIN của module
        _school_year_start = max(SCHOOL_YEAR_MIN, _school_year_start)
        school_year = f"{_school_year_start}-{_school_year_start + 1}"

        # Các field đã bỏ khỏi UI — gán default an toàn để code downstream không vỡ
        department = ""           # tổ/khối chuyên môn — không in trong template
        class_name = ""           # lớp cụ thể — fallback dùng grade
        location = ""             # địa danh — không in trong template
        digital_level = "Vừa đủ, tự nhiên"
        diff_level = "Chi tiết theo 3 nhóm"
        teacher_note = ""
        official_links = ""
        use_metadata = True       # luôn ON
        book_series = BOOK_SERIES_LOCKED

        # === Nút duy nhất — gộp AI + demo, demo logic xử lý ở callback nếu cần ===
        st.markdown("")
        submit_ai = st.form_submit_button(
            f"⚡ Tạo giáo án ({point_cost} điểm)",
            type="primary", use_container_width=True,
        )
        submit_demo = False  # nút demo đã bỏ — biến giữ để khớp signature downstream

    if submit_ai or submit_demo:
        if not lesson_title.strip():
            st.error("Vui lòng nhập tên bài học.")
            st.stop()

        # === Phase B: smart validation ===
        # 1) Auto-capitalize tên bài (giữ thông tin sạch cho prompt + file)
        original_title = lesson_title.strip()
        lesson_title = smart_capitalize_title(original_title)
        if lesson_title != original_title:
            st.info(f"🔤 Đã chuẩn hoá tên bài: **{lesson_title}**")

        # 2) Phát hiện môn/lớp/tên bài lệch nhau — cảnh báo nhưng không chặn
        mismatch_msg = detect_grade_subject_mismatch(lesson_title, str(grade), subject)
        if mismatch_msg:
            confirm_key = _k("confirm_mismatch")
            st.warning("⚠️ " + mismatch_msg)
            # Yêu cầu giáo viên xác nhận thay vì tự động chạy → tiết kiệm điểm AI
            already_confirmed = st.session_state.get(confirm_key) == lesson_title.lower()
            if not already_confirmed:
                if st.button(
                    f"✅ Tôi đã kiểm tra, đúng là {subject} lớp {grade} — tiếp tục",
                    key=_k("btn_confirm_mismatch"), type="primary",
                ):
                    st.session_state[confirm_key] = lesson_title.lower()
                    st.rerun()
                st.info("👆 Nhấn nút trên để xác nhận trước khi tốn điểm AI.")
                st.stop()

        # Kiểm tra điểm trước khi gọi AI
        if submit_ai and point_check is not None:
            if not point_check(point_cost, "soạn giáo án AI nâng cao"):
                st.stop()

        with st.spinner("Đang đọc nguồn học liệu và xác định bài học..."):
            meta = _search_metadata(grade, subject, lesson_title) if use_metadata else []
            upl_text, upl_notes = _extract_text_from_uploads(uploaded_files)
            link_text, link_notes = _fetch_official_links(official_links)
            img_text, img_notes = ("", [])
            if image_files:
                img_text, img_notes = _extract_image_context(image_files, api_key, model_name)

            bundle = _SourceBundle(
                metadata_matches=meta,
                uploaded_text=upl_text,
                uploaded_text_notes=upl_notes,
                image_context=img_text,
                image_notes=img_notes,
                link_context=link_text,
                link_notes=link_notes,
                teacher_note=teacher_note,
            )
            bundle = _determine_trust(bundle)

        _render_source_status(bundle)

        title_slug = f"GiaoAn_{_slugify(subject)}_lop_{_slugify(grade)}_{_slugify(lesson_title)}"

        if submit_ai:
            if not api_key:
                st.error("Chưa có Gemini API key. Bạn có thể bấm 'Tạo bản demo' để xem cấu trúc.")
                st.stop()
            prompt = _build_prompt(
                phong_gd_dt=phong_gd_dt, school=school, department=department, teacher=teacher,
                school_year=school_year, teaching_date=teaching_date, location=location,
                level=level, grade=grade, class_name=class_name,
                subject=subject, book_series=book_series, lesson_title=lesson_title.strip(),
                duration=int(duration), period_note=period_note, sgk_pages=sgk_pages,
                digital_level=digital_level, diff_level=diff_level, bundle=bundle,
            )
            with st.spinner("AI đang soạn giáo án..."):
                try:
                    raw = _call_gemini_text(api_key, prompt, model_name=model_name, temperature=0.22)
                    content_html = _clean_html(raw)
                except Exception as e:
                    st.error(f"Lỗi khi gọi AI: {e}")
                    st.stop()
        else:
            content_html = _demo_html(
                phong_gd_dt=phong_gd_dt, school=school, department=department, teacher=teacher,
                school_year=school_year, teaching_date=teaching_date, location=location,
                level=level, grade=grade, class_name=class_name,
                subject=subject, book_series=book_series, lesson_title=lesson_title.strip(),
                duration=int(duration), period_note=period_note, sgk_pages=sgk_pages,
                bundle=bundle,
            )

        # Normalizer: loại phrase cấm trước khi wrap/validate/lưu state.
        content_html, _stripped = _strip_banned_phrases(content_html)
        full_html = _wrap_a4(content_html, title_slug)
        validation = _validate(content_html, lesson_title.strip(), bundle.trust_level)
        if _stripped:
            validation["stripped_phrases"] = _stripped

        st.session_state[_k("html_content")] = content_html
        st.session_state[_k("full_html")] = full_html
        st.session_state[_k("title_slug")] = title_slug
        st.session_state[_k("validation")] = json.dumps(validation, ensure_ascii=False)
        if _stripped:
            st.info(f"🧹 Đã tự động làm sạch {len(_stripped)} cụm từ không phù hợp với KHBD nộp tổ: {', '.join(_stripped[:5])}")
        st.success("✅ Đã tạo giáo án. Kiểm tra bên dưới.")

    full_html = st.session_state.get(_k("full_html"), "")
    content_html = st.session_state.get(_k("html_content"), "")
    title_slug = st.session_state.get(_k("title_slug"), "GiaoAn")

    if full_html and content_html:
        st.markdown("## ✅ Kiểm định chất lượng")
        validation = json.loads(st.session_state.get(_k("validation"), "{}") or "{}")
        v1, v2 = st.columns(2)
        with v1:
            st.metric("Điểm kiểm định", f"{validation.get('score', 0)}/100")
        with v2:
            st.metric("Mức đánh giá", validation.get("level", "—"))
        with st.expander("Xem chi tiết tiêu chí", expanded=False):
            for name, ok, note in validation.get("checks", []):
                st.write(("✅" if ok else "⚠️") + f" {name}: {note}")

        st.markdown("## 📄 Xem trước giáo án (A4)")
        components.html(full_html, height=820, scrolling=True)

        st.markdown("## ⬇️ Tải giáo án")
        d1, d2 = st.columns(2)
        with d1:
            # Chiến lược 3 lớp đảm bảo giáo viên LUÔN tải được:
            # 1) docx_renderer (truyền từ app.py — create_docx_from_html, chỉ cần python-docx)
            # 2) _create_docx nội bộ (cần python-docx + beautifulsoup4)
            # 3) _create_doc_html_fallback (.doc HTML — zero dependency)
            docx_bytes: Optional[bytes] = None
            last_err: Optional[str] = None
            if docx_renderer is not None:
                try:
                    docx_bytes = docx_renderer(full_html, title_slug)
                except Exception as e:
                    last_err = f"renderer chính: {type(e).__name__}: {e}"
            if docx_bytes is None:
                try:
                    docx_bytes = _create_docx(full_html)
                except Exception as e:
                    last_err = (last_err + " | " if last_err else "") + f"_create_docx: {e}"

            if docx_bytes is not None:
                st.download_button("Tải Word .docx", data=docx_bytes,
                                   file_name=f"{title_slug}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   type="primary", use_container_width=True, key=_k("dl_docx"))
            else:
                # Fallback cuối: .doc HTML — không cần thư viện nào
                doc_bytes = _create_doc_html_fallback(full_html, title_slug)
                st.download_button("Tải Word .doc (dự phòng)", data=doc_bytes,
                                   file_name=f"{title_slug}.doc",
                                   mime="application/msword",
                                   type="primary", use_container_width=True, key=_k("dl_doc_fb"))
                with st.expander("Vì sao là .doc thay vì .docx?", expanded=False):
                    st.caption(
                        "Môi trường chạy hiện thiếu thư viện để tạo .docx chuẩn. "
                        "App đã tự chuyển sang .doc (HTML) — Microsoft Word/WPS/LibreOffice mở được bình thường. "
                        f"Chi tiết kỹ thuật: {last_err}"
                    )
                    if _IMPORT_ERRORS:
                        st.code("\n".join(f"{k}: {v}" for k, v in _IMPORT_ERRORS.items()))
        with d2:
            st.download_button("Tải HTML", data=full_html.encode("utf-8"),
                               file_name=f"{title_slug}.html", mime="text/html",
                               use_container_width=True, key=_k("dl_html"))

        with st.expander("Xem HTML giáo án (kỹ thuật)", expanded=False):
            st.code(content_html, language="html")

        # ====================================================================
        # 🆕 MẪU CHUẨN 2026-2027 (JSON pipeline) — kiến trúc mới, code 100%
        # quyết định layout, AI chỉ tạo nội dung. Khớp PDF mẫu KHBD tiểu học.
        # ====================================================================
        st.markdown("---")
        st.markdown("### 🆕 Tải theo mẫu chuẩn 2026-2027")
        st.caption(
            "Bản render mới theo mẫu KHBD tiểu học chuẩn 2026-2027 (7 mục I-VII, "
            "bảng 3 cột GV/HS/Sản phẩm-Đánh giá, header chỉ trang 1, không có 'AIEXAM'). "
            "Code quyết định 100% bố cục — AI chỉ sinh nội dung JSON."
        )
        if st.button("⚡ Tạo & tải theo mẫu chuẩn 2026-2027",
                     key=_k("btn_v2"), use_container_width=True):
            _run_v2_pipeline(
                api_key=api_key, point_check=point_check, point_cost=point_cost,
                model_name=model_name,
                # Lấy giá trị từ session_state (form đã submit ở trên)
                subject=st.session_state.get(_k("subject"), ""),
                grade=st.session_state.get(_k("grade"), ""),
                lesson_title=st.session_state.get(_k("title"), title_slug),
                school=st.session_state.get(_k("school"), ""),
                department=st.session_state.get(_k("department"), ""),
                teacher=st.session_state.get(_k("teacher"), ""),
                school_year=st.session_state.get(_k("school_year"), "2026 - 2027"),
                teaching_date=st.session_state.get(_k("teaching_date"), ""),
                duration_minutes=int(st.session_state.get(_k("duration"), 35) or 35),
                week=st.session_state.get(_k("period_note"), ""),
                sgk_pages=st.session_state.get(_k("sgk_pages"), ""),
                source_text=content_html[:8000],  # cung cấp ngữ cảnh từ HTML đã có
            )


def _run_v2_pipeline(
    *,
    api_key: str,
    point_check: Optional[Callable[[int, str], bool]],
    point_cost: int,
    model_name: str,
    subject: str,
    grade: str,
    lesson_title: str,
    school: str,
    department: str,
    teacher: str,
    school_year: str,
    teaching_date: Any,
    duration_minutes: int,
    week: str,
    sgk_pages: str,
    source_text: str = "",
) -> None:
    """Wrapper gọi lesson_docx pipeline + nút tải. Tách riêng để dễ test."""
    try:
        from lesson_docx.pipeline import LessonContext, generate_docx_from_ai_response
    except Exception as e:
        st.error(f"❌ Không import được lesson_docx: {e}")
        return

    # Tính phí trước khi gọi AI (tránh tốn điểm khi pipeline lỗi parse)
    if point_check is not None:
        if not point_check(point_cost, "lesson_plan_advanced_v2"):
            return

    # Convert teaching_date dt → str
    if hasattr(teaching_date, "strftime"):
        teaching_date = teaching_date.strftime("%d/%m/%Y")

    ctx = LessonContext(
        subject=subject,
        grade=str(grade),
        lesson_title=lesson_title,
        textbook_series=BOOK_SERIES_LOCKED,
        duration_minutes=duration_minutes,
        school=school,
        department=department,
        teacher_name=teacher,
        school_year=school_year,
        teaching_date=str(teaching_date) if teaching_date else "",
        week=week,
        period="",
        sgk_pages=sgk_pages,
        source_text=source_text or "",
    )

    with st.spinner("Đang sinh JSON từ AI và render DOCX theo mẫu chuẩn..."):
        try:
            from .pipeline import build_prompt  # type: ignore
        except Exception:
            from lesson_docx.pipeline import build_prompt
        prompt = build_prompt(ctx)
        try:
            raw = _call_gemini_text(api_key, prompt, model_name=model_name, temperature=0.2)
        except Exception as e:
            st.error(f"❌ Lỗi gọi AI: {e}")
            return

        result = generate_docx_from_ai_response(ctx, raw)

    if not result.ok:
        st.error(f"❌ Pipeline fail tại stage '{result.stage}':")
        for err in result.errors:
            st.write(f"- {err}")
        with st.expander("Xem JSON AI trả về (kỹ thuật)", expanded=False):
            st.code(raw[:5000], language="json")
        return

    # Thành công — show download
    st.success(result.summary())
    if result.warnings:
        for w in result.warnings:
            st.caption(f"⚠️ {w}")
    if result.replaced_phrases:
        st.caption(f"🧹 Đã thay phrase cấm: {', '.join(result.replaced_phrases)}")
    if result.review_fixes:
        with st.expander(f"🔧 {len(result.review_fixes)} chỉnh sửa tự động trước khi xuất file", expanded=False):
            for fix in result.review_fixes:
                st.write(f"- {fix}")

    safe_title = re.sub(r"[^a-zA-Z0-9_\-]+", "_", lesson_title) or "GiaoAn"
    assert result.docx_bytes is not None
    st.download_button(
        "⬇️ Tải file .docx (mẫu chuẩn 2026-2027)",
        data=result.docx_bytes,
        file_name=f"{safe_title}_mau_chuan.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
        key=f"dl_v2_{safe_title}",
    )
    with st.expander("Xem JSON đã chuẩn hoá (kỹ thuật)", expanded=False):
        import json as _json
        st.code(_json.dumps(result.json_data, ensure_ascii=False, indent=2), language="json")


def _demo_html(*, phong_gd_dt, school, department, teacher, school_year, teaching_date, location,
               level, grade, class_name, subject, book_series, lesson_title, duration,
               period_note, sgk_pages, bundle: _SourceBundle) -> str:
    best = bundle.metadata_matches[0] if bundle.metadata_matches else {}
    objs = best.get("learning_objectives") or [
        "Nhận biết được nội dung trọng tâm của bài học (theo SGK KNTT).",
        "Thực hiện được nhiệm vụ học tập phù hợp với yêu cầu cần đạt.",
        "Vận dụng kiến thức/kĩ năng vào tình huống đơn giản gần gũi.",
    ]
    comps = best.get("digital_competencies") or DIGITAL_COMPETENCIES[:2]
    summary = best.get("main_content_summary") or "Chưa có tóm tắt; giáo viên cần kiểm tra nội dung trong SGK KNTT."

    # NLĐT theo môn (chuẩn CTGDPT 2018) — bỏ câu chung chung
    subject_competencies = SUBJECT_SPECIFIC_COMPETENCIES.get(subject, [])
    if subject_competencies:
        sc_html = "".join(f"<li>{html.escape(x)}</li>" for x in subject_competencies)
        sc_block = f"<ul>{sc_html}</ul><p><em>Biểu hiện cụ thể trong bài thể hiện qua các hoạt động ở mục IX.</em></p>"
    else:
        sc_block = f"<p>Năng lực đặc thù môn <strong>{html.escape(subject)}</strong> theo chuẩn đầu ra CTGDPT 2018 cấp {html.escape(level)}. <em>[GV bổ sung danh mục cụ thể theo Thông tư 32/2018/TT-BGDĐT]</em></p>"

    # 4 pha CV 2345/BGDĐT-GDTH: Mở đầu / Hình thành KT mới / Luyện tập, thực hành / Vận dụng, trải nghiệm
    t1 = max(3, round(duration * 0.10))
    # Spec: 3 hoạt động — Mở đầu / (LT-thực hành HOẶC HTKT, tuỳ kiểu bài) / Vận dụng.
    # Fallback dùng "Hình thành kiến thức mới / Luyện tập, thực hành" — phù hợp cả 2 kiểu bài.
    t2 = max(15, round(duration * 0.60))  # hoạt động giữa chiếm ~60% thời gian
    t3 = max(5, duration - t1 - t2)        # vận dụng = phần còn lại
    phases = [
        {
            "title": "1. Hoạt động Mở đầu",
            "time": t1,
            "objective": "Tạo tâm thế, kết nối kinh nghiệm sẵn có của HS với bài học mới.",
            "product": "Câu trả lời/chia sẻ ban đầu của HS được nói trước lớp.",
            "assessment": "GV quan sát thái độ tham gia, nhận xét miệng, ghi nhận câu trả lời.",
            "gv": "GV nêu tình huống/câu hỏi gắn với thực tiễn liên quan đến bài. Mời 2-3 HS chia sẻ. GV dẫn dắt vào tên bài và viết tên bài lên bảng.",
            "hs": "HS lắng nghe, suy nghĩ, xung phong trả lời. Ghi tên bài học vào vở.",
        },
        {
            "title": "2. Hoạt động Hình thành kiến thức mới / Luyện tập, thực hành",
            "time": t2,
            "objective": f"HS khám phá và tiếp thu (hoặc củng cố) nội dung trọng tâm: {html.escape(summary[:180])}",
            "product": "Bài làm trong vở / phiếu học tập đã hoàn thành; câu trả lời các câu hỏi của GV.",
            "assessment": "GV hỏi đáp trực tiếp, chấm nhanh / đối chiếu đáp án, đánh giá theo tiêu chí Đạt / Cần cố gắng.",
            "gv": "GV hướng dẫn HS đọc SGK KNTT, đặt câu hỏi gợi mở; tổ chức thảo luận cặp đôi/nhóm 4; giao 2-3 nhiệm vụ luyện tập theo mức độ tăng dần; quan sát, hỗ trợ HS gặp khó khăn; tổ chức báo cáo, chốt kiến thức bằng sơ đồ/bảng.",
            "hs": "HS đọc SGK, gạch chân từ khóa; thảo luận theo cặp/nhóm; làm bài cá nhân vào vở/phiếu; đổi vở kiểm tra chéo; đại diện nhóm trình bày; ghi kết luận vào vở.",
        },
        {
            "title": "3. Hoạt động Vận dụng, trải nghiệm",
            "time": t3,
            "objective": "HS vận dụng kiến thức vào tình huống thực tiễn gần gũi và chuẩn bị bài sau.",
            "product": "Lời giải / ý tưởng vận dụng do HS đề xuất; ghi chú nhiệm vụ về nhà.",
            "assessment": "GV nhận xét sự sáng tạo, tính phù hợp; HS tự đánh giá vào phiếu cuối bài.",
            "gv": "GV đưa 1 tình huống thực tiễn gắn với bài học, yêu cầu HS đề xuất cách giải quyết. Chốt bài học. Giao nhiệm vụ về nhà & chuẩn bị bài tiếp theo.",
            "hs": "HS suy nghĩ, trình bày cách vận dụng. Tự đánh giá vào phiếu cuối bài. Ghi nhiệm vụ về nhà vào vở.",
        },
    ]

    def render_phase(p):
        return f"""
    <h3>{html.escape(p['title'])} <span style="font-weight:normal;font-size:.9em;color:#475569">(khoảng {p['time']} phút)</span></h3>
    <p><strong>Mục tiêu:</strong> {p['objective']}</p>
    <table>
      <tr>
        <th style="width:30%">Hoạt động của giáo viên</th>
        <th style="width:30%">Hoạt động của học sinh</th>
        <th style="width:20%">Sản phẩm học tập</th>
        <th style="width:20%">Đánh giá</th>
      </tr>
      <tr>
        <td>{p['gv']}</td>
        <td>{p['hs']}</td>
        <td>{p['product']}</td>
        <td>{p['assessment']}</td>
      </tr>
    </table>"""

    phases_html = "".join(render_phase(p) for p in phases)
    obj_html = "".join(f"<li>{html.escape(x)}</li>" for x in objs)
    comp_html = "".join(f"<li>{html.escape(x)}</li>" for x in comps)
    total_time = t1 + t2 + t3

    enforced_book = BOOK_SERIES_LOCKED
    enforced_year = school_year if school_year and school_year >= SCHOOL_YEAR_DEFAULT else SCHOOL_YEAR_DEFAULT
    day_str = teaching_date.strftime("%d")
    month_str = teaching_date.strftime("%m")
    year_str = teaching_date.strftime("%Y")
    # Khung 4 mục La Mã chuẩn — fallback dùng khi AI lỗi hoặc demo.
    return f"""
<article class="lesson-plan">
  <table class="hdr" style="border:none;margin-bottom:8px">
    <tr style="border:none">
      <td style="border:none;width:50%;text-align:center;font-weight:700">
        PHÒNG GD&amp;ĐT: {html.escape(phong_gd_dt or '...')}<br>
        TRƯỜNG: <span style="text-transform:uppercase">{html.escape(school or '...')}</span>
      </td>
      <td style="border:none;width:50%;text-align:center;font-weight:700">
        CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM<br>
        <span style="text-decoration:underline">Độc lập – Tự do – Hạnh phúc</span>
      </td>
    </tr>
  </table>

  <h1>KẾ HOẠCH BÀI DẠY</h1>

  <p><strong>Môn học/HĐGD:</strong> {html.escape(subject)}</p>
  <p><strong>Lớp:</strong> {html.escape(class_name or grade)}</p>
  <p><strong>Tên bài học:</strong> {html.escape(lesson_title)}</p>
  <p><strong>Thời lượng:</strong> {duration} phút</p>
  <p><strong>Tuần / Tiết PPCT:</strong> {html.escape(period_note or '...')}</p>
  <p><strong>Ngày dạy:</strong> {teaching_date.strftime('%d/%m/%Y')}</p>
  <p><strong>Giáo viên:</strong> {html.escape(teacher or '...')}</p>
  <p><strong>Bộ sách:</strong> {html.escape(enforced_book)} — {PUBLISHER}</p>

  <section><h2>I. YÊU CẦU CẦN ĐẠT</h2>
    <h3>1. Kiến thức, kĩ năng</h3>
    <p><em>Sau bài học, học sinh:</em></p>
    <ul>{obj_html}</ul>

    <h3>2. Năng lực</h3>
    <p><strong>a) Năng lực chung:</strong></p>
    <ul>
      <li><strong>Tự chủ và tự học:</strong> HS chủ động đọc SGK, tự hoàn thành nhiệm vụ học tập theo hướng dẫn.</li>
      <li><strong>Giao tiếp và hợp tác:</strong> HS trao đổi với bạn, trình bày ý kiến rõ ràng trong hoạt động nhóm/cặp.</li>
    </ul>
    <p><strong>b) Năng lực đặc thù môn {html.escape(subject)}:</strong></p>
    {sc_block}

    <h3>3. Phẩm chất</h3>
    <ul>
      <li><strong>Chăm chỉ:</strong> tích cực tham gia hoạt động học tập, hoàn thành nhiệm vụ được giao.</li>
      <li><strong>Trách nhiệm:</strong> giữ gìn đồ dùng học tập, hoàn thành phần việc trong nhóm.</li>
    </ul>
  </section>

  <section><h2>II. ĐỒ DÙNG DẠY HỌC</h2>
    <h3>1. Giáo viên</h3>
    <ul>
      <li>SGK {html.escape(enforced_book)} — môn {html.escape(subject)}, lớp {html.escape(class_name or grade)}.</li>
      <li>Phiếu học tập, bảng phụ, đồ dùng trực quan phù hợp bài.</li>
      <li>Máy chiếu/tivi (nếu có).</li>
    </ul>
    <h3>2. Học sinh</h3>
    <ul>
      <li>SGK, vở ghi, đồ dùng học tập.</li>
      <li>Đồ dùng được phân công chuẩn bị trước (nếu có).</li>
    </ul>
  </section>

  <section><h2>III. CÁC HOẠT ĐỘNG DẠY HỌC CHỦ YẾU</h2>
    <p><em>Tổng thời gian dự kiến: <strong>{total_time} phút</strong>.</em></p>
    {phases_html}
  </section>

  <section><h2>IV. ĐIỀU CHỈNH SAU BÀI DẠY</h2>
    <p>.....................................................................................................................</p>
    <p>.....................................................................................................................</p>
    <p>.....................................................................................................................</p>
  </section>

  <p style="text-align:right;margin-top:14px;font-style:italic">
    {html.escape(location or '...')}, ngày {day_str} tháng {month_str} năm {year_str}
  </p>
  <table style="border:none;margin-top:6px">
    <tr style="border:none">
      <td style="border:none;width:50%;text-align:center;font-weight:700;text-transform:uppercase">Tổ trưởng chuyên môn</td>
      <td style="border:none;width:50%;text-align:center;font-weight:700;text-transform:uppercase">Người soạn</td>
    </tr>
    <tr style="border:none">
      <td style="border:none;text-align:center;font-style:italic">(Kí, ghi rõ họ tên)</td>
      <td style="border:none;text-align:center;font-style:italic">(Kí, ghi rõ họ tên)</td>
    </tr>
    <tr style="border:none">
      <td style="border:none;height:60px"></td>
      <td style="border:none;height:60px"></td>
    </tr>
    <tr style="border:none">
      <td style="border:none;text-align:center;font-weight:700">&nbsp;</td>
      <td style="border:none;text-align:center;font-weight:700">{html.escape(teacher or '')}</td>
    </tr>
  </table>
</article>
""".strip()
